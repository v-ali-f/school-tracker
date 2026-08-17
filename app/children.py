INCIDENT_CATEGORIES = [
    "Драка/конфликт",
    "Нарушение дисциплины",
    "Трудности в обучении",
    "Буллинг",
    "Травма/вызов скорой",
    "Жалоба родителей",
    "Психологическая проблема",
    "Другое",
]

# Легаси-коды (старые записи в БД с латинскими тех-кодами) → русские лейблы.
INCIDENT_CATEGORY_LEGACY = {
    "conflict": "Драка/конфликт",
    "discipline": "Нарушение дисциплины",
    "absence": "Пропуски/опоздания",
    "bullying": "Буллинг",
    "trauma": "Травма/вызов скорой",
    "complaint": "Жалоба родителей",
    "psych": "Психологическая проблема",
    "property_damage": "Порча имущества",
    "other": "Другое",
}


def _category_label(value):
    """Русский лейбл для категории. Возвращает '—' для пустых, маппит легаси-коды."""
    if not value:
        return "Без категории"
    return INCIDENT_CATEGORY_LEGACY.get(value, value)
OVZ_LEVEL_LABELS = {
    "NOO": "Начальное общее образование",
    "OOO": "Основное общее образование",
    "SOO": "Среднее общее образование",
}

CLASS_REGISTRY_LEVEL_OPTIONS = (
    ("NOO", "НОО · 1–4 классы", tuple(range(1, 5))),
    ("OOO", "ООО · 5–9 классы", tuple(range(5, 10))),
    ("SOO", "СОО · 10–11 классы", (10, 11)),
)
CLASS_REGISTRY_LEVEL_GRADES = {
    code: set(grades)
    for code, _label, grades in CLASS_REGISTRY_LEVEL_OPTIONS
}

OVZ_NOZOLOGY_LABELS = {
    "VISION": "Нарушения зрения",
    "TNR": "Тяжёлые нарушения речи",
    "NODA": "Нарушения опорно-двигательного аппарата",
    "ZPR": "Задержка психического развития",
    "INT": "Интеллектуальные нарушения",
}
AOOP_FULL_NAMES = {
    "4.1": "АООП НОО для обучающихся с нарушениями зрения (вариант 4.1)",
    "4.2": "АООП НОО для обучающихся с нарушениями зрения (вариант 4.2)",
    "4.3": "АООП НОО для обучающихся с нарушениями зрения (вариант 4.3)",

    "5.1": "АООП НОО для обучающихся с тяжёлыми нарушениями речи (вариант 5.1)",
    "5.2": "АООП НОО для обучающихся с тяжёлыми нарушениями речи (вариант 5.2)",

    "6.1": "АООП НОО для обучающихся с нарушениями опорно-двигательного аппарата (вариант 6.1)",
    "6.2": "АООП НОО для обучающихся с нарушениями опорно-двигательного аппарата (вариант 6.2)",
    "6.3": "АООП НОО для обучающихся с нарушениями опорно-двигательного аппарата (вариант 6.3)",
    "6.4": "АООП НОО для обучающихся с нарушениями опорно-двигательного аппарата (вариант 6.4)",

    "7.1": "АООП НОО для обучающихся с задержкой психического развития (вариант 7.1)",
    "7.2": "АООП НОО для обучающихся с задержкой психического развития (вариант 7.2)",

    "8.1": "АООП НОО для обучающихся с интеллектуальными нарушениями (вариант 8.1)",
    "8.2": "АООП НОО для обучающихся с интеллектуальными нарушениями (вариант 8.2)",
    "8.3": "АООП НОО для обучающихся с интеллектуальными нарушениями (вариант 8.3)",
    "8.4": "АООП НОО для обучающихся с интеллектуальными нарушениями (вариант 8.4)",
}

AOOP_TO_OVZ = {
    "4": {"nosology": "VISION", "label": "Нарушения зрения"},
    "5": {"nosology": "TNR", "label": "Тяжёлые нарушения речи"},
    "6": {"nosology": "NODA", "label": "Нарушения опорно-двигательного аппарата"},
    "7": {"nosology": "ZPR", "label": "Задержка психического развития"},
    "8": {"nosology": "INT", "label": "Интеллектуальные нарушения"},
}
from datetime import datetime, date, timedelta, timezone as _dt_tz
try:
    from zoneinfo import ZoneInfo as _MSK_ZI
    _MSK_TZ_INC = _MSK_ZI("Europe/Moscow")
except Exception:
    _MSK_TZ_INC = _dt_tz(timedelta(hours=3))

def _now_msk_naive():
    return datetime.now(_MSK_TZ_INC).replace(tzinfo=None)
import os
import shutil
import re
import mimetypes
from uuid import uuid4
from html import escape

from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    current_app,
    flash,
    abort,
    jsonify,
    send_file,
    session,
    g,
)
from flask_login import login_required, current_user
from sqlalchemy import func
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import aliased, joinedload, subqueryload, contains_eager, selectinload
from openpyxl import load_workbook, Workbook
from io import BytesIO

from app.core.extensions import db
from app.core.cache import cache, view_response_cache, make_key
from app.core.pagination import paginate_list, resolve_pagination, SimplePagination
from app.services.education_activity_service import (
    assign_subject_activity,
    get_or_create_subject_activity,
    get_subject_activity,
    list_subject_activities,
)
from .models import (
    AcademicYear,
    Building,
    User,
    Role,
    UserRole,
    SchoolClass,
    Child,
    ChildEnrollment,
    Parent,
    ChildParent,
    ChildSocial,
    Subject,
    Debt,
    Document,
    ChildComment,
    ChildEvent,
    ChildTransferHistory,
    ChildMovement,
    Incident,
    IncidentChild,
    IncidentNote,
    IncidentNoteAttachment,
    ControlWorkResult,
    OlympiadResult,
    Task,
    EducationPlan,
    EducationPlanBinding,
    PopulationSnapshotClass,
    TariffCycle,
    TariffVersion,
)
from app.services.teaching_group_service import current_population_snapshot
from app.utils.building_matrix_tones import (
    BUILDING_MATRIX_TONE_CHOICES,
    building_matrix_tone,
    normalize_building_matrix_tone,
)
from .ovz_rules import OVZ_LEVELS, OVZ_NOZOLOGIES, allowed_variants, is_allowed
from .roles import require_roles
from app.service_staff import child_support_assignments, can_view_child_service_block
from .permissions import (
    can_view_child_basic,
    build_child_card_flags,
    should_limit_children_to_own_class,
    has_permission,
    has_role,
    can_view_documents,
    can_upload_documents,
    is_admin,
    can_edit_social_passport,
)

children_bp = Blueprint("children", __name__)


# =========================================================
# HELPERS
# =========================================================
def as_checkbox(form, name: str) -> bool:
    vals = form.getlist(name)
    return ("1" in vals) or ("on" in vals) or ("true" in vals) or ("True" in vals)


def parse_date(s: str):
    s = (s or "").strip()
    if not s:
        return None
    try:
        return datetime.strptime(s, "%Y-%m-%d").date()
    except Exception:
        return None




def _document_abs_path(stored_path: str) -> str:
    if not stored_path:
        return ""
    if os.path.isabs(stored_path):
        return stored_path
    upload_root = current_app.config.get("UPLOAD_FOLDER") or os.path.abspath(os.path.join("data", "uploads"))
    return os.path.join(os.path.abspath(upload_root), stored_path)


def _user_can_manage_document(child) -> bool:
    return is_admin() or can_upload_documents(child)


def _render_docx_preview(path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    parts = ['<div class="container-fluid py-3">']
    for p in doc.paragraphs:
        text = (p.text or '').strip()
        if text:
            parts.append(f'<p style="margin-bottom:.5rem; white-space:pre-wrap;">{escape(text)}</p>')
    for table in doc.tables:
        parts.append('<div class="table-responsive"><table class="table table-sm table-bordered">')
        for row in table.rows:
            parts.append('<tr>')
            for cell in row.cells:
                parts.append(f'<td>{escape((cell.text or "").strip())}</td>')
            parts.append('</tr>')
        parts.append('</table></div>')
    if len(parts) == 1:
        parts.append('<div class="text-muted">В документе нет читаемого текста.</div>')
    parts.append('</div>')
    return ''.join(parts)


def _render_xlsx_preview(path: str) -> str:
    wb = load_workbook(path, data_only=True)
    parts = ['<div class="container-fluid py-3">']
    for ws in wb.worksheets:
        parts.append(f'<h6 class="mt-2">{escape(ws.title)}</h6>')
        parts.append('<div class="table-responsive mb-3"><table class="table table-sm table-bordered">')
        max_row = min(ws.max_row or 0, 50)
        max_col = min(ws.max_column or 0, 12)
        if max_row == 0 or max_col == 0:
            parts.append('<tr><td class="text-muted">Пустой лист</td></tr>')
        else:
            for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col, values_only=True):
                parts.append('<tr>')
                for val in row:
                    txt = '' if val is None else str(val)
                    parts.append(f'<td>{escape(txt)}</td>')
                parts.append('</tr>')
        parts.append('</table></div>')
        if (ws.max_row or 0) > 50 or (ws.max_column or 0) > 12:
            parts.append('<div class="small text-muted mb-3">Показана только часть таблицы для предпросмотра.</div>')
    parts.append('</div>')
    return ''.join(parts)


def _render_text_preview(path: str) -> str:
    for enc in ('utf-8', 'cp1251', 'latin-1'):
        try:
            with open(path, 'r', encoding=enc) as f:
                text = f.read()
            break
        except Exception:
            text = None
    if text is None:
        text = 'Не удалось прочитать текст документа.'
    return f'<div class="container-fluid py-3"><pre style="white-space:pre-wrap;">{escape(text[:200000])}</pre></div>'

def _registry_filter_state(year=None, allow_only_own_class: bool = False):
    q_text = (request.args.get("q") or "").strip()
    selected_grade_raw = (request.args.get("grade") or "").strip()
    selected_class_id = request.args.get("class_id", type=int)

    selected_grade = None
    if selected_grade_raw:
        try:
            selected_grade = int(selected_grade_raw)
        except ValueError:
            selected_grade = None

    classes_query = SchoolClass.query
    if year:
        classes_query = classes_query.filter(SchoolClass.academic_year_id == year.id)
    if allow_only_own_class:
        classes_query = classes_query.filter(SchoolClass.teacher_user_id == current_user.id)

    all_classes = (
        classes_query
        .order_by(
            SchoolClass.grade.asc().nullslast(),
            SchoolClass.letter.asc().nullslast(),
            SchoolClass.name.asc(),
        )
        .all()
    )

    grades = []
    for c in all_classes:
        if c.grade is not None and c.grade not in grades:
            grades.append(c.grade)

    classes = [c for c in all_classes if selected_grade is None or c.grade == selected_grade]

    selected_class = None
    if selected_class_id:
        for c in all_classes:
            if c.id == selected_class_id:
                selected_class = c
                break

    if selected_class and selected_grade is None and selected_class.grade is not None:
        selected_grade = selected_class.grade
        classes = [c for c in all_classes if c.grade == selected_grade]

    selected_class_name = selected_class.name if selected_class else ""

    return {
        "q_text": q_text,
        "selected_grade": selected_grade,
        "selected_grade_raw": str(selected_grade) if selected_grade is not None else "",
        "selected_class_id": selected_class_id,
        "selected_class_name": selected_class_name,
        "classes": classes,
        "grades": grades,
    }


def _match_fio_query(child: Child, q_text: str) -> bool:
    if not q_text:
        return True
    hay = " ".join([
        child.last_name or "",
        child.first_name or "",
        child.middle_name or "",
        child.current_class_name or "",
    ]).lower()
    return q_text.lower() in hay


def parse_int(s: str):
    s = (s or "").strip()
    if not s:
        return None
    try:
        return int(s)
    except Exception:
        return None


def split_class_name(raw: str):
    raw = (raw or "").strip()
    if not raw:
        return (None, None)

    m = re.match(r"^(\d{1,2})\s*(.*)$", raw)
    if not m:
        return (None, raw or None)

    grade = parse_int(m.group(1))
    letter = (m.group(2) or "").strip() or None
    return (grade, letter)

def normalize_class_name(raw: str):
    s = (raw or "").strip().upper()
    s = s.replace(" ", "")
    s = s.replace("-", "")
    return s or None


def promoted_class_identity(school_class):
    grade = school_class.grade
    letter = school_class.letter
    if grade is None:
        grade, letter = split_class_name(school_class.name)
    if grade not in {*range(1, 9), 10}:
        return None
    target_grade = grade + 1
    target_name = normalize_class_name(
        f"{target_grade}{letter or ''}"
    )
    return target_name, target_grade, letter


def _class_name_exists(academic_year_id, name, *, exclude_class_id=None):
    normalized = normalize_class_name(name)
    if not normalized:
        return False
    query = SchoolClass.query.filter(
        SchoolClass.academic_year_id == academic_year_id,
        db.func.upper(
            db.func.replace(
                db.func.replace(SchoolClass.name, " ", ""),
                "-",
                "",
            )
        ) == normalized,
    )
    if exclude_class_id is not None:
        query = query.filter(SchoolClass.id != exclude_class_id)
    return query.first() is not None

def _ensure_can_edit():
    if getattr(current_user, "role", "VIEWER") == "VIEWER":
        abort(403)


def _get_current_year():
    return AcademicYear.query.filter_by(is_current=True).first()


def _calc_retention_until(academic_year):
    if academic_year and academic_year.end_date:
        try:
            return academic_year.end_date.replace(year=academic_year.end_date.year + 7)
        except Exception:
            return None
    return None

def _sync_class_teacher_role(user_id):
    if not user_id:
        return

    user = User.query.get(user_id)
    if not user:
        return

    role = Role.query.filter_by(code="CLASS_TEACHER").first()
    if not role:
        return

    has_any_class = (
        SchoolClass.query
        .filter(SchoolClass.teacher_user_id == user.id)
        .first()
        is not None
    )

    existing_link = UserRole.query.filter_by(user_id=user.id, role_id=role.id).first()

    if has_any_class and not existing_link:
        db.session.add(UserRole(user_id=user.id, role_id=role.id))

    if not has_any_class and existing_link:
        db.session.delete(existing_link)

def _get_class_teacher_id(child: Child):
    if child.current_class:
        return child.current_class.teacher_user_id
    return None


def _can_edit_profile_admin_only(child: Child) -> bool:
    return getattr(current_user, "role", None) == "ADMIN"


def _can_edit_social_passport(child: Child) -> bool:
    return can_edit_social_passport(child)


def _sync_child_az_flag(child: Child):
    has_open = (
        Debt.query
        .filter_by(child_id=child.id, status="OPEN")
        .first()
        is not None
    )
    child.is_az = bool(has_open)


def _get_or_create_social(child: Child) -> ChildSocial:
    if child.social:
        return child.social

    social = ChildSocial(child_id=child.id)
    db.session.add(social)
    db.session.flush()
    return social


def _get_parent_by_relation(child: Child, relation_type: str):
    for link in (child.parent_links or []):
        if link.relation_type == relation_type:
            return link.parent
    return None


def _set_parent_relation(child: Child, relation_type: str, fio: str, phone: str):
    fio = (fio or "").strip()
    phone = (phone or "").strip()

    existing_link = None
    for link in (child.parent_links or []):
        if link.relation_type == relation_type:
            existing_link = link
            break

    if not fio and not phone:
        if existing_link:
            db.session.delete(existing_link)
        return

    if existing_link and existing_link.parent:
        existing_link.parent.fio = fio or existing_link.parent.fio
        existing_link.parent.phone = phone or None
        return

    parent = Parent(
        fio=fio or relation_type,
        phone=phone or None,
    )
    db.session.add(parent)
    db.session.flush()

    link = ChildParent(
        child_id=child.id,
        parent_id=parent.id,
        relation_type=relation_type,
        is_legal_representative=True,
    )
    db.session.add(link)


def _export_children_xlsx(title: str, children):
    wb = Workbook()
    ws = wb.active
    ws.title = "Реестр"

    ws.append(["№", "ФИО", "Класс", "Дата рождения", "Мама", "Телефон мамы", "Папа", "Телефон папы"])

    for idx, ch in enumerate(children, start=1):
        mother = _get_parent_by_relation(ch, "mother")
        father = _get_parent_by_relation(ch, "father")

        ws.append([
            idx,
            ch.fio,
            ch.current_class_name or "—",
            ch.birth_date.strftime("%d.%m.%Y") if ch.birth_date else "",
            mother.fio if mother else "",
            mother.phone if mother else "",
            father.fio if father else "",
            father.phone if father else "",
        ])

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    safe_name = re.sub(r"[^0-9A-Za-zА-Яа-я_\- ]+", "", title).strip().replace(" ", "_")
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"{safe_name}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


def _children_base_query_for_current_year():
    year = _get_current_year()

    q = db.session.query(Child).filter(
        Child.status == "ACTIVE",
        ~_active_expel_exists(Child.id),
    )

    if year:
        q = (
            q.outerjoin(
                ChildEnrollment,
                (ChildEnrollment.child_id == Child.id)
                & (ChildEnrollment.academic_year_id == year.id)
                & (ChildEnrollment.status == "ACTIVE")
                & (ChildEnrollment.ended_at.is_(None))
            )
            .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        )
    else:
        q = (
            q.outerjoin(
                ChildEnrollment,
                (ChildEnrollment.child_id == Child.id)
                & (ChildEnrollment.status == "ACTIVE")
                & (ChildEnrollment.ended_at.is_(None))
            )
            .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        )

    return q, year


def _matching_transfer_history_for_enrollment(enrollment):
    """Return the active transfer record that created an enrollment."""
    if enrollment is None or enrollment.enrolled_at is None:
        return None
    candidates = (
        ChildTransferHistory.query
        .filter(
            ChildTransferHistory.child_id == enrollment.child_id,
            ChildTransferHistory.to_academic_year_id
            == enrollment.academic_year_id,
            ChildTransferHistory.to_class_id
            == enrollment.school_class_id,
            ChildTransferHistory.reversed_at.is_(None),
        )
        .order_by(
            ChildTransferHistory.created_at.desc(),
            ChildTransferHistory.id.desc(),
        )
        .all()
    )
    for history in candidates:
        if history.created_at is None:
            continue
        if abs(
            (history.created_at - enrollment.enrolled_at).total_seconds()
        ) <= 10 * 60:
            return history
    return None


def _source_enrollment_for_undo(target_enrollment, history=None):
    if target_enrollment is None:
        return None
    query = ChildEnrollment.query.filter(
        ChildEnrollment.child_id == target_enrollment.child_id,
        ChildEnrollment.id != target_enrollment.id,
        ChildEnrollment.status != "CANCELLED",
    )
    if history is not None:
        if (
            history.from_academic_year_id is None
            or history.from_class_id is None
        ):
            return None
        query = query.filter(
            ChildEnrollment.academic_year_id
            == history.from_academic_year_id,
            ChildEnrollment.school_class_id == history.from_class_id,
        )
    else:
        # Older manual changes from the child card did not have a transfer
        # history row and always replaced a class in the same academic year.
        query = query.filter(
            ChildEnrollment.academic_year_id
            == target_enrollment.academic_year_id,
        )
    if target_enrollment.enrolled_at is not None:
        query = query.filter(
            ChildEnrollment.enrolled_at
            <= target_enrollment.enrolled_at,
        )
    return query.order_by(
        ChildEnrollment.enrolled_at.desc(),
        ChildEnrollment.id.desc(),
    ).first()


def _latest_reversible_enrollment(child_id):
    target = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.child_id == child_id,
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.status != "CANCELLED",
        )
        .order_by(
            ChildEnrollment.enrolled_at.desc(),
            ChildEnrollment.id.desc(),
        )
        .first()
    )
    if target is None:
        return None
    history = _matching_transfer_history_for_enrollment(target)
    source = _source_enrollment_for_undo(target, history)
    if source is None:
        return None
    return {
        "target": target,
        "source": source,
        "history": history,
    }


def _active_expel_event(child_id):
    """Return the last non-reversed expulsion event for a child."""
    latest = (
        ChildEvent.query
        .filter(
            ChildEvent.child_id == child_id,
            ChildEvent.event_type.in_(("EXPEL", "EXPEL_UNDO")),
            db.or_(
                ChildEvent.promotion_kind.is_(None),
                ChildEvent.promotion_kind != "ARCHIVED",
            ),
        )
        .order_by(
            ChildEvent.created_at.desc(),
            ChildEvent.id.desc(),
        )
        .first()
    )
    return latest if latest and latest.event_type == "EXPEL" else None


def _expelled_enrollments_for_event(event):
    if event is None or event.created_at is None:
        return []
    window = timedelta(minutes=10)
    return (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.child_id == event.child_id,
            ChildEnrollment.status == "EXPELLED",
            ChildEnrollment.ended_at.isnot(None),
            ChildEnrollment.ended_at >= event.created_at - window,
            ChildEnrollment.ended_at <= event.created_at + window,
        )
        .order_by(
            ChildEnrollment.enrolled_at.desc(),
            ChildEnrollment.id.desc(),
        )
        .all()
    )


def _active_expel_events_query():
    """Base query containing only expellings that were not later undone."""
    return (
        ChildEvent.query
        .join(Child, ChildEvent.child_id == Child.id)
        .filter(
            ChildEvent.event_type == "EXPEL",
            db.or_(
                ChildEvent.promotion_kind.is_(None),
                ChildEvent.promotion_kind != "ARCHIVED",
            ),
            ~_later_expel_undo_exists(
                ChildEvent.child_id,
                ChildEvent.created_at,
                ChildEvent.id,
            ),
        )
    )


def _later_expel_undo_exists(child_id, created_at, event_id):
    undo_event = aliased(ChildEvent)
    return (
        db.session.query(undo_event.id)
        .filter(
            undo_event.child_id == child_id,
            undo_event.event_type == "EXPEL_UNDO",
            db.or_(
                undo_event.created_at > created_at,
                db.and_(
                    undo_event.created_at == created_at,
                    undo_event.id > event_id,
                ),
            ),
        )
        .exists()
    )


def _active_expel_exists(child_id):
    expel_event = aliased(ChildEvent)
    return (
        db.session.query(expel_event.id)
        .filter(
            expel_event.child_id == child_id,
            expel_event.event_type == "EXPEL",
            db.or_(
                expel_event.promotion_kind.is_(None),
                expel_event.promotion_kind != "ARCHIVED",
            ),
            ~_later_expel_undo_exists(
                expel_event.child_id,
                expel_event.created_at,
                expel_event.id,
            ),
        )
        .exists()
    )


def _active_class_enrollment_filters(academic_year_id):
    """Canonical filters for pupils who belong to the live contingent."""
    return (
        ChildEnrollment.academic_year_id == academic_year_id,
        ChildEnrollment.status == "ACTIVE",
        ChildEnrollment.ended_at.is_(None),
        Child.status == "ACTIVE",
        ~_active_expel_exists(Child.id),
    )

def parse_aoop_variant(raw_value: str):
    text = (raw_value or "").strip().upper()
    if not text:
        return None

    # Ищем вариант вроде 7.1, 8.2, 6.4
    m = re.search(r'(\d\.\d)', text)
    if m:
        return m.group(1)

    # Ищем одиночные коды вроде "12"
    m2 = re.search(r'\b(\d{1,2})\b', text)
    if m2:
        return m2.group(1)

    return text

def apply_aoop_to_child(child, social, aoop_raw: str):
    variant_code = parse_aoop_variant(aoop_raw)
    if not variant_code:
        return

    social.aoop_variant_text = AOOP_FULL_NAMES.get(variant_code, str(aoop_raw).strip())

    if variant_code not in AOOP_FULL_NAMES:
        return

    group_code = variant_code.split(".")[0]
    if group_code not in AOOP_TO_OVZ:
        return

    child.is_ovz = True
    child.ovz_level = "NOO"
    child.ovz_nosology = AOOP_TO_OVZ[group_code]["nosology"]

    try:
        child.ovz_variant = int(variant_code.split(".")[1])
    except Exception:
        child.ovz_variant = None

# =========================================================
# HOME
# =========================================================
@children_bp.route("/")
@login_required
def home():
    return redirect(url_for("main.dashboard"))


# =========================================================
# CHILDREN LIST
# =========================================================
@children_bp.route("/children")
@login_required
def list_children():
    can_view_school = has_role("ADMIN") or has_role("METHODIST") or has_role("PSYCHOLOGIST") or has_role("SOCIAL_PEDAGOG")
    if not can_view_school and not has_role("CLASS_TEACHER"):
        abort(403)
    query, year = _children_base_query_for_current_year()
    own_class_only = not can_view_school
    filters = _registry_filter_state(year, allow_only_own_class=own_class_only)

    if own_class_only:
        query = query.filter(SchoolClass.teacher_user_id == current_user.id)

    if filters["selected_grade"] is not None:
        query = query.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        query = query.filter(SchoolClass.id == filters["selected_class_id"])

    if filters["q_text"]:
        like = f"%{filters['q_text']}%"
        query = query.filter(db.or_(
            Child.last_name.ilike(like),
            Child.first_name.ilike(like),
            Child.middle_name.ilike(like),
            SchoolClass.name.ilike(like),
        ))

    query = query.order_by(SchoolClass.name.asc(), Child.last_name.asc(), Child.first_name.asc())

    page, per_page = resolve_pagination()
    total = query.count()
    pagination = SimplePagination(page=page, per_page=per_page, total=total)
    if per_page == "all":
        children_page = query.all()
    else:
        offset = (pagination.page - 1) * pagination.per_page_num
        children_page = query.offset(offset).limit(pagination.per_page_num).all()

    can_add_incident = has_permission("incident_add")
    rows = [
        {
            "child": ch,
            "can_view_card": can_view_child_basic(ch),
            "can_add_incident": can_add_incident,
        }
        for ch in children_page
    ]

    return render_template(
        "children_list.html",
        children=children_page,
        rows=rows,
        pagination=pagination,
        q=filters["q_text"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        classes=filters["classes"],
        grades=filters["grades"],
        academic_year=year,
    )


_LATIN_TO_CYRILLIC_LOOKALIKES = str.maketrans(
    "ABCEHIKMNOPSTUXY",
    "АВСЕНИКМНОРСТУХУ",
)


def _normalize_class_name(s: str) -> str:
    """Нормализация названия класса для поиска: upper + убираем пробелы/дефисы +
    меняем латинские буквы на визуально похожие кириллические.
    Примеры: '6-АС' / '6 ас' / '6as' → 'АС' c префиксом '6' → '6АС'."""
    if not s:
        return ""
    s = s.upper().strip()
    for ch in (" ", "-", "_", "/", "\\", "."):
        s = s.replace(ch, "")
    return s.translate(_LATIN_TO_CYRILLIC_LOOKALIKES)


@children_bp.route("/children/search-ajax")
@login_required
def search_children_ajax():
    """AJAX-поиск учеников + классов для главной страницы. Возвращает JSON.

    Если строка запроса соответствует названию класса (с нормализацией —
    пробелы/дефисы/латиница ↔ кириллица), первым элементом идёт ссылка
    на страницу класса.
    """
    q = request.args.get("q", "").strip()
    if len(q) < 2:
        return jsonify([])
    if not (has_role("ADMIN") or has_role("METHODIST")):
        return jsonify([])

    results: list[dict] = []

    # 1. Попытка найти класс по нормализованному названию (текущий учебный год)
    year = _get_current_year()
    if year:
        norm_q = _normalize_class_name(q)
        if norm_q:
            class_candidates = (
                SchoolClass.query
                .filter(SchoolClass.academic_year_id == year.id)
                .all()
            )
            limit_to_own = should_limit_children_to_own_class()
            for sc in class_candidates:
                if _normalize_class_name(sc.name) == norm_q:
                    if limit_to_own and sc.teacher_user_id != current_user.id:
                        continue
                    results.append({
                        "type": "class",
                        "id": sc.id,
                        "name": f"Класс {sc.name}",
                        "class": "Открыть страницу класса",
                        "url": url_for("children.class_detail", class_id=sc.id),
                    })
                    if has_permission("social_passport_registry_view"):
                        results.append({
                            "type": "class_social",
                            "id": sc.id,
                            "name": f"Соц. паспорт · {sc.name}",
                            "class": "Открыть соц. паспорт класса",
                            "url": url_for("children.social_passport_registry", class_id=sc.id),
                        })
                    break

    # 2. Поиск учеников по ФИО / названию класса.
    # Любая роль с правом создавать/смотреть инциденты находит любого ученика.
    # Доступ к карточке ограничивается отдельно, через can_view_child_basic
    # (класс.рук видит карточку только своих — для чужих в результатах будет
    # только кнопка «+ инцидент»).
    query, _ = _children_base_query_for_current_year()
    like = f"%{q}%"
    query = query.filter(db.or_(
        Child.last_name.ilike(like),
        Child.first_name.ilike(like),
        Child.middle_name.ilike(like),
        SchoolClass.name.ilike(like),
    ))
    children = query.order_by(Child.last_name.asc(), Child.first_name.asc()).limit(12).all()
    can_add_incident = has_permission("incident_add")
    for child in children:
        can_view = can_view_child_basic(child)
        item = {
            "type": "child",
            "id": child.id,
            "name": f"{child.last_name} {child.first_name} {child.middle_name or ''}".strip(),
            "class": child.current_class_name or "—",
            "url": url_for("children.child_card", child_id=child.id) if can_view else None,
            "can_view_card": can_view,
            "can_add_incident": can_add_incident,
            "incident_url": url_for("children.incident_new", student_id=child.id) if can_add_incident else None,
        }
        results.append(item)
    return jsonify(results)


# =========================================================
# NEW CHILD
# =========================================================
@children_bp.route("/children/new", methods=["GET", "POST"])
@login_required
def new_child():
    if not has_permission("child_create"):
        abort(403)
    year = _get_current_year()
    school_classes = []
    if year:
        school_classes = (
            SchoolClass.query
            .filter(SchoolClass.academic_year_id == year.id)
            .order_by(SchoolClass.name.asc())
            .all()
        )

    if request.method == "POST":
        _ensure_can_edit()

        last_name = (request.form.get("last_name") or "").strip()
        first_name = (request.form.get("first_name") or "").strip()
        middle_name = (request.form.get("middle_name") or "").strip() or None
        birth_date = parse_date(request.form.get("birth_date"))
        reg_address = (request.form.get("reg_address") or "").strip() or None
        notes = (request.form.get("notes") or "").strip() or None

        if not last_name or not first_name:
            flash("Укажите фамилию и имя", "danger")
            return render_template("child_new.html", school_classes=school_classes)

        child = Child(
            last_name=last_name,
            first_name=first_name,
            middle_name=middle_name,
            birth_date=birth_date,
            reg_address=reg_address,
            notes=notes,
        )
        db.session.add(child)
        db.session.flush()

        _set_parent_relation(
            child,
            "mother",
            request.form.get("mother_fio"),
            request.form.get("mother_phone"),
        )
        _set_parent_relation(
            child,
            "father",
            request.form.get("father_fio"),
            request.form.get("father_phone"),
        )

        _get_or_create_social(child)

        school_class_id = request.form.get("school_class_id", type=int)
        if year and school_class_id:
            sc = (
                SchoolClass.query
                .filter(
                    SchoolClass.id == school_class_id,
                    SchoolClass.academic_year_id == year.id
                )
                .first()
            )
            if sc:
                en = ChildEnrollment(
                    child_id=child.id,
                    academic_year_id=year.id,
                    school_class_id=sc.id,
                    status="ACTIVE",
                )
                db.session.add(en)

        db.session.commit()
        flash("Ребёнок добавлен", "success")
        return redirect(url_for("children.child_card", child_id=child.id))

    return render_template("child_new.html", school_classes=school_classes)


# =========================================================
# CHILD CARD
# =========================================================
@children_bp.route("/children/<int:child_id>")
@login_required
def child_card(child_id: int):
    child = Child.query.get_or_404(child_id)

    if not can_view_child_basic(child):
        abort(403)

    requested_return_class_id = request.args.get("return_class_id", type=int)
    return_class = None
    if requested_return_class_id:
        return_enrollment = (
            ChildEnrollment.query
            .filter(
                ChildEnrollment.child_id == child.id,
                ChildEnrollment.school_class_id == requested_return_class_id,
                ChildEnrollment.status == "ACTIVE",
                ChildEnrollment.ended_at.is_(None),
            )
            .first()
        )
        if return_enrollment is not None:
            return_class = return_enrollment.school_class
    if return_class is None:
        return_class = child.current_class

    can_return_to_class = (
        return_class is not None
        and has_permission("children_registry_view")
        and (
            not should_limit_children_to_own_class()
            or return_class.teacher_user_id == current_user.id
        )
    )
    return_class_url = None
    return_contingent_url = None
    if can_return_to_class:
        contingent_year_id = (
            request.args.get("contingent_year_id", type=int)
            or return_class.academic_year_id
        )
        contingent_building_id = request.args.get(
            "contingent_building_id",
            type=int,
        )
        return_class_args = {
            "class_id": return_class.id,
            "contingent_year_id": contingent_year_id,
        }
        contingent_args = {"year_id": contingent_year_id}
        if contingent_building_id is not None:
            return_class_args["contingent_building_id"] = (
                contingent_building_id
            )
            contingent_args["building_id"] = contingent_building_id
        return_class_url = url_for(
            "children.class_detail",
            **return_class_args,
        )
        return_contingent_url = url_for(
            "children.contingent",
            **contingent_args,
        )

    social = child.social or _get_or_create_social(child)
    db.session.flush()

    mother = _get_parent_by_relation(child, "mother")
    father = _get_parent_by_relation(child, "father")

    selected_year_id = request.args.get("academic_year_id", type=int)
    year = _get_current_year()
    if not selected_year_id and year:
        selected_year_id = year.id

    docs = [d for d in (child.documents or []) if not getattr(d, "is_deleted_soft", False) and not getattr(d, "is_hidden_by_retention", False)]
    if selected_year_id:
        docs = [d for d in docs if (d.academic_year_id == selected_year_id or d.academic_year_id is None)]

    def dt(x):
        return (x.doc_type or "").strip().upper()

    documents_ovz = [d for d in docs if dt(d) == "OVZ"]
    documents_vshu = [d for d in docs if dt(d) == "VSHU"]
    documents_low = [d for d in docs if dt(d) == "LOW"]
    documents_az = [d for d in docs if dt(d) == "AZ"]
    documents_mse = [d for d in docs if dt(d) == "MSE"]
    documents_ipra = [d for d in docs if dt(d) == "IPRA"]
    documents_general = [d for d in docs if dt(d) == "GENERAL"]
    documents_disabled = [d for d in docs if dt(d) == "DISABLED"]

    ovz_allowed = allowed_variants(child.ovz_level, child.ovz_nosology)

    open_debts = (
        Debt.query
        .filter(Debt.child_id == child.id, Debt.status == "OPEN")
        .order_by(Debt.due_date.is_(None), Debt.due_date.asc(), Debt.created_at.desc())
        .all()
    )

    closed_debts = (
        Debt.query
        .filter(Debt.child_id == child.id, Debt.status == "CLOSED")
        .order_by(Debt.closed_at.is_(None), Debt.closed_at.desc(), Debt.created_at.desc())
        .all()
    )

    _sync_child_az_flag(child)
    db.session.commit()

    comments = (
        ChildComment.query
        .filter_by(child_id=child.id)
        .order_by(ChildComment.created_at.desc())
        .all()
    )

    events = (
        ChildEvent.query
        .filter_by(child_id=child.id)
        .order_by(ChildEvent.created_at.desc())
        .all()
    )

    incidents = (
        db.session.query(Incident)
        .join(IncidentChild, IncidentChild.incident_id == Incident.id)
        .filter(IncidentChild.child_id == child.id)
        .order_by(Incident.occurred_at.desc())
        .all()
    )

    school_classes = []
    current_year_id = selected_year_id
    if year:
        school_classes = (
            SchoolClass.query
            .filter(SchoolClass.academic_year_id == year.id)
            .order_by(SchoolClass.name.asc())
            .all()
        )

    show_ovz = child.is_ovz or bool(documents_ovz)
    show_disabled = (
        child.is_disabled
        or bool(child.disability_mse)
        or bool(child.disability_from)
        or bool(child.disability_to)
        or bool(child.disability_ipra)
        or bool(documents_mse)
        or bool(documents_ipra)
        or bool(documents_disabled)
    )
    show_vshu = child.is_vshu or bool(social.vshu_since and not social.vshu_removed_at) or bool(documents_vshu)
    show_low = child.is_low or bool(child.low_subjects) or bool(child.low_notes) or bool(documents_low)
    show_az = bool(open_debts) or bool(closed_debts) or bool(documents_az)
    show_general = bool(documents_general)

    ovz_level_label = OVZ_LEVEL_LABELS.get(child.ovz_level, child.ovz_level)
    ovz_nosology_label = OVZ_NOZOLOGY_LABELS.get(child.ovz_nosology, child.ovz_nosology)

    transfer_history = (
        ChildTransferHistory.query
        .filter_by(child_id=child.id)
        .order_by(ChildTransferHistory.transfer_date.desc().nullslast(), ChildTransferHistory.created_at.desc())
        .all()
    )
    enrollment_history = (
        ChildEnrollment.query
        .filter_by(child_id=child.id)
        .order_by(ChildEnrollment.enrolled_at.desc())
        .all()
    )
    expel_event = (
        _active_expel_event(child.id)
        if is_admin(current_user)
        else None
    )
    expel_undo = (
        {
            "event": expel_event,
            "enrollments": _expelled_enrollments_for_event(expel_event),
        }
        if expel_event is not None
        else None
    )
    active_enrollment_count = sum(
        1
        for enrollment in enrollment_history
        if enrollment.status == "ACTIVE" and enrollment.ended_at is None
    )
    transfer_undo = (
        _latest_reversible_enrollment(child.id)
        if is_admin(current_user) and expel_undo is None
        else None
    )
    control_results_q = ControlWorkResult.query.filter_by(child_id=child.id)
    if selected_year_id:
        control_results_q = control_results_q.filter_by(academic_year_id=selected_year_id)
    control_results = (
        control_results_q
        .order_by(ControlWorkResult.created_at.desc())
        .limit(50)
        .all()
    )
    olympiad_q = OlympiadResult.query.filter_by(child_id=child.id, is_archived=False)
    if selected_year_id:
        olympiad_q = olympiad_q.filter_by(academic_year_id=selected_year_id)
    olympiad_results = olympiad_q.order_by(OlympiadResult.created_at.desc()).limit(100).all()
    academic_years = AcademicYear.query.order_by(AcademicYear.start_date.desc().nullslast(), AcademicYear.name.desc()).all()

    flags = build_child_card_flags(child)
    child_service_assignments = child_support_assignments(child, user=current_user)
    can_view_service_assignments = can_view_child_service_block(child, user=current_user)

    task_query = Task.query.filter_by(child_id=child.id)
    if getattr(current_user, 'role', None) not in {'ADMIN', 'DEPUTY_DIRECTOR', 'METHODIST'}:
        task_query = task_query.filter(db.or_(Task.creator_user_id == current_user.id, Task.responsible_user_id == current_user.id))
    child_tasks = task_query.order_by(Task.deadline_at.is_(None), Task.deadline_at.asc(), Task.created_at.desc()).limit(20).all()
    child_task_stats = {
        'active': sum(1 for t in child_tasks if t.status not in {Task.STATUS_DONE, Task.STATUS_CLOSED, Task.STATUS_CANCELLED} and not t.is_overdue),
        'overdue': sum(1 for t in child_tasks if t.is_overdue),
        'done': sum(1 for t in child_tasks if t.status in {Task.STATUS_DONE, Task.STATUS_CLOSED}),
    }

    return render_template(
        "child_card.html",
        child=child,
        return_class=return_class if can_return_to_class else None,
        return_class_url=return_class_url,
        return_contingent_url=return_contingent_url,
        social=social,
        mother=mother,
        father=father,
        debts=child.debts,
        open_debts=open_debts,
        closed_debts=closed_debts,
        documents=docs,
        documents_ovz=documents_ovz,
        documents_vshu=documents_vshu,
        documents_low=documents_low,
        documents_az=documents_az,
        documents_mse=documents_mse,
        documents_ipra=documents_ipra,
        documents_general=documents_general,
        documents_disabled=documents_disabled,
        ovz_levels=OVZ_LEVELS,
        ovz_nozologies=OVZ_NOZOLOGIES,
        ovz_allowed=ovz_allowed,
        school_classes=school_classes,
        current_year_id=current_year_id,
        current_year_name=year.name if year else None,
        academic_years=academic_years,
        selected_year_id=selected_year_id,
        incidents=incidents,
        comments=comments,
        events=events,
        transfer_history=transfer_history,
        enrollment_history=enrollment_history,
        transfer_undo=transfer_undo,
        expel_undo=expel_undo,
        can_manage_active_enrollment=(
            is_admin(current_user)
            and child.status == "ACTIVE"
            and active_enrollment_count > 0
            and expel_undo is None
        ),
        control_results=control_results,
        olympiad_results=olympiad_results,
        today=date.today(),
        datetime=datetime,
        show_ovz=show_ovz,
        show_disabled=show_disabled,
        show_vshu=show_vshu,
        show_low=show_low,
        show_az=show_az,
        show_general=show_general,
        ovz_level_label=ovz_level_label,
        ovz_nosology_label=ovz_nosology_label,
        child_service_assignments=child_service_assignments,
        can_view_service_assignments=can_view_service_assignments,
        child_tasks=child_tasks,
        child_task_stats=child_task_stats,
        **flags
    )
# =========================================================
# DOCUMENTS
# =========================================================
@children_bp.route("/children/<int:child_id>/documents/upload", methods=["POST"])
@login_required
def upload_child_document(child_id: int):
    _ensure_can_edit()

    child = Child.query.get_or_404(child_id)

    file = request.files.get("file")
    doc_type = (request.form.get("doc_type") or "").strip().upper()
    debt_id = request.form.get("debt_id", type=int)

    if not file or not file.filename:
        flash("Выберите файл", "danger")
        return redirect(url_for("children.child_card", child_id=child.id))

    if not doc_type:
        doc_type = "GENERAL"

    allowed_types = {"OVZ", "MSE", "IPRA", "VSHU", "LOW", "AZ", "GENERAL", "DISABLED"}
    if doc_type not in allowed_types:
        doc_type = "GENERAL"

    upload_root = current_app.config.get("UPLOAD_FOLDER")
    if not upload_root:
        flash("Не настроена папка для загрузки файлов", "danger")
        return redirect(url_for("children.child_card", child_id=child.id))

    child_folder = os.path.join(upload_root, str(child.id))
    os.makedirs(child_folder, exist_ok=True)

    original_name = file.filename
    safe_name = re.sub(r"[^0-9A-Za-zА-Яа-я._ -]+", "_", original_name).strip()
    if not safe_name:
        safe_name = f"document_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"

    stored_name = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S_%f')}_{safe_name}"
    stored_path = os.path.join(child_folder, stored_name)

    file.save(stored_path)

    current_year = _get_current_year()
    retention_until = None
    if current_year and current_year.end_date:
        try:
            retention_until = current_year.end_date.replace(year=current_year.end_date.year + 7)
        except Exception:
            retention_until = None

    doc = Document(
        child_id=child.id,
        debt_id=debt_id if debt_id else None,
        academic_year_id=current_year.id if current_year else None,
        doc_type=doc_type,
        original_name=original_name,
        stored_path=stored_path,
        filename=stored_name,
        title=original_name,
        uploaded_by_user_id=getattr(current_user, "id", None),
        uploaded_at=datetime.utcnow(),
        retention_until=retention_until,
    )

    db.session.add(doc)
    db.session.commit()

    flash("Документ загружен", "success")
    return redirect(url_for("children.child_card", child_id=child.id))


@children_bp.route("/documents/<int:doc_id>/download")
@login_required
def download_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)

    if not (can_view_documents(child) or can_upload_documents(child) or is_admin()):
        abort(403)

    abs_path = _document_abs_path(doc.stored_path)
    if not abs_path or not os.path.isfile(abs_path):
        flash("Файл не найден", "danger")
        return redirect(url_for("children.child_card", child_id=doc.child_id))

    return send_file(
        abs_path,
        as_attachment=True,
        download_name=doc.original_name or os.path.basename(abs_path)
    )


@children_bp.route("/documents/<int:doc_id>/view")
@login_required
def view_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)

    if not (can_view_documents(child) or can_upload_documents(child) or is_admin()):
        abort(403)

    abs_path = _document_abs_path(doc.stored_path)
    if not abs_path or not os.path.isfile(abs_path):
        abort(404)

    mime, _ = mimetypes.guess_type(abs_path)
    mime = mime or "application/octet-stream"
    resp = send_file(abs_path, mimetype=mime, as_attachment=False)
    resp.headers["Content-Disposition"] = "inline"
    resp.headers["Cache-Control"] = "no-store"
    return resp


@children_bp.route("/documents/<int:doc_id>/preview")
@login_required
def preview_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)

    if not (can_view_documents(child) or can_upload_documents(child) or is_admin()):
        abort(403)

    abs_path = _document_abs_path(doc.stored_path)
    if not abs_path or not os.path.isfile(abs_path):
        abort(404)

    ext = os.path.splitext(doc.original_name or abs_path)[1].lower()
    html = None
    inline_url = None
    mode = "html"
    preview_error = None

    try:
        if ext in {'.pdf', '.png', '.jpg', '.jpeg', '.gif', '.webp', '.txt', '.csv'}:
            if ext in {'.txt', '.csv'}:
                html = _render_text_preview(abs_path)
            else:
                mode = "iframe"
                inline_url = url_for('children.view_document', doc_id=doc.id)
        elif ext == '.docx':
            html = _render_docx_preview(abs_path)
        elif ext == '.xlsx':
            html = _render_xlsx_preview(abs_path)
        else:
            preview_error = 'Для этого формата доступно скачивание. Полноценный просмотр в окне не поддерживается.'
    except Exception:
        preview_error = 'Не удалось построить предпросмотр документа.'

    return render_template(
        'document_preview.html',
        doc=doc,
        mode=mode,
        inline_url=inline_url,
        preview_html=html,
        preview_error=preview_error,
    )


@children_bp.route("/documents/<int:doc_id>/delete", methods=["POST"])
@login_required
def delete_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child_id = doc.child_id
    child = Child.query.get_or_404(child_id)

    if not _user_can_manage_document(child):
        abort(403)

    doc.is_deleted_soft = True
    doc.deleted_at = datetime.utcnow()
    doc.deleted_by = getattr(current_user, "id", None)
    db.session.commit()

    flash("Документ скрыт из карточки. Файл сохранён в архиве.", "success")
    return redirect(url_for("children.child_card", child_id=child_id))

# =========================================================
# DELETE CHILD
# =========================================================
@children_bp.route("/children/<int:child_id>/delete", methods=["POST"])
@require_roles("ADMIN")
def delete_child(child_id: int):
    child = Child.query.get_or_404(child_id)
    child.status = "ARCHIVED"
    child.archived_at = datetime.utcnow()

    active = (
        ChildEnrollment.query
        .filter(ChildEnrollment.child_id == child.id, ChildEnrollment.ended_at.is_(None))
        .all()
    )
    for en in active:
        en.status = "ARCHIVED"
        en.ended_at = datetime.utcnow()
        db.session.add(ChildTransferHistory(
            child_id=child.id,
            from_academic_year_id=en.academic_year_id,
            to_academic_year_id=None,
            from_class_id=en.school_class_id,
            to_class_id=None,
            transfer_type="ARCHIVED",
            transfer_date=date.today(),
            comment="Архивирование карточки вместо физического удаления",
            created_by=getattr(current_user, "id", None),
        ))

    db.session.add(ChildEvent(
        child_id=child.id,
        author_id=getattr(current_user, "id", None),
        event_type="EXPEL",
        from_class=child.current_class_name,
        to_class=None,
        promotion_kind="ARCHIVED",
        reason="Архивирование карточки",
        created_at=datetime.utcnow(),
    ))

    db.session.commit()
    flash("Карточка ученика переведена в архив. Документы и история сохранены.", "success")
    return redirect(url_for("children.list_children"))


# =========================================================
# FLAGS
# =========================================================
@children_bp.route("/children/<int:child_id>/flags", methods=["POST"])
@login_required
def update_child_flags(child_id: int):
    _ensure_can_edit()

    child = Child.query.get_or_404(child_id)

    child.is_ovz = as_checkbox(request.form, "is_ovz")
    child.is_disabled = as_checkbox(request.form, "is_disabled")
    child.is_vshu = as_checkbox(request.form, "is_vshu")
    child.is_low = as_checkbox(request.form, "is_low")

    child.disability_mse = (request.form.get("disability_mse") or "").strip() or None
    child.disability_from = parse_date(request.form.get("disability_from"))
    child.disability_to = parse_date(request.form.get("disability_to"))

    if not child.is_disabled:
        child.disability_mse = None
        child.disability_from = None
        child.disability_to = None

    child.low_subjects = (request.form.get("low_subjects") or "").strip() or None
    child.low_notes = (request.form.get("low_notes") or "").strip() or None

    if child.is_ovz:
        child.ovz_level = (request.form.get("ovz_level") or "").strip().upper() or None
        child.ovz_nosology = (request.form.get("ovz_nosology") or "").strip().upper() or None

        v_raw = (request.form.get("ovz_variant") or "").strip()
        child.ovz_variant = int(v_raw) if v_raw.isdigit() else None

        if child.ovz_variant and not is_allowed(child.ovz_level, child.ovz_nosology, child.ovz_variant):
            child.ovz_variant = None
    else:
        child.ovz_level = None
        child.ovz_nosology = None
        child.ovz_variant = None

    _sync_child_az_flag(child)

    db.session.commit()
    flash("Сохранено", "success")
    return redirect(url_for("children.child_card", child_id=child_id))


# =========================================================
# PROFILE UPDATE
# =========================================================
@children_bp.route("/children/<int:child_id>/profile", methods=["POST"])
@login_required
def update_child_profile(child_id: int):
    child = Child.query.get_or_404(child_id)

    if not _can_edit_profile_admin_only(child):
        abort(403)

    child.last_name = (request.form.get("last_name") or "").strip()
    child.first_name = (request.form.get("first_name") or "").strip()
    child.middle_name = (request.form.get("middle_name") or "").strip() or None
    child.birth_date = parse_date(request.form.get("birth_date"))
    child.reg_address = (request.form.get("reg_address") or "").strip() or None
    child.notes = (request.form.get("notes") or "").strip() or None
    child.education_form = request.form.get("education_form") or None
    child.reg_address = request.form.get("reg_address") or None
    child.temporary_address = request.form.get("temporary_address") or None
    child.actual_address = request.form.get("actual_address") or None

    db.session.commit()
    flash("Основные данные обновлены", "success")
    return redirect(url_for("children.child_card", child_id=child.id))


# =========================================================
# SOCIAL PASSPORT
# =========================================================
@children_bp.route("/children/<int:child_id>/social-passport", methods=["POST"])
@login_required
def update_child_social_passport(child_id: int):
    child = Child.query.get_or_404(child_id)

    if not _can_edit_social_passport(child):
        abort(403)

    social = _get_or_create_social(child)

    _set_parent_relation(
        child,
        "mother",
        request.form.get("mother_fio"),
        request.form.get("mother_phone"),
    )
    _set_parent_relation(
        child,
        "father",
        request.form.get("father_fio"),
        request.form.get("father_phone"),
    )

    child.reg_address = (request.form.get("reg_address") or "").strip() or None

    social.family_status = (request.form.get("family_status") or "").strip() or None
    social.living_conditions = (request.form.get("living_conditions") or "").strip() or None
    social.social_risk = (request.form.get("social_risk") or "").strip() or None
    social.aoop_variant_text = (request.form.get("aoop_variant_text") or "").strip() or None

    child.is_ovz = as_checkbox(request.form, "is_ovz")
    child.ovz_level = (request.form.get("ovz_level") or "").strip().upper() or None
    child.ovz_nosology = (request.form.get("ovz_nosology") or "").strip().upper() or None
    v_raw = (request.form.get("ovz_variant") or "").strip()
    child.ovz_variant = int(v_raw) if v_raw.isdigit() else None
    child.ovz_doc_number = (request.form.get("ovz_doc_number") or "").strip() or None
    child.ovz_doc_date = parse_date(request.form.get("ovz_doc_date"))
    if child.is_ovz and child.ovz_variant and not is_allowed(child.ovz_level, child.ovz_nosology, child.ovz_variant):
        child.ovz_variant = None
    if not child.is_ovz:
        child.ovz_level = None
        child.ovz_nosology = None
        child.ovz_variant = None
        child.ovz_doc_number = None
        child.ovz_doc_date = None

    child.is_disabled = as_checkbox(request.form, "is_disabled")
    child.disability_mse = (request.form.get("disability_mse") or "").strip() or None
    child.disability_from = parse_date(request.form.get("disability_from"))
    child.disability_to = parse_date(request.form.get("disability_to"))
    child.disability_ipra = (request.form.get("disability_ipra") or "").strip() or None
    if not child.is_disabled:
        child.disability_mse = None
        child.disability_from = None
        child.disability_to = None
        child.disability_ipra = None

    social.vshu_since = parse_date(request.form.get("vshu_since"))
    social.vshu_reason = (request.form.get("vshu_reason") or "").strip() or None

    social.kdn_since = parse_date(request.form.get("kdn_since"))
    social.kdn_reason = (request.form.get("kdn_reason") or "").strip() or None

    social.pdn_since = parse_date(request.form.get("pdn_since"))
    social.pdn_reason = (request.form.get("pdn_reason") or "").strip() or None

    social.vshu_removed_at = parse_date(request.form.get("vshu_removed_at"))
    social.vshu_remove_reason = (request.form.get("vshu_remove_reason") or "").strip() or None

    child.is_vshu = bool(social.vshu_since and not social.vshu_removed_at)

    social.has_disability_parents = as_checkbox(request.form, "has_disability_parents")
    social.has_large_family = as_checkbox(request.form, "has_large_family")
    social.has_low_income_family = as_checkbox(request.form, "has_low_income_family")
    social.has_guardianship = as_checkbox(request.form, "has_guardianship")
    social.has_orphan_status = as_checkbox(request.form, "has_orphan_status")
    social.has_refugee_status = as_checkbox(request.form, "has_refugee_status")

    social.is_socially_dangerous = as_checkbox(request.form, "is_socially_dangerous")
    social.is_hard_life = as_checkbox(request.form, "is_hard_life")
    social.is_single_mother = as_checkbox(request.form, "is_single_mother")
    social.is_single_father = as_checkbox(request.form, "is_single_father")
    social.is_repeat_year = as_checkbox(request.form, "is_repeat_year")
    social.is_svo_family = as_checkbox(request.form, "is_svo_family")

    social.notes = (request.form.get("social_notes") or "").strip() or None
    social.updated_at = datetime.utcnow()

    db.session.commit()
    flash("Данные социального паспорта обновлены", "success")
    return redirect(url_for("children.child_card", child_id=child.id))


# =========================================================
# DEBTS
# =========================================================
@children_bp.route("/children/<int:child_id>/debts/add", methods=["POST"])
@login_required
def add_debt(child_id: int):
    _ensure_can_edit()

    child = Child.query.get_or_404(child_id)

    subject_id = (request.form.get("subject_id") or "").strip()
    detected_date = parse_date(request.form.get("detected_date")) or date.today()
    due_date = parse_date(request.form.get("due_date"))

    if not subject_id.isdigit():
        flash("Не выбран предмет", "danger")
        return redirect(url_for("children.child_card", child_id=child.id))

    activity = get_subject_activity(int(subject_id))
    if activity is None:
        flash("Выберите предмет из единого каталога.", "danger")
        return redirect(url_for("children.child_card", child_id=child.id))

    debt = Debt(
        child_id=child.id,
        detected_date=detected_date,
        due_date=due_date,
        status="OPEN",
        created_at=datetime.utcnow(),
        created_by_user_id=current_user.id,
    )
    assign_subject_activity(debt, activity)

    db.session.add(debt)
    db.session.flush()

    _sync_child_az_flag(child)

    db.session.commit()
    flash("Задолженность добавлена", "success")
    return redirect(url_for("children.child_card", child_id=child.id))


@children_bp.route("/debts/<int:debt_id>/close", methods=["POST"])
@login_required
def close_debt(debt_id: int):
    _ensure_can_edit()

    debt = Debt.query.get_or_404(debt_id)
    child = Child.query.get_or_404(debt.child_id)

    if debt.status != "CLOSED":
        debt.status = "CLOSED"
        debt.closed_at = datetime.utcnow()
        debt.closed_by_user_id = current_user.id

    _sync_child_az_flag(child)

    db.session.commit()
    flash("Задолженность закрыта", "success")
    return redirect(url_for("children.child_card", child_id=child.id))


@children_bp.route("/debts/<int:debt_id>/reopen", methods=["POST"])
@login_required
def reopen_debt(debt_id: int):
    _ensure_can_edit()

    debt = Debt.query.get_or_404(debt_id)
    child = Child.query.get_or_404(debt.child_id)

    if debt.status != "OPEN":
        debt.status = "OPEN"
        debt.closed_at = None
        debt.closed_by_user_id = None

    _sync_child_az_flag(child)

    db.session.commit()
    flash("Задолженность возвращена в открытые", "success")
    return redirect(url_for("children.child_card", child_id=child.id))


# =========================================================
# COMMENTS
# =========================================================
@children_bp.route("/children/<int:child_id>/comments", methods=["POST"])
@login_required
def add_child_comment(child_id: int):
    child = Child.query.get_or_404(child_id)

    text = (request.form.get("comment_text") or "").strip()
    if not text:
        return redirect(url_for("children.child_card", child_id=child_id))

    c = ChildComment(
        child_id=child.id,
        author_id=current_user.id,
        text=text
    )

    db.session.add(c)
    db.session.commit()

    return redirect(url_for("children.child_card", child_id=child_id))


@children_bp.route("/comments/<int:comment_id>/delete", methods=["POST"])
@login_required
def delete_child_comment(comment_id: int):
    c = ChildComment.query.get_or_404(comment_id)

    if not (is_admin(current_user) or c.author_id == current_user.id):
        abort(403)

    child_id = c.child_id
    db.session.delete(c)
    db.session.commit()

    return redirect(url_for("children.child_card", child_id=child_id))


# =========================================================
# TRANSFER / EXPEL
# =========================================================
@children_bp.route(
    "/children/<int:child_id>/enrollments/<int:enrollment_id>/undo",
    methods=["POST"],
)
@require_roles("ADMIN")
def undo_child_enrollment(child_id: int, enrollment_id: int):
    child = Child.query.get_or_404(child_id)
    requested_target = ChildEnrollment.query.filter_by(
        id=enrollment_id,
        child_id=child.id,
    ).first_or_404()
    undo_state = _latest_reversible_enrollment(child.id)
    if (
        undo_state is None
        or undo_state["target"].id != requested_target.id
    ):
        flash(
            "Эту запись нельзя отменить: сначала отмените более позднее "
            "изменение обучения ученика.",
            "warning",
        )
        return redirect(url_for("children.child_card", child_id=child.id))

    target = undo_state["target"]
    source = undo_state["source"]
    history = undo_state["history"]
    conflict = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.child_id == child.id,
            ChildEnrollment.academic_year_id
            == source.academic_year_id,
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.id.notin_([target.id, source.id]),
        )
        .first()
    )
    if conflict is not None:
        flash(
            "Отмена невозможна: в исходном учебном году уже есть другое "
            "активное зачисление. Сначала исправьте более позднюю запись.",
            "danger",
        )
        return redirect(url_for("children.child_card", child_id=child.id))

    operation_at = datetime.utcnow()
    target_class = target.school_class
    source_class = source.school_class
    target.status = "CANCELLED"
    target.ended_at = operation_at
    target.note = (
        f"{target.note + '. ' if target.note else ''}"
        "Зачисление отменено администратором."
    )
    # Close the target first so restoring a source from the same academic year
    # cannot violate the one-active-enrollment-per-year constraint.
    db.session.flush()

    source.status = "ACTIVE"
    source.ended_at = None
    child.status = "ACTIVE"
    child.archived_at = None

    reversal_reason = (
        request.form.get("reason") or "Исправление ошибочного перевода"
    ).strip()
    if history is None:
        # Preserve an audit trail for legacy manual class changes that created
        # enrollments but did not write ChildTransferHistory.
        history = ChildTransferHistory(
            child_id=child.id,
            from_academic_year_id=source.academic_year_id,
            to_academic_year_id=target.academic_year_id,
            from_class_id=source.school_class_id,
            to_class_id=target.school_class_id,
            transfer_type="MANUAL",
            transfer_date=(
                target.enrolled_at.date()
                if target.enrolled_at
                else date.today()
            ),
            comment="Восстановлена запись старого ручного перевода",
            created_by=getattr(current_user, "id", None),
            created_at=target.enrolled_at or operation_at,
        )
        db.session.add(history)

    history.reversed_at = operation_at
    history.reversed_by = getattr(current_user, "id", None)
    history.reversal_reason = reversal_reason

    db.session.add(ChildEvent(
        child_id=child.id,
        author_id=current_user.id,
        event_type="TRANSFER_UNDO",
        from_class=target_class.name if target_class else None,
        to_class=source_class.name if source_class else None,
        promotion_kind="UNDO",
        reason=reversal_reason,
        created_at=operation_at,
    ))
    db.session.add(ChildMovement(
        child_id=child.id,
        academic_year_id=source.academic_year_id,
        movement_type="correction",
        movement_date=date.today(),
        from_class_id=target.school_class_id,
        to_class_id=source.school_class_id,
        reason=reversal_reason,
        created_by=current_user.id,
        created_at=operation_at,
    ))
    db.session.commit()

    flash(
        "Последнее изменение отменено. Восстановлен класс "
        f"{source_class.name if source_class else '—'} "
        f"({source.academic_year.name if source.academic_year else '—'}).",
        "success",
    )
    return redirect(url_for("children.child_card", child_id=child.id))


@children_bp.route(
    "/children/<int:child_id>/expel/undo",
    methods=["POST"],
)
@require_roles("ADMIN")
def undo_child_expulsion(child_id: int):
    child = Child.query.get_or_404(child_id)
    expel_event = _active_expel_event(child.id)
    if expel_event is None:
        flash(
            "Отмена невозможна: последнее отчисление уже отменено "
            "или после него выполнено другое действие.",
            "warning",
        )
        return redirect(url_for("children.child_card", child_id=child.id))

    enrollments = _expelled_enrollments_for_event(expel_event)
    legacy_active_enrollments = []
    if not enrollments:
        # Older expulsion code could write the event but leave a future-year
        # enrollment active. Cancelling that event must keep this enrollment.
        legacy_active_enrollments = (
            ChildEnrollment.query
            .filter(
                ChildEnrollment.child_id == child.id,
                ChildEnrollment.status == "ACTIVE",
                ChildEnrollment.ended_at.is_(None),
                ChildEnrollment.enrolled_at <= expel_event.created_at,
            )
            .order_by(
                ChildEnrollment.enrolled_at.desc(),
                ChildEnrollment.id.desc(),
            )
            .all()
        )
        if not legacy_active_enrollments:
            flash(
                "Не найдена запись обучения, связанная с этим отчислением. "
                "Данные не изменены.",
                "danger",
            )
            return redirect(
                url_for("children.child_card", child_id=child.id)
            )

    enrollment_ids = [item.id for item in enrollments]
    year_ids = {item.academic_year_id for item in enrollments}
    conflict = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.child_id == child.id,
            ChildEnrollment.academic_year_id.in_(year_ids),
            ChildEnrollment.status == "ACTIVE",
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.id.notin_(enrollment_ids),
        )
        .first()
    )
    if conflict is not None:
        flash(
            "Отмена невозможна: после отчисления уже создано новое "
            "активное зачисление. Сначала отмените более позднее действие.",
            "danger",
        )
        return redirect(url_for("children.child_card", child_id=child.id))

    operation_at = datetime.utcnow()
    for enrollment in enrollments:
        enrollment.status = "ACTIVE"
        enrollment.ended_at = None
        enrollment.note = (
            f"{enrollment.note + '. ' if enrollment.note else ''}"
            "Ошибочное отчисление отменено администратором."
        )

    child.status = "ACTIVE"
    child.archived_at = None
    reversal_reason = (
        request.form.get("reason") or "Исправление ошибочного отчисления"
    ).strip()
    histories = (
        ChildTransferHistory.query
        .filter(
            ChildTransferHistory.child_id == child.id,
            ChildTransferHistory.transfer_type == "EXPELLED",
            ChildTransferHistory.reversed_at.is_(None),
            ChildTransferHistory.created_at
            >= expel_event.created_at - timedelta(minutes=10),
            ChildTransferHistory.created_at
            <= expel_event.created_at + timedelta(minutes=10),
        )
        .all()
    )
    for history in histories:
        history.reversed_at = operation_at
        history.reversed_by = current_user.id
        history.reversal_reason = reversal_reason

    restored = (enrollments or legacy_active_enrollments)[0]
    restored_class = restored.school_class
    db.session.add(ChildEvent(
        child_id=child.id,
        author_id=current_user.id,
        event_type="EXPEL_UNDO",
        from_class=None,
        to_class=restored_class.name if restored_class else None,
        promotion_kind="UNDO",
        reason=reversal_reason,
        created_at=operation_at,
    ))
    db.session.add(ChildMovement(
        child_id=child.id,
        academic_year_id=restored.academic_year_id,
        movement_type="correction",
        movement_date=date.today(),
        from_class_id=None,
        to_class_id=restored.school_class_id,
        reason=reversal_reason,
        created_by=current_user.id,
        created_at=operation_at,
    ))
    db.session.commit()

    flash(
        "Отчисление отменено. Ученик восстановлен в классе "
        f"{restored_class.name if restored_class else '—'}.",
        "success",
    )
    return redirect(url_for("children.child_card", child_id=child.id))


@children_bp.route("/children/<int:child_id>/transfer", methods=["POST"])
@require_roles("ADMIN")
def transfer_child(child_id: int):
    child = Child.query.get_or_404(child_id)

    is_repeat = (request.form.get("is_repeat") == "1")
    note = (request.form.get("note") or "").strip() or None

    year = _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.child_card", child_id=child_id))

    en = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.child_id == child.id,
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None)
        )
        .first()
    )

    old_class = (
        en.school_class.name
        if en and en.school_class
        else child.current_class_name or "—"
    )

    if is_repeat:
        db.session.add(ChildEvent(
            child_id=child.id,
            author_id=current_user.id,
            event_type="REPEAT",
            from_class=old_class,
            to_class=old_class,
            reason=note,
            created_at=datetime.utcnow(),
        ))
        db.session.commit()
        flash("Сохранено", "success")
        return redirect(url_for("children.child_card", child_id=child_id))

    school_class_id = request.form.get("school_class_id")
    if not (school_class_id and str(school_class_id).isdigit()):
        flash("Выберите класс из реестра", "danger")
        return redirect(url_for("children.child_card", child_id=child_id))

    sc = (
        SchoolClass.query
        .filter(
            SchoolClass.id == int(school_class_id),
            SchoolClass.academic_year_id == year.id
        )
        .first()
    )
    if not sc:
        flash("Выбранный класс не найден в текущем учебном году", "danger")
        return redirect(url_for("children.child_card", child_id=child_id))

    if en and en.school_class_id == sc.id:
        flash("Ученик уже находится в выбранном классе.", "info")
        return redirect(url_for("children.child_card", child_id=child_id))

    operation_at = datetime.utcnow()
    if en:
        en.ended_at = operation_at
        en.status = "TRANSFERRED"

    new_en = ChildEnrollment(
        child_id=child.id,
        academic_year_id=year.id,
        school_class_id=sc.id,
        status="ACTIVE",
        enrolled_at=operation_at,
    )
    db.session.add(new_en)
    child.status = "ACTIVE"
    child.archived_at = None

    db.session.add(ChildEvent(
        child_id=child.id,
        author_id=current_user.id,
        event_type="TRANSFER",
        from_class=old_class,
        to_class=sc.name,
        promotion_kind="MANUAL",
        reason=note,
        created_at=operation_at,
    ))
    db.session.add(ChildTransferHistory(
        child_id=child.id,
        from_academic_year_id=en.academic_year_id if en else None,
        to_academic_year_id=year.id,
        from_class_id=en.school_class_id if en else None,
        to_class_id=sc.id,
        transfer_type="MANUAL",
        transfer_date=date.today(),
        comment=note,
        created_by=current_user.id,
        created_at=operation_at,
    ))
    db.session.add(ChildMovement(
        child_id=child.id,
        academic_year_id=year.id,
        movement_type="transfer",
        movement_date=date.today(),
        from_class_id=en.school_class_id if en else None,
        to_class_id=sc.id,
        reason=note,
        created_by=current_user.id,
        created_at=operation_at,
    ))

    db.session.commit()
    flash("Сохранено", "success")
    return redirect(url_for("children.child_card", child_id=child_id))


@children_bp.route("/children/<int:child_id>/expel", methods=["POST"])
@require_roles("ADMIN")
def expel_child(child_id: int):
    child = Child.query.get_or_404(child_id)

    note = (request.form.get("note") or "").strip() or None
    to_where = (request.form.get("to_where") or "").strip() or None
    active_enrollments = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.child_id == child.id,
            ChildEnrollment.status == "ACTIVE",
            ChildEnrollment.ended_at.is_(None),
        )
        .order_by(
            ChildEnrollment.enrolled_at.desc(),
            ChildEnrollment.id.desc(),
        )
        .all()
    )
    if not active_enrollments or _active_expel_event(child.id) is not None:
        flash(
            "Ученик уже отчислен или не имеет активного зачисления.",
            "warning",
        )
        return redirect(url_for("children.child_card", child_id=child_id))

    operation_at = datetime.utcnow()
    primary_enrollment = active_enrollments[0]
    old_class = (
        primary_enrollment.school_class.name
        if primary_enrollment.school_class
        else "—"
    )
    for enrollment in active_enrollments:
        enrollment.ended_at = operation_at
        enrollment.status = "EXPELLED"
        enrollment.note = note

    child.status = "EXPELLED"
    db.session.add(ChildTransferHistory(
        child_id=child.id,
        from_academic_year_id=primary_enrollment.academic_year_id,
        to_academic_year_id=None,
        from_class_id=primary_enrollment.school_class_id,
        to_class_id=None,
        transfer_type="EXPELLED",
        transfer_date=date.today(),
        comment=note,
        created_by=current_user.id,
        created_at=operation_at,
    ))
    db.session.add(ChildMovement(
        child_id=child.id,
        academic_year_id=primary_enrollment.academic_year_id,
        movement_type="leave",
        movement_date=date.today(),
        from_class_id=primary_enrollment.school_class_id,
        to_class_id=None,
        reason=note,
        created_by=current_user.id,
        created_at=operation_at,
    ))

    db.session.add(
        ChildEvent(
            child_id=child.id,
            author_id=current_user.id,
            event_type="EXPEL",
            from_class=old_class,
            to_class=to_where,
            reason=note,
            created_at=operation_at,
        )
    )

    db.session.commit()
    flash("Ребёнок отчислен", "success")
    return redirect(url_for("children.child_card", child_id=child_id))


# =========================================================
# CONTINGENT
# =========================================================
@children_bp.route("/contingent")
@login_required
def contingent():
    if not (has_role("ADMIN") or has_role("METHODIST")):
        abort(403)
    year_id = request.args.get("year_id", type=int)
    building_id = request.args.get("building_id", type=int)

    years = AcademicYear.query.order_by(AcademicYear.created_at.desc()).all()
    buildings = Building.query.order_by(Building.name.asc()).all()

    if not year_id:
        y = _get_current_year()
        year_id = y.id if y else None

    q = SchoolClass.query
    if year_id:
        q = q.filter(SchoolClass.academic_year_id == year_id)
    if building_id:
        q = q.filter(SchoolClass.building_id == building_id)

    classes = (
        q.outerjoin(
            Building,
            SchoolClass.building_id == Building.id,
        )
        .order_by(
            SchoolClass.grade.asc().nullslast(),
            Building.name.asc().nullslast(),
            SchoolClass.letter.asc().nullslast(),
            SchoolClass.name.asc(),
        )
        .all()
    )

    teachers = User.query.order_by(User.last_name.asc(), User.first_name.asc()).all()
    teachers_map = {u.id: u for u in teachers}
    buildings_map = {b.id: b for b in buildings}
    building_tone_map = {
        building.id: building_matrix_tone(building)
        for building in buildings
    }

    class_counts = dict(
        db.session.query(
            ChildEnrollment.school_class_id,
            db.func.count(ChildEnrollment.id)
        )
        .join(Child, Child.id == ChildEnrollment.child_id)
        .filter(
            *_active_class_enrollment_filters(year_id),
        )
        .group_by(ChildEnrollment.school_class_id)
        .all()
    ) if year_id else {}

    transfer_counts = {
        "PROMOTED": 0,
        "CONDITIONAL": 0,
        "REPEAT": 0,
        "EXPELLED": 0,
        "TRANSFERRED_OUT": 0,
        "ARCHIVED": 0,
    }
    if year_id:
        for t_type, cnt in (
            db.session.query(ChildTransferHistory.transfer_type, db.func.count(ChildTransferHistory.id))
            .filter(
                ChildTransferHistory.from_academic_year_id == year_id,
                ChildTransferHistory.reversed_at.is_(None),
            )
            .group_by(ChildTransferHistory.transfer_type)
            .all()
        ):
            transfer_counts[t_type] = int(cnt or 0)

    totals = {
        "school": 0,
        "grades_1_4": 0,
        "grades_5_9": 0,
        "grades_10_11": 0,
        "boys": 0,
        "girls": 0,

        "ovz": 0,
        "vshu": 0,
        "kdn": 0,

        "by_grade": {},
        "by_building": {},
        "education_forms": {},
        "level_stats": {},

        "classes_total": 0,
        "classes_1_4": 0,
        "classes_5_9": 0,
        "classes_10_11": 0,

        "parallel_stats": {},
        "pending_transfer": 0,
        "transferred_out": 0,
        "repeat_total": 0,
        "conditional_total": 0,
    }

    # --- Батч-подсчёт через SQL GROUP BY вместо загрузки всех Child (perf s51) ---
    from collections import defaultdict

    class_ids = [c.id for c in classes]
    profiles_by_class: dict[int, list[str]] = {}
    if year_id and class_ids:
        versions = (
            TariffVersion.query
            .join(TariffCycle)
            .filter(TariffCycle.academic_year_id == year_id)
            .all()
        )
        version_priority = {
            "DRAFT": 7,
            "EFFECTIVE": 6,
            "APPROVED": 5,
            "APPROVAL": 4,
            "VALIDATION": 3,
            "SUPERSEDED": 1,
            "ARCHIVED": 0,
        }
        selected_version = (
            max(
                versions,
                key=lambda item: (
                    version_priority.get(item.status, -1),
                    item.version_no,
                    item.id,
                ),
            )
            if versions else None
        )
        snapshot = (
            current_population_snapshot(selected_version.id)
            if selected_version else None
        )
        if snapshot is not None:
            profile_rows = (
                db.session.query(
                    PopulationSnapshotClass.source_school_class_id,
                    EducationPlan.name,
                    EducationPlan.profile_name,
                )
                .join(
                    EducationPlanBinding,
                    EducationPlanBinding.population_snapshot_class_id
                    == PopulationSnapshotClass.id,
                )
                .join(
                    EducationPlan,
                    EducationPlan.id
                    == EducationPlanBinding.education_plan_id,
                )
                .filter(
                    PopulationSnapshotClass.population_snapshot_id
                    == snapshot.id,
                    PopulationSnapshotClass.source_school_class_id.in_(
                        class_ids
                    ),
                    EducationPlan.plan_kind == "CURRICULUM",
                    EducationPlan.root_plan_id.is_(None),
                )
                .all()
            )
            profile_sets = defaultdict(set)
            for school_class_id, plan_name, profile_name in profile_rows:
                display_name = profile_name or plan_name
                if display_name:
                    profile_sets[school_class_id].add(display_name)
            profiles_by_class = {
                class_id: sorted(
                    profile_names,
                    key=str.casefold,
                )
                for class_id, profile_names in profile_sets.items()
            }
    boys_by_class = {}
    girls_by_class = {}
    ovz_by_class = {}
    vshu_by_class = {}
    kdn_by_class = {}
    edu_form_totals = {}
    transfer_by_class = defaultdict(dict)

    if year_id and class_ids:
        enrollment_base = (
            db.session.query(
                ChildEnrollment.school_class_id, db.func.count(ChildEnrollment.id)
            )
            .join(Child, Child.id == ChildEnrollment.child_id)
            .filter(
                *_active_class_enrollment_filters(year_id),
                ChildEnrollment.school_class_id.in_(class_ids),
            )
            .group_by(ChildEnrollment.school_class_id)
        )

        boys_by_class = dict(
            enrollment_base.filter(db.func.upper(Child.gender) == "М").all()
        )
        girls_by_class = dict(
            enrollment_base.filter(db.func.upper(Child.gender) == "Ж").all()
        )
        ovz_by_class = dict(
            enrollment_base.filter(Child.is_ovz.is_(True)).all()
        )

        # ВШУ: Child.is_vshu OR (social.vshu_since IS NOT NULL AND social.vshu_removed_at IS NULL)
        vshu_by_class = dict(
            db.session.query(
                ChildEnrollment.school_class_id, db.func.count(db.distinct(ChildEnrollment.id))
            )
            .join(Child, Child.id == ChildEnrollment.child_id)
            .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
            .filter(
                *_active_class_enrollment_filters(year_id),
                ChildEnrollment.school_class_id.in_(class_ids),
                db.or_(
                    Child.is_vshu.is_(True),
                    db.and_(
                        ChildSocial.vshu_since.isnot(None),
                        ChildSocial.vshu_removed_at.is_(None),
                    ),
                ),
            )
            .group_by(ChildEnrollment.school_class_id)
            .all()
        )

        # КДН: social.kdn_since IS NOT NULL
        kdn_by_class = dict(
            db.session.query(
                ChildEnrollment.school_class_id, db.func.count(db.distinct(ChildEnrollment.id))
            )
            .join(Child, Child.id == ChildEnrollment.child_id)
            .join(ChildSocial, ChildSocial.child_id == Child.id)
            .filter(
                *_active_class_enrollment_filters(year_id),
                ChildEnrollment.school_class_id.in_(class_ids),
                ChildSocial.kdn_since.isnot(None),
            )
            .group_by(ChildEnrollment.school_class_id)
            .all()
        )

        # education_form по школе (агрегат, по всем классам сразу)
        for form_raw, cnt in (
            db.session.query(Child.education_form, db.func.count(ChildEnrollment.id))
            .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
            .filter(
                *_active_class_enrollment_filters(year_id),
                ChildEnrollment.school_class_id.in_(class_ids),
            )
            .group_by(Child.education_form)
            .all()
        ):
            form_name = (form_raw or "Не указана").strip() or "Не указана"
            edu_form_totals[form_name] = edu_form_totals.get(form_name, 0) + int(cnt or 0)

        # Трансферы по классам (как было)
        for class_id, t_type, cnt in (
            db.session.query(
                ChildEnrollment.school_class_id,
                ChildTransferHistory.transfer_type,
                db.func.count(ChildTransferHistory.id),
            )
            .join(ChildTransferHistory, ChildTransferHistory.child_id == ChildEnrollment.child_id)
            .join(Child, Child.id == ChildEnrollment.child_id)
            .filter(
                *_active_class_enrollment_filters(year_id),
                ChildEnrollment.school_class_id.in_(class_ids),
                ChildTransferHistory.from_academic_year_id == year_id,
                ChildTransferHistory.reversed_at.is_(None),
            )
            .group_by(ChildEnrollment.school_class_id, ChildTransferHistory.transfer_type)
            .all()
        ):
            transfer_by_class[class_id][t_type] = int(cnt or 0)
    sc_by_class: dict[int, int] = {}
    do_by_class: dict[int, int] = {}

    rows = []

    for c in classes:
        total = int(class_counts.get(c.id, 0))

        boys_count = int(boys_by_class.get(c.id, 0))
        girls_count = int(girls_by_class.get(c.id, 0))
        ovz_count = int(ovz_by_class.get(c.id, 0))
        vshu_count = int(vshu_by_class.get(c.id, 0))
        kdn_count = int(kdn_by_class.get(c.id, 0))
        sc_count = int(sc_by_class.get(c.id, 0))
        do_count = int(do_by_class.get(c.id, 0))

        free = int((c.max_students or 0) - total)

        teacher = teachers_map.get(c.teacher_user_id)
        teacher_fio = teacher.fio if teacher else None
        teacher_phone = teacher.phone if teacher else None

        class_transfer_types = transfer_by_class.get(c.id, {})
        transferred_total = sum(class_transfer_types.get(k, 0) for k in ["PROMOTED", "CONDITIONAL", "REPEAT"])
        pending_transfer = max(total - transferred_total, 0)

        grade = c.grade
        if grade is not None and 1 <= grade <= 4:
            level_code = "NOO"
            level_display_code = "НОО"
            level_label = "Начальное общее образование"
            level_range = "1–4"
        elif grade is not None and 5 <= grade <= 9:
            level_code = "OOO"
            level_display_code = "ООО"
            level_label = "Основное общее образование"
            level_range = "5–9"
        elif grade is not None and 10 <= grade <= 11:
            level_code = "SOO"
            level_display_code = "СОО"
            level_label = "Среднее общее образование"
            level_range = "10–11"
        else:
            level_code = "OTHER"
            level_display_code = "—"
            level_label = "Классы без уровня образования"
            level_range = ""

        rows.append({
            "class": c,
            "total": total,
            "free": free,
            "boys": boys_count,
            "girls": girls_count,
            "teacher_fio": teacher_fio,
            "teacher_phone": teacher_phone,
            "profile_names": profiles_by_class.get(c.id, []),
            "applications_count": int(c.applications_count or 0),
            "building_tone": building_tone_map.get(c.building_id, 0),
            "level_code": level_code,
            "level_display_code": level_display_code,
            "level_label": level_label,
            "level_range": level_range,
            "sc_in_club": sc_count,
            "do_count": do_count,
            "pending_transfer": pending_transfer,
            "promoted": class_transfer_types.get("PROMOTED", 0),
            "conditional": class_transfer_types.get("CONDITIONAL", 0),
            "repeat": class_transfer_types.get("REPEAT", 0),
        })

        totals["school"] += total
        totals["boys"] += boys_count
        totals["girls"] += girls_count

        totals["ovz"] += ovz_count
        totals["vshu"] += vshu_count
        totals["kdn"] += kdn_count

        totals["classes_total"] += 1
        totals["pending_transfer"] += pending_transfer
        totals["transferred_out"] += transferred_total
        totals["repeat_total"] += class_transfer_types.get("REPEAT", 0)
        totals["conditional_total"] += class_transfer_types.get("CONDITIONAL", 0)

        building = buildings_map.get(c.building_id)
        bname = building.name if building else "Без здания"
        totals["by_building"][bname] = totals["by_building"].get(bname, 0) + total

        if level_code not in totals["level_stats"]:
            totals["level_stats"][level_code] = {
                "code": level_code,
                "display_code": level_display_code,
                "label": level_label,
                "range": level_range,
                "classes": 0,
                "children": 0,
                "boys": 0,
                "girls": 0,
                "applications": 0,
                "capacity": 0,
                "free": 0,
            }

        level_stats = totals["level_stats"][level_code]
        level_stats["classes"] += 1
        level_stats["children"] += total
        level_stats["boys"] += boys_count
        level_stats["girls"] += girls_count
        level_stats["applications"] += int(c.applications_count or 0)
        level_stats["capacity"] += int(c.max_students or 0)
        level_stats["free"] += free

        if grade is not None:
            totals["by_grade"][grade] = totals["by_grade"].get(grade, 0) + total

            if grade not in totals["parallel_stats"]:
                totals["parallel_stats"][grade] = {
                    "classes": 0,
                    "children": 0,
                    "boys": 0,
                    "girls": 0,
                    "applications": 0,
                    "capacity": 0,
                    "free": 0,
                }

            totals["parallel_stats"][grade]["classes"] += 1
            totals["parallel_stats"][grade]["children"] += total
            totals["parallel_stats"][grade]["boys"] += boys_count
            totals["parallel_stats"][grade]["girls"] += girls_count
            totals["parallel_stats"][grade]["applications"] += int(
                c.applications_count or 0
            )
            totals["parallel_stats"][grade]["capacity"] += int(
                c.max_students or 0
            )
            totals["parallel_stats"][grade]["free"] += free

            if 1 <= grade <= 4:
                totals["grades_1_4"] += total
                totals["classes_1_4"] += 1
            elif 5 <= grade <= 9:
                totals["grades_5_9"] += total
                totals["classes_5_9"] += 1
            elif 10 <= grade <= 11:
                totals["grades_10_11"] += total
                totals["classes_10_11"] += 1

    totals["education_forms"] = edu_form_totals

    return render_template(
        "contingent.html",
        rows=rows,
        years=years,
        buildings=buildings,
        year_id=year_id,
        building_id=building_id,
        totals=totals,
        transfer_counts=transfer_counts,
        can_edit_applications=(
            getattr(current_user, "role", None)
            in {
                "ADMIN",
                "DIRECTOR",
                "DEPUTY_DIRECTOR",
                "SECRETARY",
                "SECRETARY_ACADEMIC",
            }
        ),
    )
# =========================================================
# IMPORT CHILDREN
# =========================================================
@children_bp.route("/children/import", methods=["GET", "POST"])
@require_roles("ADMIN")
def children_import():
    print("=== NEW CHILDREN IMPORT ROUTE ===")
    year = _get_current_year()

    if request.method == "POST":
        f = request.files.get("file")

        if not f or not f.filename:
            flash("Выберите Excel файл", "danger")
            return redirect(url_for("children.children_import"))

        if not year:
            flash("Не найден текущий учебный год", "danger")
            return redirect(url_for("children.children_import"))

        wb = load_workbook(f, data_only=True)
        ws = wb.active

        headers = [(str(c.value).strip() if c.value else "") for c in ws[1]]
        idx = {h: i for i, h in enumerate(headers)}

        required = [
            "ФИО",
            "Пол",
            "Родился",
            "Номер и буква класса"
        ]

        missing = [c for c in required if c not in idx]

        if missing:
             flash(f"НОВЫЙ ИМПОРТ: не хватает колонок: {', '.join(missing)}", "danger")
             return redirect(url_for("children.children_import"))

        created = 0
        skipped = 0

        def parse_birth(x):
            if not x:
                return None
            if isinstance(x, datetime):
                return x.date()
            if isinstance(x, date):
                return x

            s = str(x).strip()

            m = re.match(r"^(\d{2})\.(\d{2})\.(\d{4})$", s)
            if m:
                return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))

            try:
                return datetime.strptime(s, "%Y-%m-%d").date()
            except Exception:
                return None

        def split_fio(fio):
            parts = str(fio).strip().split()
            last = parts[0] if len(parts) > 0 else ""
            first = parts[1] if len(parts) > 1 else ""
            middle = parts[2] if len(parts) > 2 else None
            return last, first, middle

        for r in range(2, ws.max_row + 1):

            row = [ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)]

            fio = str(row[idx["ФИО"]] or "").strip()

            if not fio:
                skipped += 1
                continue

            last_name, first_name, middle_name = split_fio(fio)

            birth_date = parse_birth(row[idx["Родился"]])

            gender_raw = str(row[idx["Пол"]] or "").strip().lower()

            if gender_raw in ["м", "муж", "мужской"]:
                gender = "М"
            elif gender_raw in ["ж", "жен", "женский"]:
                gender = "Ж"
            else:
                gender = None

            class_name = normalize_class_name(row[idx["Номер и буква класса"]])

            education_form = None
            if "Сведения о форме обучения" in idx:
                education_form = str(row[idx["Сведения о форме обучения"]] or "").strip() or None

            reg_address = None
            if "Регистрация по месту жительства" in idx:
                reg_address = str(row[idx["Регистрация по месту жительства"]] or "").strip() or None

            temporary_address = None
            if "Регистрация по месту пребывания" in idx:
                temporary_address = str(row[idx["Регистрация по месту пребывания"]] or "").strip() or None

            actual_address = None
            if "Адрес фактического проживания" in idx:
                actual_address = str(row[idx["Адрес фактического проживания"]] or "").strip() or None

            child = Child(
                last_name=last_name,
                first_name=first_name,
                middle_name=middle_name,
                birth_date=birth_date,
                gender=gender,
                education_form=education_form,
                reg_address=reg_address,
                temporary_address=temporary_address,
                actual_address=actual_address,
            )

            db.session.add(child)
            db.session.flush()

            social = _get_or_create_social(child)

            if "Вариант АООП" in idx:
                social.aoop_variant_text = str(row[idx["Вариант АООП"]] or "").strip() or None
                apply_aoop_to_child(child, social, social.aoop_variant_text)

            if "На ВШУ с" in idx:
                social.vshu_since = parse_birth(row[idx["На ВШУ с"]])

            if "Основание(я) постановки на ВШУ" in idx:
                social.vshu_reason = str(row[idx["Основание(я) постановки на ВШУ"]] or "").strip() or None

            if "На учете КДН с" in idx:
                social.kdn_since = parse_birth(row[idx["На учете КДН с"]])

            if "Основание(я) постановки на учет КДН" in idx:
                social.kdn_reason = str(row[idx["Основание(я) постановки на учет КДН"]] or "").strip() or None

            if "На учете ПДН с" in idx:
                social.pdn_since = parse_birth(row[idx["На учете ПДН с"]])

            if "Основание(я) постановки на учет ПДН" in idx:
                social.pdn_reason = str(row[idx["Основание(я) постановки на учет ПДН"]] or "").strip() or None

            if "Снят с ВШУ" in idx:
                social.vshu_removed_at = parse_birth(row[idx["Снят с ВШУ"]])

            if "Основание снятия с ВШУ" in idx:
                social.vshu_remove_reason = str(row[idx["Основание снятия с ВШУ"]] or "").strip() or None

            if class_name:

                sc = (
                    SchoolClass.query
                    .filter(
                        SchoolClass.academic_year_id == year.id,
                        SchoolClass.name == class_name
                    )
                    .first()
                )

                if not sc:

                    g, l = split_class_name(class_name)

                    sc = SchoolClass(
                        academic_year_id=year.id,
                        name=class_name,
                        grade=g,
                        letter=l,
                        max_students=25,
                    )

                    db.session.add(sc)
                    db.session.flush()

                en = ChildEnrollment(
                    child_id=child.id,
                    academic_year_id=year.id,
                    school_class_id=sc.id,
                    status="ACTIVE"
                )

                db.session.add(en)

            created += 1

        db.session.commit()

        flash(f"Импорт завершён. Добавлено детей: {created}, пропущено строк: {skipped}", "success")

        return redirect(url_for("children.list_children"))

    return render_template("children_import.html")

@children_bp.route("/classes/<int:class_id>/update", methods=["POST"])
@require_roles("ADMIN")
def update_class(class_id: int):
    c = SchoolClass.query.get_or_404(class_id)
    wants_json = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    old_teacher_user_id = c.teacher_user_id

    max_students = request.form.get("max_students", type=int)
    if max_students is None or max_students < 1:
        message = "Максимальное количество учеников должно быть больше нуля."
        if wants_json:
            return jsonify({"ok": False, "error": message}), 400
        flash(message, "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(c.academic_year_id),
        ))

    applications_count = request.form.get("applications_count", type=int)
    if applications_count is None:
        applications_count = 0
    if applications_count < 0:
        message = "Количество заявлений не может быть отрицательным."
        if wants_json:
            return jsonify({"ok": False, "error": message}), 400
        flash(message, "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(c.academic_year_id),
        ))

    name = normalize_class_name(request.form.get("name"))
    if not name:
        message = "Укажите название класса."
        if wants_json:
            return jsonify({"ok": False, "error": message}), 400
        flash(message, "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(c.academic_year_id),
        ))
    if _class_name_exists(
        c.academic_year_id,
        name,
        exclude_class_id=c.id,
    ):
        db.session.rollback()
        message = (
            f"Класс «{name}» уже существует в этом учебном году. "
            "Названия и буквы классов не должны совпадать."
        )
        if wants_json:
            return jsonify({"ok": False, "error": message}), 409
        flash(message, "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(c.academic_year_id),
        ))

    building_id = request.form.get("building_id")
    teacher_user_id = request.form.get("teacher_user_id")
    c.building_id = (
        int(building_id)
        if building_id and str(building_id).isdigit()
        else None
    )
    c.max_students = max_students
    c.applications_count = applications_count
    c.teacher_user_id = (
        int(teacher_user_id)
        if teacher_user_id and teacher_user_id.isdigit()
        else None
    )
    c.name = name
    g, l = split_class_name(name)
    c.grade = g
    c.letter = l

    try:
        db.session.flush()
        _sync_class_teacher_role(old_teacher_user_id)
        _sync_class_teacher_role(c.teacher_user_id)
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        message = "Такой класс уже существует в выбранном учебном году."
        if wants_json:
            return jsonify({"ok": False, "error": message}), 409
        flash(message, "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(c.academic_year_id),
        ))
    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    if wants_json:
        teacher = db.session.get(User, c.teacher_user_id) if c.teacher_user_id else None
        building = db.session.get(Building, c.building_id) if c.building_id else None
        return jsonify({
            "ok": True,
            "school_class": {
                "id": c.id,
                "name": c.name,
                "max_students": c.max_students,
                "applications_count": c.applications_count,
                "building_id": c.building_id,
                "building_name": building.name if building else None,
                "teacher_user_id": c.teacher_user_id,
                "teacher_fio": teacher.fio if teacher else None,
                "teacher_phone": teacher.phone if teacher else None,
            },
        })
    flash("Сохранено", "success")
    return redirect(url_for(
        "children.classes_registry",
        **_class_registry_return_args(c.academic_year_id),
    ))


def _normalized_class_registry_filters(level, grade):
    selected_level = (level or "").strip().upper()
    if selected_level not in CLASS_REGISTRY_LEVEL_GRADES:
        selected_level = ""

    try:
        selected_grade = int(grade) if grade not in (None, "") else None
    except (TypeError, ValueError):
        selected_grade = None
    if selected_grade not in range(1, 12):
        selected_grade = None
    if (
        selected_level
        and selected_grade is not None
        and selected_grade not in CLASS_REGISTRY_LEVEL_GRADES[selected_level]
    ):
        selected_grade = None
    return selected_level, selected_grade


def _class_registry_return_args(academic_year_id):
    selected_level, selected_grade = _normalized_class_registry_filters(
        request.form.get("return_level"),
        request.form.get("return_grade"),
    )
    q_text = (request.form.get("return_q") or "").strip()
    args = {"academic_year_id": academic_year_id}
    if selected_level:
        args["level"] = selected_level
    if selected_grade is not None:
        args["grade"] = selected_grade
    if q_text:
        args["q"] = q_text
    return args


@children_bp.route("/classes")
@require_roles("ADMIN")
def classes_registry():
    year_id = request.args.get("academic_year_id", type=int)
    year = AcademicYear.query.get(year_id) if year_id else _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.contingent"))

    all_years = AcademicYear.query.order_by(AcademicYear.start_date.desc().nullslast(), AcademicYear.name.desc()).all()
    q_text = (request.args.get("q") or "").strip()
    selected_level, selected_grade = _normalized_class_registry_filters(
        request.args.get("level"),
        request.args.get("grade"),
    )
    grade_options = (
        next(
            grades
            for code, _label, grades in CLASS_REGISTRY_LEVEL_OPTIONS
            if code == selected_level
        )
        if selected_level
        else tuple(range(1, 12))
    )
    last_copy = session.get("last_class_copy") or {}
    copy_undo = (
        last_copy
        if last_copy.get("target_year_id") == year.id
        else None
    )
    last_delete = session.get("last_class_delete") or {}
    delete_undo = None
    if last_delete.get("academic_year_id") == year.id:
        delete_token = last_delete.get("token")
        deleted_records = (
            cache.get(make_key("deleted_classes", delete_token))
            if delete_token else None
        )
        if deleted_records:
            delete_undo = last_delete
        else:
            session.pop("last_class_delete", None)
    older_years = [
        candidate
        for candidate in all_years
        if candidate.id != year.id
        and (
            candidate.start_date is None
            or year.start_date is None
            or candidate.start_date < year.start_date
        )
    ]
    copy_source_year = older_years[0] if older_years else None

    # Response-кеш для /classes (~4.6 МБ HTML, ~650 мс рендера из-за inline-форм редактирования).
    # Все ADMIN видят одинаковое (данных привязанных к user_id нет), поэтому общий ключ.
    cache_key = make_key(
        "classes_registry",
        year.id,
        selected_level,
        selected_grade,
        q_text,
    )
    cached_html = (
        None
        if copy_undo or delete_undo
        else view_response_cache.get(cache_key)
    )
    if cached_html is not None:
        from flask import Response
        return Response(cached_html, mimetype="text/html; charset=utf-8")

    raw_query = SchoolClass.query.filter(SchoolClass.academic_year_id == year.id)

    if selected_level:
        raw_query = raw_query.filter(
            SchoolClass.grade.in_(
                CLASS_REGISTRY_LEVEL_GRADES[selected_level]
            )
        )
    if selected_grade is not None:
        raw_query = raw_query.filter(SchoolClass.grade == selected_grade)

    if q_text:
        like = f"%{q_text.lower()}%"
        # match по имени класса ИЛИ по ФИО классрука
        teacher_match = db.session.query(User.id).filter(
            db.or_(
                db.func.lower(User.last_name).like(like),
                db.func.lower(User.first_name).like(like),
                db.func.lower(User.middle_name).like(like),
            )
        )
        raw_query = raw_query.filter(
            db.or_(
                db.func.lower(SchoolClass.name).like(like),
                SchoolClass.teacher_user_id.in_(teacher_match),
            )
        )

    raw_classes = (
        raw_query
        .order_by(
            SchoolClass.grade.asc().nullslast(),
            SchoolClass.letter.asc().nullslast(),
            SchoolClass.name.asc()
        )
        .all()
    )

    buildings = Building.query.order_by(Building.name.asc()).all()
    teachers = User.query.order_by(User.last_name.asc(), User.first_name.asc()).all()

    teachers_map = {u.id: u for u in teachers}

    class_ids = [c.id for c in raw_classes]

    # Batch: active enrollment counts per class
    enrollment_counts = {}
    if class_ids:
        for class_id, cnt in db.session.query(
            ChildEnrollment.school_class_id, db.func.count(ChildEnrollment.id)
        ).filter(
            ChildEnrollment.school_class_id.in_(class_ids),
            ChildEnrollment.ended_at.is_(None)
        ).group_by(ChildEnrollment.school_class_id).all():
            enrollment_counts[class_id] = cnt

    # Batch: task counts per class
    task_total_map = {}
    task_overdue_map = {}
    if class_ids:
        for class_id, cnt in db.session.query(
            Task.class_id, db.func.count(Task.id)
        ).filter(Task.class_id.in_(class_ids)).group_by(Task.class_id).all():
            task_total_map[class_id] = cnt
        terminal = [Task.STATUS_DONE, Task.STATUS_CLOSED, Task.STATUS_CANCELLED]
        for class_id, cnt in db.session.query(
            Task.class_id, db.func.count(Task.id)
        ).filter(
            Task.class_id.in_(class_ids),
            Task.status.notin_(terminal),
            Task.deadline_at.isnot(None),
            Task.deadline_at < datetime.utcnow()
        ).group_by(Task.class_id).all():
            task_overdue_map[class_id] = cnt
    sc_in_club_by_class: dict[int, int] = {}
    do_in_class_by_class: dict[int, int] = {}

    classes = []
    for c in raw_classes:
        teacher = teachers_map.get(c.teacher_user_id)
        teacher_fio = teacher.fio if teacher else None
        teacher_phone = teacher.phone if teacher else None
        active_count = enrollment_counts.get(c.id, 0)
        classes.append((
            c, teacher_fio, teacher_phone, active_count,
            task_total_map.get(c.id, 0),
            task_overdue_map.get(c.id, 0),
            int(sc_in_club_by_class.get(c.id, 0)),
            int(do_in_class_by_class.get(c.id, 0)),
        ))

    html = render_template(
        "classes_list.html",
        classes=classes,
        teachers=teachers,
        buildings=buildings,
        year=year,
        all_years=all_years,
        q=q_text,
        level_options=CLASS_REGISTRY_LEVEL_OPTIONS,
        selected_level=selected_level,
        selected_grade=selected_grade,
        grade_options=grade_options,
        is_admin=is_admin(current_user),
        copy_source_year=copy_source_year,
        copy_undo=copy_undo,
        delete_undo=delete_undo,
    )
    if not copy_undo and not delete_undo:
        view_response_cache.set(cache_key, html, timeout=60)
    return html


@children_bp.route("/classes/<int:class_id>")
@login_required
def class_detail(class_id):
    """Страница класса: список учеников + блок Кубка (для ADMIN)."""
    if not has_permission("children_registry_view"):
        abort(403)

    sc = SchoolClass.query.get_or_404(class_id)

    if should_limit_children_to_own_class():
        if sc.teacher_user_id != current_user.id:
            abort(403)

    enrollments = (
        ChildEnrollment.query
        .join(Child, Child.id == ChildEnrollment.child_id)
        .options(joinedload(ChildEnrollment.child))
        .filter(
            ChildEnrollment.school_class_id == class_id,
            ChildEnrollment.status == "ACTIVE",
            ChildEnrollment.ended_at.is_(None),
            Child.status == "ACTIVE",
            ~_active_expel_exists(Child.id),
        )
        .all()
    )
    students = sorted(
        [e.child for e in enrollments if e.child],
        key=lambda c: ((c.last_name or "").lower(), (c.first_name or "").lower()),
    )

    teacher = User.query.get(sc.teacher_user_id) if sc.teacher_user_id else None

    contingent_year_id = (
        request.args.get("contingent_year_id", type=int)
        or sc.academic_year_id
    )
    contingent_building_id = request.args.get(
        "contingent_building_id",
        type=int,
    )
    contingent_args = {"year_id": contingent_year_id}
    student_return_args = {
        "return_class_id": sc.id,
        "contingent_year_id": contingent_year_id,
    }
    if contingent_building_id is not None:
        contingent_args["building_id"] = contingent_building_id
        student_return_args["contingent_building_id"] = (
            contingent_building_id
        )

    return render_template(
        "class_detail.html",
        school_class=sc,
        students=students,
        teacher=teacher,
        contingent_url=url_for(
            "children.contingent",
            **contingent_args,
        ),
        student_return_args=student_return_args,
    )


@children_bp.route("/classes/new", methods=["POST"])
@require_roles("ADMIN")
def classes_new():
    requested_year_id = request.form.get("academic_year_id", type=int)
    year = AcademicYear.query.get(requested_year_id) if requested_year_id else _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(year.id if year else None),
        ))

    name = normalize_class_name(request.form.get("name"))
    max_students = request.form.get("max_students", type=int) or 25
    teacher_user_id = request.form.get("teacher_user_id", type=int)
    building_id = request.form.get("building_id", type=int)

    if not name:
        flash("Укажите класс", "danger")
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(year.id if year else None),
        ))
    if _class_name_exists(year.id, name):
        flash(
            f"Класс «{name}» уже существует в учебном году {year.name}. "
            "Используйте другую букву или название.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(year.id),
        ))

    g, l = split_class_name(name)

    c = SchoolClass(
        academic_year_id=year.id,
        building_id=building_id,
        name=name,
        grade=g,
        letter=l,
        max_students=max_students,
        applications_count=0,
        teacher_user_id=teacher_user_id
    )

    try:
        db.session.add(c)
        db.session.flush()
        _sync_class_teacher_role(teacher_user_id)
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        flash(
            "Такой класс уже существует в выбранном учебном году.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            **_class_registry_return_args(year.id),
        ))
    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    flash("Класс добавлен", "success")
    return redirect(url_for(
        "children.classes_registry",
        **_class_registry_return_args(year.id),
    ))

# =========================================================
# REGISTRIES
# =========================================================
@children_bp.route("/registry/vshu")
@login_required
def registry_vshu():
    q, year = _children_base_query_for_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    q = q.outerjoin(ChildSocial, ChildSocial.child_id == Child.id).filter(
        db.or_(
            Child.is_vshu.is_(True),
            ChildSocial.vshu_since.isnot(None)
        )
    ).filter(
        db.or_(
            ChildSocial.vshu_removed_at.is_(None),
            Child.is_vshu.is_(True)
        )
    )

    if filters["selected_grade"] is not None:
        q = q.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        q = q.filter(SchoolClass.id == filters["selected_class_id"])

    children = (
        q.options(joinedload(Child.parent_links).joinedload(ChildParent.parent))
        .order_by(SchoolClass.name.asc(), Child.last_name.asc(), Child.first_name.asc())
        .all()
    )

    if filters["q_text"]:
        children = [ch for ch in children if _match_fio_query(ch, filters["q_text"])]

    page, per_page = resolve_pagination()
    children_page, pagination = paginate_list(children, page=page, per_page=per_page)

    return render_template(
        "registry_children.html",
        title="Реестр ВШУ",
        children=children_page,
        pagination=pagination,
        q_text=filters["q_text"],
        classes=filters["classes"],
        grades=filters["grades"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        export_url=url_for("children.registry_vshu_export", grade=filters["selected_grade_raw"], class_id=filters["selected_class_id"], q=filters["q_text"])
    )


@children_bp.route("/registry/vshu/export")
@login_required
def registry_vshu_export():
    q, year = _children_base_query_for_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    q = q.outerjoin(ChildSocial, ChildSocial.child_id == Child.id).filter(
        db.or_(
            Child.is_vshu.is_(True),
            ChildSocial.vshu_since.isnot(None)
        )
    ).filter(
        db.or_(
            ChildSocial.vshu_removed_at.is_(None),
            Child.is_vshu.is_(True)
        )
    )
    if filters["selected_grade"] is not None:
        q = q.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        q = q.filter(SchoolClass.id == filters["selected_class_id"])

    children = (
        q.options(joinedload(Child.parent_links).joinedload(ChildParent.parent))
        .order_by(Child.last_name.asc(), Child.first_name.asc())
        .all()
    )
    if filters["q_text"]:
        children = [ch for ch in children if _match_fio_query(ch, filters["q_text"])]

    return _export_children_xlsx("Реестр_ВШУ", children)


@children_bp.route("/registry/ovz")
@login_required
def registry_ovz():
    q, year = _children_base_query_for_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    q = q.filter(Child.is_ovz.is_(True))
    if filters["selected_grade"] is not None:
        q = q.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        q = q.filter(SchoolClass.id == filters["selected_class_id"])

    children = q.order_by(SchoolClass.name.asc(), Child.last_name.asc(), Child.first_name.asc()).all()
    if filters["q_text"]:
        children = [ch for ch in children if _match_fio_query(ch, filters["q_text"])]

    page, per_page = resolve_pagination()
    children_page, pagination = paginate_list(children, page=page, per_page=per_page)

    return render_template(
        "registry_ovz.html",
        title="Реестр ОВЗ",
        children=children_page,
        pagination=pagination,
        q_text=filters["q_text"],
        classes=filters["classes"],
        grades=filters["grades"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        export_url=url_for("children.registry_ovz_export", grade=filters["selected_grade_raw"], class_id=filters["selected_class_id"], q=filters["q_text"])
    )


@children_bp.route("/registry/ovz/export")
@login_required
def registry_ovz_export():
    q, year = _children_base_query_for_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    q = q.filter(Child.is_ovz.is_(True))
    if filters["selected_grade"] is not None:
        q = q.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        q = q.filter(SchoolClass.id == filters["selected_class_id"])

    children = q.order_by(Child.last_name.asc(), Child.first_name.asc()).all()
    if filters["q_text"]:
        children = [ch for ch in children if _match_fio_query(ch, filters["q_text"])]

    return _export_children_xlsx("Реестр_ОВЗ", children)


@children_bp.route("/registry/az")
@login_required
def registry_az():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())

    debts = (
        Debt.query
        .join(Child, Debt.child_id == Child.id)
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == year.id if year else True)
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .filter(Debt.status == "OPEN")
    )

    if filters["selected_grade"] is not None:
        debts = debts.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        debts = debts.filter(SchoolClass.id == filters["selected_class_id"])

    debts = debts.order_by(
        SchoolClass.name.asc(),
        Child.last_name.asc(),
        Child.first_name.asc(),
        Debt.education_activity_id.asc(),
    ).all()

    m = {}
    for d in debts:
        ch = d.child
        if not _match_fio_query(ch, filters["q_text"]):
            continue

        if ch.id not in m:
            m[ch.id] = {"child": ch, "subjects": []}
        subj = d.subject_name
        if subj and subj not in m[ch.id]["subjects"]:
            m[ch.id]["subjects"].append(subj)

    for item in m.values():
        item["subjects"].sort(key=lambda value: value.casefold())
    rows = list(m.values())

    page, per_page = resolve_pagination()
    rows_page, pagination = paginate_list(rows, page=page, per_page=per_page)

    return render_template(
        "registry_az.html",
        title="Реестр АЗ (открытые задолженности)",
        rows=rows_page,
        pagination=pagination,
        q_text=filters["q_text"],
        classes=filters["classes"],
        grades=filters["grades"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        export_url=url_for("children.registry_az_export", grade=filters["selected_grade_raw"], class_id=filters["selected_class_id"], q=filters["q_text"])
    )


@children_bp.route("/registry/az/export")
@login_required
def registry_az_export():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())

    debts = (
        Debt.query
        .join(Child, Debt.child_id == Child.id)
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == year.id if year else True)
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .filter(Debt.status == "OPEN")
    )

    if filters["selected_grade"] is not None:
        debts = debts.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        debts = debts.filter(SchoolClass.id == filters["selected_class_id"])

    debts = debts.order_by(
        Child.last_name.asc(),
        Child.first_name.asc(),
        Debt.education_activity_id.asc(),
    ).all()

    m = {}
    for d in debts:
        ch = d.child
        if not _match_fio_query(ch, filters["q_text"]):
            continue
        if ch.id not in m:
            m[ch.id] = {"child": ch, "subjects": []}
        subj = d.subject_name
        if subj and subj not in m[ch.id]["subjects"]:
            m[ch.id]["subjects"].append(subj)

    for item in m.values():
        item["subjects"].sort(key=lambda value: value.casefold())

    wb = Workbook()
    ws = wb.active
    ws.title = "АЗ"
    ws.append(["№", "ФИО", "Класс", "Предметы (открытые)"])

    for idx, item in enumerate(m.values(), start=1):
        ch = item["child"]
        ws.append([
            idx,
            ch.fio,
            ch.current_class_name or "—",
            ", ".join(item["subjects"])
        ])

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="Реестр_АЗ.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@children_bp.route("/registry/enrolled")
@login_required
def registry_enrolled():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    ens = (
        ChildEnrollment.query
        .join(Child, ChildEnrollment.child_id == Child.id)
        .join(SchoolClass, ChildEnrollment.school_class_id == SchoolClass.id)
        .options(
            contains_eager(ChildEnrollment.child),
            contains_eager(ChildEnrollment.school_class),
        )
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.status == "ACTIVE"
        )
    )

    if filters["selected_grade"] is not None:
        ens = ens.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        ens = ens.filter(SchoolClass.id == filters["selected_class_id"])

    ens = ens.order_by(SchoolClass.name.asc(), Child.last_name.asc(), Child.first_name.asc()).all()

    rows = []
    for en in ens:
        ch = en.child
        if not _match_fio_query(ch, filters["q_text"]):
            continue
        rows.append({
            "child": ch,
            "class_name": en.school_class.name if en.school_class else None,
            "en": en
        })

    return render_template(
        "registry_enrolled.html",
        title=f"Реестр зачисленных ({year.name})",
        rows=rows,
        q_text=filters["q_text"],
        classes=filters["classes"],
        grades=filters["grades"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        export_url=url_for("children.registry_enrolled_export", grade=filters["selected_grade_raw"], class_id=filters["selected_class_id"], q=filters["q_text"])
    )


@children_bp.route("/registry/enrolled/export")
@login_required
def registry_enrolled_export():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    ens = (
        ChildEnrollment.query
        .join(Child, ChildEnrollment.child_id == Child.id)
        .join(SchoolClass, ChildEnrollment.school_class_id == SchoolClass.id)
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.status == "ACTIVE"
        )
    )

    if filters["selected_grade"] is not None:
        ens = ens.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        ens = ens.filter(SchoolClass.id == filters["selected_class_id"])

    ens = ens.order_by(Child.last_name.asc(), Child.first_name.asc()).all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Зачисленные"
    ws.append(["№", "ФИО", "Класс", "Дата зачисления"])

    row_num = 1
    for en in ens:
        ch = en.child
        if not _match_fio_query(ch, filters["q_text"]):
            continue
        ws.append([
            row_num,
            ch.fio,
            en.school_class.name if en.school_class else "",
            en.enrolled_at.strftime("%d.%m.%Y %H:%M") if en.enrolled_at else ""
        ])
        row_num += 1

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="Реестр_зачисленных.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@children_bp.route("/registry/expelled")
@login_required
def registry_expelled():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())

    events = _active_expel_events_query()

    if filters["selected_grade"] is not None:
        grade_classes = [c.name for c in filters["classes"]]
        if grade_classes:
            events = events.filter(ChildEvent.from_class.in_(grade_classes))
        else:
            events = events.filter(db.text("1=0"))
    if filters["selected_class_name"]:
        events = events.filter(ChildEvent.from_class == filters["selected_class_name"])

    events = events.order_by(ChildEvent.created_at.desc()).all()

    # s85: батч-загрузка детей вместо .query.get() в цикле
    child_ids = list({ev.child_id for ev in events if ev.child_id})
    children_by_id = {
        c.id: c
        for c in Child.query.filter(Child.id.in_(child_ids)).all()
    } if child_ids else {}

    rows = []
    for ev in events:
        ch = children_by_id.get(ev.child_id)
        if not ch:
            continue
        if not _match_fio_query(ch, filters["q_text"]):
            continue
        rows.append({"child": ch, "ev": ev})

    return render_template(
        "registry_expelled.html",
        title="Реестр отчисленных",
        rows=rows,
        q_text=filters["q_text"],
        classes=filters["classes"],
        grades=filters["grades"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        export_url=url_for("children.registry_expelled_export", grade=filters["selected_grade_raw"], class_id=filters["selected_class_id"], q=filters["q_text"])
    )


@children_bp.route("/registry/expelled/export")
@login_required
def registry_expelled_export():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())

    events = _active_expel_events_query()

    if filters["selected_grade"] is not None:
        grade_classes = [c.name for c in filters["classes"]]
        if grade_classes:
            events = events.filter(ChildEvent.from_class.in_(grade_classes))
        else:
            events = events.filter(db.text("1=0"))
    if filters["selected_class_name"]:
        events = events.filter(ChildEvent.from_class == filters["selected_class_name"])

    events = events.order_by(ChildEvent.created_at.desc()).all()

    # s85: батч-загрузка детей
    child_ids = list({ev.child_id for ev in events if ev.child_id})
    children_by_id = {
        c.id: c
        for c in Child.query.filter(Child.id.in_(child_ids)).all()
    } if child_ids else {}

    wb = Workbook()
    ws = wb.active
    ws.title = "Отчисленные"
    ws.append(["№", "ФИО", "Класс (откуда)", "Дата", "Причина/основание", "Куда"])

    row_num = 1
    for ev in events:
        ch = children_by_id.get(ev.child_id)
        if not ch:
            continue
        if not _match_fio_query(ch, filters["q_text"]):
            continue

        ws.append([
            row_num,
            ch.fio,
            ev.from_class or "",
            ev.created_at.strftime("%d.%m.%Y %H:%M") if ev.created_at else "",
            ev.reason or "",
            ev.to_class or "",
        ])
        row_num += 1

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="Реестр_отчисленных.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# =========================================================
# INCIDENTS
# =========================================================
@children_bp.route("/incidents/new", methods=["GET", "POST"])
@login_required
def incident_new():
    if not has_permission("incident_add"):
        abort(403)
    if request.method == "POST":
        occurred_date = (request.form.get("occurred_date") or "").strip()
        occurred_hour = (request.form.get("occurred_hour") or "").strip()
        occurred_minute = (request.form.get("occurred_minute") or "").strip()
        occurred_time = f"{occurred_hour}:{occurred_minute}"
        category = (request.form.get("category") or "").strip()
        description = (request.form.get("description") or "").strip() or None
        initial_work = (request.form.get("initial_work") or "").strip() or None

        raw_ids = request.form.getlist("child_ids")
        child_ids = []
        for x in raw_ids:
            if str(x).isdigit():
                child_ids.append(int(x))
        child_ids = list(dict.fromkeys(child_ids))

        def _incident_new_error(msg):
            flash(msg, "danger")
            participants = []
            try:
                if child_ids:
                    year = _get_current_year()
                    by_class = {}
                    for cid in child_ids:
                        ch = Child.query.get(cid)
                        if not ch:
                            continue
                        en = None
                        if year:
                            en = (
                                ChildEnrollment.query
                                .filter(ChildEnrollment.child_id == ch.id)
                                .filter(ChildEnrollment.academic_year_id == year.id)
                                .filter(ChildEnrollment.status == "ACTIVE")
                                .order_by(ChildEnrollment.id.desc())
                                .first()
                            )
                        sc = SchoolClass.query.get(en.school_class_id) if (en and en.school_class_id) else None
                        key = sc.id if sc else 0
                        if key not in by_class:
                            by_class[key] = {
                                "grade": sc.grade if sc else None,
                                "class_id": sc.id if sc else None,
                                "child_ids": [],
                            }
                        by_class[key]["child_ids"].append(ch.id)
                    participants = [v for v in by_class.values() if v.get("class_id")]
            except Exception:
                participants = []
            return render_template(
                "incident_new.html",
                categories=INCIDENT_CATEGORIES,
                preselected_student=None,
                form_data={
                    "occurred_date": occurred_date,
                    "occurred_hour": occurred_hour,
                    "occurred_minute": occurred_minute,
                    "category": category,
                    "description": description or "",
                    "initial_work": initial_work or "",
                },
                preselected_participants=participants,
            )

        if not occurred_date or not occurred_time:
            return _incident_new_error("Укажите дату и время")

        if category not in INCIDENT_CATEGORIES:
            return _incident_new_error("Выберите категорию инцидента")

        if not description:
            return _incident_new_error("Заполните описание инцидента")

        if not child_ids:
            return _incident_new_error("Добавьте хотя бы одного ребёнка")

        try:
            occurred_at = datetime.strptime(f"{occurred_date} {occurred_time}", "%Y-%m-%d %H:%M")
        except Exception:
            return _incident_new_error("Неверный формат даты/времени")

        now = _now_msk_naive()
        if occurred_at < now - timedelta(hours=48):
            return _incident_new_error("Дата инцидента не должна быть раньше, чем за 48 часов до момента подачи заявки")
        if occurred_at > now + timedelta(minutes=5):
            return _incident_new_error("Дата инцидента не может быть в будущем")

        inc = Incident(
            occurred_at=occurred_at,
            category=category,
            description=description,
            status="new",
            author_id=getattr(current_user, "id", None),
            created_at=datetime.utcnow(),
        )
        db.session.add(inc)
        db.session.flush()

        for cid in child_ids:
            ch = Child.query.get(cid)
            if ch:
                db.session.add(IncidentChild(incident_id=inc.id, child_id=ch.id))

        # Проделанная автором работа на момент подачи — опциональный первый
        # пункт журнала работы. Идёт в IncidentNote с префиксом [Сделано автором].
        if initial_work:
            from app.models_legacy import IncidentNote
            db.session.add(IncidentNote(
                incident_id=inc.id,
                author_id=getattr(current_user, "id", None),
                text=f"[Сделано автором] {initial_work}",
            ))

        db.session.commit()
        flash("Инцидент сохранён", "success")
        return redirect(url_for("children.incidents_my"))

    preselected_student = None
    student_id = request.args.get("student_id", type=int)
    if student_id:
        try:
            ch = Child.query.get(student_id)
        except Exception:
            ch = None
        if ch:
            year = _get_current_year()
            en = None
            if year:
                en = (
                    ChildEnrollment.query
                    .filter(ChildEnrollment.child_id == ch.id)
                    .filter(ChildEnrollment.academic_year_id == year.id)
                    .filter(ChildEnrollment.status == "ACTIVE")
                    .order_by(ChildEnrollment.id.desc())
                    .first()
                )
            sc = None
            if en and en.school_class_id:
                sc = SchoolClass.query.get(en.school_class_id)
            preselected_student = {
                "child_id": ch.id,
                "fio": f"{ch.last_name or ''} {ch.first_name or ''} {ch.middle_name or ''}".strip(),
                "grade": sc.grade if sc else None,
                "class_id": sc.id if sc else None,
                "class_name": sc.name if sc else None,
            }

    return render_template(
        "incident_new.html",
        categories=INCIDENT_CATEGORIES,
        preselected_student=preselected_student,
        form_data=None,
        preselected_participants=None,
    )


@children_bp.route("/api/classes/by-grade")
@login_required
def api_classes_by_grade():
    if not has_permission("incident_add"):
        abort(403)
    grade = request.args.get("grade", type=int)
    if not grade:
        return jsonify([])

    year = _get_current_year()
    if not year:
        return jsonify([])

    q = SchoolClass.query.filter(
        SchoolClass.academic_year_id == year.id,
        SchoolClass.grade == grade
    )

    classes = q.order_by(SchoolClass.name.asc()).all()
    return jsonify([{"id": c.id, "name": c.name} for c in classes])


@children_bp.route("/api/children/by-class")
@login_required
def api_children_by_class():
    if not has_permission("incident_add"):
        abort(403)
    class_id = request.args.get("class_id", type=int)
    if not class_id:
        return jsonify([])

    year = _get_current_year()
    if not year:
        return jsonify([])

    ens = (
        ChildEnrollment.query
        .join(Child, ChildEnrollment.child_id == Child.id)
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.school_class_id == class_id,
            ChildEnrollment.ended_at.is_(None)
        )
        .order_by(Child.last_name.asc(), Child.first_name.asc(), Child.middle_name.asc())
        .all()
    )

    return jsonify([{"id": en.child.id, "fio": en.child.fio} for en in ens])


_ROLE_LABELS = {
    "ADMIN": "Администратор",
    "CLASS_TEACHER": "Классный руководитель",
    "TEACHER": "Учитель",
    "PSYCHOLOGIST": "Психолог",
    "SOCIAL_PEDAGOG": "Социальный педагог",
    "METHODIST": "Методист",
    "KPP": "КПП",
    "VIEWER": "Наблюдатель",
}


def _get_author_label(user):
    """Return human-readable 'ФИО, должность [класс]' for incident author.

    Кэш живёт ровно один HTTP-запрос (flask.g) — переименование/смена роли
    видны сразу, а в пределах одного рендера реестра одного автора не считаем
    повторно.
    """
    if not user:
        return "—"
    uid = user.id
    try:
        cache = g._author_label_cache  # type: ignore[attr-defined]
    except (AttributeError, RuntimeError):
        cache = None
    if cache is not None and uid in cache:
        return cache[uid]
    fio = (user.fio or "").strip() or "—"
    role = getattr(user, "role", None)
    if role == "CLASS_TEACHER":
        cls = (
            SchoolClass.query
            .filter_by(teacher_user_id=user.id)
            .order_by(SchoolClass.academic_year_id.desc())
            .first()
        )
        if cls:
            label = f"{fio}, кл. рук. {cls.name}"
        else:
            label = f"{fio}, классный руководитель"
    else:
        role_label = _ROLE_LABELS.get(role, role or "")
        label = f"{fio}, {role_label}" if role_label else fio
    if cache is None:
        try:
            cache = {}
            g._author_label_cache = cache  # type: ignore[attr-defined]
        except RuntimeError:
            return label
    cache[uid] = label
    return label


def _can_change_status():
    """Управлять ходом инцидента (смена статуса, назначение исполнителя,
    принудительное закрытие) могут: ADMIN, DEPUTY_DIRECTOR, SOCIAL_PEDAGOG.
    SOCIAL_PEDAGOG получили права директорским решением (s87) — они активно
    ведут инциденты и должны менять статусы и назначать исполнителей наравне
    с управляющими. PSYCHOLOGIST/METHODIST остаются read-only."""
    return has_role("ADMIN") or has_role("DEPUTY_DIRECTOR") or has_role("SOCIAL_PEDAGOG")


def _can_mark_resolved(inc):
    """Пометить инцидент как «Отработан» может любой из назначенных исполнителей
    или управляющий (ADMIN/DEPUTY_DIRECTOR)."""
    if _can_change_status():
        return True
    uid = getattr(current_user, "id", None)
    return _uid_is_assignee(inc, uid)


# =========================================================
# IncidentNote attachments (файлы к комментариям инцидента)
# =========================================================
INC_NOTE_ATT_ALLOWED = {
    'pdf', 'doc', 'docx', 'xls', 'xlsx',
    'jpg', 'jpeg', 'png', 'gif', 'webp',
    'zip', 'txt',
    'mp4', 'mov', 'webm',
}
INC_NOTE_ATT_MAX_SIZE = 100 * 1024 * 1024  # 100 МБ — потолок (видео)
INC_NOTE_ATT_MAX_SIZE_NONVIDEO = 30 * 1024 * 1024  # документы/картинки 30 МБ
INC_NOTE_VIDEO_EXT = {'mp4', 'mov', 'webm'}
INC_NOTE_ATT_MAX_FILES = 10

# Whitelist реальных MIME-типов для вложений из MAX-бота. Проверяется через
# libmagic по сигнатуре файла, не по расширению/Content-Type из бота.
_MAX_MIME_WHITELIST = {
    "image/jpeg", "image/png", "image/gif", "image/webp",
    "application/pdf",
    "video/mp4", "video/quicktime", "video/webm", "video/x-matroska",
    "application/zip", "application/x-zip-compressed",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
}
_MAX_MIME_TO_EXT = {
    "image/jpeg": "jpg", "image/png": "png", "image/gif": "gif", "image/webp": "webp",
    "application/pdf": "pdf",
    "video/mp4": "mp4", "video/quicktime": "mov", "video/webm": "webm",
    "application/zip": "zip", "application/x-zip-compressed": "zip",
    "application/msword": "doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.ms-excel": "xls",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "text/plain": "txt",
}


def _inc_note_upload_root():
    from flask import current_app as _ca
    from pathlib import Path as _Path
    root = _ca.config.get("UPLOAD_FOLDER") or os.path.abspath(os.path.join("data", "uploads"))
    path = os.path.join(root, "incident_notes")
    os.makedirs(path, exist_ok=True)
    return root, path


def _inc_note_abs_path(stored_rel_path):
    if not stored_rel_path:
        return ""
    if os.path.isabs(stored_rel_path):
        return stored_rel_path
    from flask import current_app as _ca
    return os.path.join(_ca.config.get("UPLOAD_FOLDER") or os.path.abspath(os.path.join("data", "uploads")), stored_rel_path)


def _save_incident_note_attachments(note, files):
    """Сохраняет загруженные файлы в каталог uploads/incident_notes/<note_id>/
    и пишет строки в incident_note_attachment. Возвращает список сохранённых."""
    from pathlib import Path as _Path
    from werkzeug.utils import secure_filename as _sfn
    import uuid as _uuid

    files = [f for f in (files or []) if f and getattr(f, "filename", "")]
    if not files:
        return []
    if len(files) > INC_NOTE_ATT_MAX_FILES:
        raise ValueError(f"К одному комментарию можно прикрепить не более {INC_NOTE_ATT_MAX_FILES} файлов.")

    upload_root, _ = _inc_note_upload_root()
    note_dir = os.path.join(upload_root, "incident_notes", str(note.id))
    os.makedirs(note_dir, exist_ok=True)

    saved = []
    written_paths = []  # физические файлы, чтобы откатить при ошибке на середине
    try:
        for storage in files:
            original_name = (storage.filename or "").strip()
            safe_name = _sfn(original_name) or "file"
            ext = (_Path(safe_name).suffix or "").lower().lstrip(".")
            if ext not in INC_NOTE_ATT_ALLOWED:
                raise ValueError(f"Формат файла не поддерживается: {original_name}")

            storage.stream.seek(0, os.SEEK_END)
            size = storage.stream.tell()
            storage.stream.seek(0)
            max_size = INC_NOTE_ATT_MAX_SIZE if ext in INC_NOTE_VIDEO_EXT else INC_NOTE_ATT_MAX_SIZE_NONVIDEO
            if size > max_size:
                limit_mb = max_size // (1024 * 1024)
                raise ValueError(f"Файл {original_name} превышает ограничение {limit_mb} МБ.")

            stored_filename = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{_uuid.uuid4().hex[:12]}.{ext}" if ext else f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{_uuid.uuid4().hex[:12]}"
            abs_path = os.path.join(note_dir, stored_filename)
            storage.save(abs_path)
            written_paths.append(abs_path)
            rel_path = os.path.relpath(abs_path, upload_root).replace("\\", "/")

            row = IncidentNoteAttachment(
                note_id=note.id,
                filename=original_name,
                stored_filename=stored_filename,
                file_path=rel_path,
                content_type=(storage.mimetype or None),
                file_size=size,
                uploaded_by_user_id=getattr(current_user, "id", None),
            )
            db.session.add(row)
            saved.append(row)
    except Exception:
        # Если что-то упало — удалить уже сохранённые физические файлы.
        # Транзакцию БД откатывает вызывающий (см. mark_resolved/add_note).
        for p in written_paths:
            try: os.remove(p)
            except OSError: pass
        raise
    return saved


def _save_max_attachment(note, original_name, raw_bytes, hint_mime=None):
    """Сохраняет байты вложения из MAX-бота в incident_note_attachment.
    Валидация типа через libmagic (whitelist `_MAX_MIME_WHITELIST`).
    Возвращает строку IncidentNoteAttachment или None при отказе.
    Не делает commit — вызывающий собирает в одну транзакцию.
    """
    from pathlib import Path as _Path
    from werkzeug.utils import secure_filename as _sfn
    import uuid as _uuid

    if not raw_bytes:
        return None

    size = len(raw_bytes)
    if size > INC_NOTE_ATT_MAX_SIZE:
        raise ValueError(f"Вложение {original_name}: {size} байт превышает лимит")

    detected_mime = None
    try:
        import magic  # python-magic (Linux: libmagic1) или python-magic-bin (Windows)
        detected_mime = magic.from_buffer(raw_bytes[:8192], mime=True)
    except Exception:
        detected_mime = (hint_mime or "").split(";")[0].strip() or None

    if not detected_mime or detected_mime not in _MAX_MIME_WHITELIST:
        raise ValueError(f"Тип файла не разрешён: {detected_mime!r}")

    ext = _MAX_MIME_TO_EXT.get(detected_mime, "bin")
    if ext not in INC_NOTE_VIDEO_EXT and size > INC_NOTE_ATT_MAX_SIZE_NONVIDEO:
        raise ValueError(f"Не-видео файл превышает лимит {INC_NOTE_ATT_MAX_SIZE_NONVIDEO // (1024*1024)} МБ")

    safe_name = _sfn(original_name or "") or f"attachment.{ext}"
    if not safe_name.lower().endswith("." + ext):
        safe_name = f"{_Path(safe_name).stem or 'attachment'}.{ext}"

    upload_root, _ = _inc_note_upload_root()
    note_dir = os.path.join(upload_root, "incident_notes", str(note.id))
    os.makedirs(note_dir, exist_ok=True)

    stored_filename = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{_uuid.uuid4().hex[:12]}.{ext}"
    abs_path = os.path.join(note_dir, stored_filename)
    with open(abs_path, "wb") as f:
        f.write(raw_bytes)
    rel_path = os.path.relpath(abs_path, upload_root).replace("\\", "/")

    row = IncidentNoteAttachment(
        note_id=note.id,
        filename=safe_name,
        stored_filename=stored_filename,
        file_path=rel_path,
        content_type=detected_mime,
        file_size=size,
        uploaded_by_user_id=None,
    )
    db.session.add(row)
    return row


def _inc_note_att_to_dict(att):
    ext = (att.filename or "").rsplit(".", 1)[-1].lower() if "." in (att.filename or "") else ""
    kind = "video" if ext in INC_NOTE_VIDEO_EXT else ("image" if ext in {"jpg","jpeg","png","gif","webp"} else "file")
    return {
        "id": att.id,
        "filename": att.filename,
        "size": att.file_size or 0,
        "kind": kind,
        "content_type": att.content_type or "",
        "url": url_for("children.incident_note_attachment_download", attachment_id=att.id),
    }


@children_bp.route("/incidents/note-attachments/<int:attachment_id>/download")
@login_required
def incident_note_attachment_download(attachment_id):
    att = IncidentNoteAttachment.query.get_or_404(attachment_id)
    note = att.note
    if not note:
        abort(404)
    inc = note.incident
    if not inc or not _can_view_incident(inc):
        abort(403)
    abs_path = _inc_note_abs_path(att.file_path)
    if not abs_path or not os.path.exists(abs_path):
        abort(404)
    mime = att.content_type or mimetypes.guess_type(att.filename)[0] or "application/octet-stream"
    # Для видео/картинок — inline, для остального — attachment.
    ext = (att.filename or "").rsplit(".", 1)[-1].lower() if "." in (att.filename or "") else ""
    inline = ext in INC_NOTE_VIDEO_EXT or ext in {"jpg","jpeg","png","gif","webp"} or ext == "pdf"
    return send_file(
        abs_path,
        mimetype=mime,
        as_attachment=(not inline),
        download_name=att.filename,
    )


@children_bp.route("/incidents/note-attachments/<int:attachment_id>/delete", methods=["POST"])
@login_required
def incident_note_attachment_delete(attachment_id):
    att = IncidentNoteAttachment.query.get_or_404(attachment_id)
    note = att.note
    if not note:
        abort(404)
    inc = note.incident
    if not inc:
        abort(404)
    # Удалить может: автор комментария или управляющий (ADMIN/DEPUTY).
    uid = getattr(current_user, "id", None)
    if not (_can_change_status() or (uid and note.author_id == uid)):
        return jsonify({"error": "forbidden"}), 403

    abs_path = _inc_note_abs_path(att.file_path)
    try:
        if abs_path and os.path.exists(abs_path):
            os.remove(abs_path)
    except Exception:
        pass
    db.session.delete(att)
    db.session.commit()
    return jsonify({"ok": True})


# Классификация типов уведомлений для гибких режимов подписки
# (см. User.notify_incident_mode и /profile/notifications).
#   open    — инцидент появился у пользователя (вам назначили / автору: назначен исполнитель)
#   status  — смена статуса (не финальная)
#   close   — финальное закрытие (resolved / closed)
#   note    — заметки и ответы в журнале инцидента
# Тип incident_status_change раскрывается в status или close по new_status,
# поэтому он не указан в таблице явно.
_NOTIFY_EVENT_CLASS = {
    "incident_assigned": "open",
    "incident_author_update": "open",
    "incident_resolved": "close",
    "incident_note": "note",
    "incident_note_reply": "note",
}
_NOTIFY_CLOSING_STATUSES = {"resolved", "closed"}
_NOTIFY_MODE_ALLOWED = {
    "all": {"open", "status", "close", "note"},
    "status": {"status", "close"},
    "open_close": {"open", "close"},
    "close_only": {"close"},
}


def _notify_user(user_id, incident_id, notification_type, title, message, new_status=None):
    """Создать запись уведомления (для колокольчика в шапке). Безопасно:
    если user_id пустой — ничего не делает.
    Фильтрует по User.notify_incident_mode: all / status / open_close / close_only.
    new_status передаётся для incident_status_change — если статус закрывающий
    (resolved/closed), событие считается «закрытием»."""
    if not user_id:
        return
    from app.models_legacy import IncidentNotification, User as _UserModel
    try:
        u = _UserModel.query.get(user_id)
    except Exception:
        u = None
    mode = (getattr(u, "notify_incident_mode", None) or "all").strip() or "all"
    if mode != "all":
        if notification_type == "incident_status_change":
            ev = "close" if (new_status in _NOTIFY_CLOSING_STATUSES) else "status"
        else:
            ev = _NOTIFY_EVENT_CLASS.get(notification_type)
        allowed = _NOTIFY_MODE_ALLOWED.get(mode, _NOTIFY_MODE_ALLOWED["all"])
        if ev is None or ev not in allowed:
            return
    db.session.add(IncidentNotification(
        incident_id=incident_id,
        user_id=user_id,
        notification_type=notification_type,
        title=title[:255],
        message=message,
    ))
    try:
        from app.services.mobile_push import send_mobile_push_to_user
        from app.services.notification_channels import allows_mobile_app

        if allows_mobile_app(u):
            send_mobile_push_to_user(
                user_id,
                title,
                message,
                data={
                    "kind": "incident",
                    "incident_id": incident_id,
                    "notification_type": notification_type,
                },
            )
    except Exception:
        current_app.logger.exception(
            "Failed to send mobile push for incident_id=%s user_id=%s",
            incident_id,
            user_id,
        )


def _log_status_change(inc, old_status, new_status, comment=None):
    from app.models_legacy import IncidentStatusHistory
    if old_status == new_status:
        return
    db.session.add(IncidentStatusHistory(
        incident_id=inc.id,
        from_status=old_status,
        to_status=new_status,
        changed_by_id=getattr(current_user, "id", None),
        comment=comment,
    ))


def _log_assignment_change(inc, old_assignee_id, new_assignee_id, note=None):
    """Закрыть текущую запись и открыть новую при смене assignee."""
    from app.models_legacy import IncidentAssignment
    if old_assignee_id == new_assignee_id:
        return
    now = datetime.utcnow()
    current = (
        IncidentAssignment.query
        .filter_by(incident_id=inc.id, ended_at=None)
        .order_by(IncidentAssignment.assigned_at.desc())
        .first()
    )
    if current:
        current.ended_at = now
    db.session.add(IncidentAssignment(
        incident_id=inc.id,
        from_user_id=old_assignee_id,
        to_user_id=new_assignee_id,
        assigned_by_id=getattr(current_user, "id", None),
        note=note,
        assigned_at=now,
    ))


def _auto_create_task_for_incident(inc, assignee_user):
    """Автосоздание задачи для исполнителя инцидента.
    Создаётся только для PSYCHOLOGIST/SOCIAL_PEDAGOG/METHODIST.
    Если задача на этот инцидент уже существует и не закрыта — меняем responsible.
    Если закрыта — создаём новую."""
    if not assignee_user:
        return None
    role = getattr(assignee_user, "role", None) or ""
    role_codes = set()
    try:
        role_codes = {r.code for r in (assignee_user.roles or [])}
    except Exception:
        pass
    eligible = {"PSYCHOLOGIST", "SOCIAL_PEDAGOG", "METHODIST"}
    if role not in eligible and not (role_codes & eligible):
        return None

    from app.models.tasks import Task
    from app.models_legacy import AcademicYear
    existing = (
        Task.query
        .filter(Task.incident_id == inc.id)
        .filter(Task.status.notin_([Task.STATUS_CLOSED, Task.STATUS_CANCELLED, Task.STATUS_DONE]))
        .order_by(Task.id.desc())
        .first()
    )
    if existing:
        if existing.responsible_user_id != assignee_user.id:
            existing.responsible_user_id = assignee_user.id
        return existing

    year = AcademicYear.query.filter_by(is_current=True).first()
    title = f"[Инцидент #{inc.id}] {inc.category or 'Инцидент'}"
    description = (inc.description or "").strip()
    if description:
        description = f"Инцидент от {inc.occurred_at.strftime('%d.%m.%Y %H:%M') if inc.occurred_at else '—'}\nКатегория: {inc.category}\n\n{description}"
    task = Task(
        title=title[:255],
        description=description or None,
        priority="обычный",
        status=Task.STATUS_NEW,
        creator_user_id=getattr(current_user, "id", None) or assignee_user.id,
        responsible_user_id=assignee_user.id,
        incident_id=inc.id,
        academic_year_id=year.id if year else None,
        deadline_at=datetime.utcnow() + timedelta(days=3),
    )
    db.session.add(task)
    return task


def _incident_assignee_ids(inc):
    """Полный набор id исполнителей инцидента (включая legacy assignee_id)."""
    ids = {a.id for a in (inc.assignees or [])}
    if inc.assignee_id:
        ids.add(inc.assignee_id)
    return ids


def _uid_is_assignee(inc, uid):
    if not uid:
        return False
    return uid in _incident_assignee_ids(inc)


def _apply_assignees_change(inc, new_ids, note_text=None):
    """Синхронизирует список исполнителей инцидента с new_ids (set/iterable).
    Журнал, заметка [Назначение], авто-Task для каждого добавленного PSY/SOC/METHODIST,
    автопереход new↔assigned, уведомления всем добавленным + автору (один раз).

    Уведомления НЕ отправляются исполняющему действие пользователю.
    Возвращает (added_ids, removed_ids).
    """
    from app.models_legacy import IncidentAssignee, User as _User, IncidentNote
    actor_id = getattr(current_user, "id", None)
    new_ids = {int(x) for x in (new_ids or []) if x}

    current_ids = {a.id for a in (inc.assignees or [])}
    if not current_ids and inc.assignee_id:
        current_ids = {inc.assignee_id}

    added = new_ids - current_ids
    removed = current_ids - new_ids
    if not added and not removed:
        return added, removed

    if removed:
        IncidentAssignee.query.filter(
            IncidentAssignee.incident_id == inc.id,
            IncidentAssignee.user_id.in_(removed),
        ).delete(synchronize_session=False)
    for uid in added:
        db.session.add(IncidentAssignee(
            incident_id=inc.id, user_id=uid, added_by_id=actor_id,
        ))

    # primary assignee_id — последнее добавление; если набор пуст — сбрасываем.
    if new_ids:
        if added:
            inc.assignee_id = next(iter(added))
        elif inc.assignee_id not in new_ids:
            inc.assignee_id = next(iter(new_ids))
    else:
        inc.assignee_id = None

    for uid in added:
        _log_assignment_change(inc, None, uid, note=note_text)
    for uid in removed:
        _log_assignment_change(inc, uid, None, note=None)

    if note_text and added:
        db.session.add(IncidentNote(
            incident_id=inc.id, author_id=actor_id,
            text=f"[Назначение] {note_text}",
        ))

    old_status = inc.status or "new"
    if new_ids and old_status == "new":
        inc.status = "assigned"
    elif not new_ids and inc.status == "assigned":
        inc.status = "new"
    if inc.status != old_status:
        _log_status_change(inc, old_status, inc.status)

    for uid in added:
        u = _User.query.get(uid)
        if u:
            _auto_create_task_for_incident(inc, u)
        if uid != actor_id:
            preview = note_text if note_text else (inc.description or "")
            _notify_user(
                uid, inc.id, "incident_assigned",
                f"Вам назначен инцидент #{inc.id}",
                f"{inc.category}: {preview[:120]}",
            )
    if added and inc.author_id and inc.author_id != actor_id and inc.author_id not in added:
        _notify_user(
            inc.author_id, inc.id, "incident_author_update",
            f"По вашему инциденту #{inc.id} назначен исполнитель",
            f"{inc.category}",
        )
    return added, removed


def _apply_assignee_change(inc, new_assignee_id, note_text=None):
    """Совместимость: replace всего набора одним исполнителем (или None)."""
    new_ids = {new_assignee_id} if new_assignee_id else set()
    added, removed = _apply_assignees_change(inc, new_ids, note_text=note_text)
    return bool(added or removed)


def _build_incident_rows(incidents, include_author=False):
    """Batch-load children for a list of incidents (eliminates N+1)."""
    if not incidents:
        return []
    inc_ids = [inc.id for inc in incidents]
    # s81: batch-load assignee/author в session cache — гасит N+1 на kanban
    # (до 1200 запросов на 600 карточек), table/list тоже выигрывает.
    db.session.query(Incident).options(
        joinedload(Incident.assignee),
        joinedload(Incident.author),
        selectinload(Incident.assignees),
    ).filter(Incident.id.in_(inc_ids)).all()
    # one query: all IncidentChild + Child for all incidents at once
    links = (
        db.session.query(IncidentChild, Child)
        .join(Child, Child.id == IncidentChild.child_id)
        .options(
            joinedload(Child.enrollments)
                .joinedload(ChildEnrollment.school_class)
                .joinedload(SchoolClass.building),
        )
        .filter(IncidentChild.incident_id.in_(inc_ids))
        .all()
    )
    # group by incident_id
    kids_map = {}
    for lk, ch in links:
        cur_cls = ch.current_class
        bld = cur_cls.building if cur_cls else None
        kids_map.setdefault(lk.incident_id, []).append({
            "id": ch.id,
            "fio": ch.fio,
            "class": ch.current_class_name or "—",
            "class_id": cur_cls.id if cur_cls else None,
            "building": (bld.short_name or bld.name) if bld else None,
        })
    rows = []
    for inc in incidents:
        row = {"inc": inc, "children": kids_map.get(inc.id, [])}
        if include_author:
            row["author_label"] = _get_author_label(inc.author)
        # Множественные исполнители: список ids + короткий лейбл для шаблонов.
        row["assignee_ids"] = sorted(_incident_assignee_ids(inc))
        row["assignees_label"] = _assignees_short_label(inc) or "—"
        rows.append(row)
    return rows


_GROUP_BY_KEYS = ("category", "class", "building", "assignee", "status")


def _group_incident_rows(rows, group_by, status_labels=None):
    """Сгруппировать построенные _build_incident_rows строки по выбранному свойству.

    Группировка идёт ВНУТРИ страницы (в кашне на 100 строк) — это Notion-style,
    без cross-page rollup. Заголовки рендерятся в порядке первого появления.
    Возвращает [{key, label, rows}].
    """
    if group_by not in _GROUP_BY_KEYS:
        return None
    groups_map = {}
    order = []

    def _push(key, label, row):
        if key not in groups_map:
            groups_map[key] = {"key": key, "label": label, "rows": []}
            order.append(key)
        groups_map[key]["rows"].append(row)

    for r in rows:
        inc = r["inc"]
        kids = r.get("children") or []
        if group_by == "category":
            label = _category_label(inc.category)
            key = label
        elif group_by == "class":
            first = next((k for k in kids if k.get("class")), None)
            label = first["class"] if first else "Без класса"
            key = str(first["class_id"]) if first and first.get("class_id") else "_none"
        elif group_by == "building":
            first = next((k for k in kids if k.get("building")), None)
            label = first["building"] if first else "Без здания"
            key = label
        elif group_by == "assignee":
            a = inc.assignee
            if a:
                label = (a.last_name or "") + " " + (a.first_name or "")
                label = label.strip() or (a.username or f"user#{a.id}")
                key = f"u{a.id}"
            else:
                label = "Без исполнителя"
                key = "_none"
        else:  # status
            s = inc.status or "new"
            label = (status_labels or {}).get(s, s)
            key = s
        _push(key, label, r)
    return [groups_map[k] for k in order]


def _can_view_incident(incident):
    """Открыть карточку инцидента (GET).
    Разрешено: управляющим (ADMIN/DEPUTY_DIRECTOR), аналитическим ролям
    (METHODIST/PSYCHOLOGIST/SOCIAL_PEDAGOG — у них уже есть реестр/дашборд),
    автору инцидента и назначенному исполнителю."""
    if _can_change_status():
        return True
    uid = getattr(current_user, "id", None)
    if not uid:
        return False
    if has_role("METHODIST") or has_role("PSYCHOLOGIST") or has_role("SOCIAL_PEDAGOG"):
        return True
    if incident.author_id == uid:
        return True
    if _uid_is_assignee(incident, uid):
        return True
    return False


def _can_edit_incident(incident):
    """Редактировать поля инцидента (дата/категория/описание/участники) и удалять.
    Разрешено: управляющим (ADMIN/DEPUTY_DIRECTOR), социальному педагогу
    (ведёт инциденты наравне с управлением), автору и исполнителю
    (для автора/исполнителя нужно право incident_add — у METHODIST/TEACHER
    оно есть только на свои)."""
    if _can_change_status():
        return True
    if has_role("SOCIAL_PEDAGOG"):
        return True
    uid = getattr(current_user, "id", None)
    if not uid:
        return False
    if not has_permission("incident_add"):
        return False
    if incident.author_id == uid:
        return True
    if _uid_is_assignee(incident, uid):
        return True
    return False


# Совместимость со старым кодом (ранее единая проверка).
def _can_manage_incident(incident):
    return _can_view_incident(incident)


@children_bp.route("/incidents/<int:incident_id>/edit", methods=["GET", "POST"])
@login_required
def incident_edit(incident_id):
    inc = Incident.query.get_or_404(incident_id)
    if not _can_view_incident(inc):
        abort(403)

    if request.method == "POST":
        if not _can_edit_incident(inc):
            abort(403)
        occurred_date = (request.form.get("occurred_date") or "").strip()
        occurred_hour = (request.form.get("occurred_hour") or "").strip()
        occurred_minute = (request.form.get("occurred_minute") or "").strip()
        occurred_time = f"{occurred_hour}:{occurred_minute}"
        category = (request.form.get("category") or "").strip()
        description = (request.form.get("description") or "").strip() or None

        raw_ids = request.form.getlist("child_ids")
        child_ids = []
        for x in raw_ids:
            if str(x).isdigit():
                child_ids.append(int(x))
        child_ids = list(dict.fromkeys(child_ids))

        if not occurred_date or not occurred_time:
            flash("Укажите дату и время", "danger")
            return redirect(url_for("children.incident_edit", incident_id=inc.id))

        if category not in INCIDENT_CATEGORIES:
            flash("Выберите категорию инцидента", "danger")
            return redirect(url_for("children.incident_edit", incident_id=inc.id))

        if not description:
            flash("Заполните описание инцидента", "danger")
            return redirect(url_for("children.incident_edit", incident_id=inc.id))

        if not child_ids:
            flash("Добавьте хотя бы одного ребёнка", "danger")
            return redirect(url_for("children.incident_edit", incident_id=inc.id))

        try:
            occurred_at = datetime.strptime(f"{occurred_date} {occurred_time}", "%Y-%m-%d %H:%M")
        except Exception:
            flash("Неверный формат даты/времени", "danger")
            return redirect(url_for("children.incident_edit", incident_id=inc.id))

        # 48-часовое окно подачи задним числом — действует и при редактировании
        # автором/исполнителем; ADMIN/DEPUTY имеют право поправить дату вне окна.
        now = _now_msk_naive()
        if not (is_admin(current_user) or has_role("DEPUTY_DIRECTOR")):
            if occurred_at < now - timedelta(hours=48):
                flash("Дата инцидента не должна быть раньше, чем за 48 часов до момента подачи заявки", "danger")
                return redirect(url_for("children.incident_edit", incident_id=inc.id))
            if occurred_at > now + timedelta(minutes=5):
                flash("Дата инцидента не может быть в будущем", "danger")
                return redirect(url_for("children.incident_edit", incident_id=inc.id))

        inc.occurred_at = occurred_at
        inc.category = category
        inc.description = description

        old_status = inc.status or "new"
        old_assignee_id = inc.assignee_id

        new_status = (request.form.get("status") or "").strip()
        if new_status in Incident.STATUS_LABELS and _can_change_status():
            inc.status = new_status

        if _can_change_status():
            raw_ids = request.form.getlist("assignee_ids")
            new_ids = set()
            for v in raw_ids:
                v = (v or "").strip()
                if v.isdigit():
                    new_ids.add(int(v))
            # обратная совместимость со старым полем
            if not raw_ids:
                _legacy = request.form.get("assignee_id", type=int)
                if _legacy:
                    new_ids.add(_legacy)
            assignment_note = (request.form.get("assignment_note") or "").strip() or None
            _apply_assignees_change(inc, new_ids, note_text=assignment_note)

        if inc.status != old_status:
            _log_status_change(inc, old_status, inc.status)
            _actor_id = getattr(current_user, "id", None)
            _notified = set()
            if inc.author_id and inc.author_id != _actor_id:
                _notify_user(
                    inc.author_id, inc.id, "incident_status_change",
                    f"Статус инцидента #{inc.id}: {Incident.STATUS_LABELS.get(inc.status, inc.status)}",
                    f"{inc.category}",
                    new_status=inc.status,
                )
                _notified.add(inc.author_id)
            for _aid in _incident_assignee_ids(inc):
                if _aid in _notified or _aid == _actor_id:
                    continue
                _notify_user(
                    _aid, inc.id, "incident_status_change",
                    f"Статус инцидента #{inc.id}: {Incident.STATUS_LABELS.get(inc.status, inc.status)}",
                    f"{inc.category}",
                    new_status=inc.status,
                )
                _notified.add(_aid)

        IncidentChild.query.filter_by(incident_id=inc.id).delete()
        for cid in child_ids:
            ch = Child.query.get(cid)
            if ch:
                db.session.add(IncidentChild(incident_id=inc.id, child_id=ch.id))

        db.session.commit()
        flash("Инцидент обновлён", "success")
        next_url = (request.form.get("next") or "").strip()
        # s81: блокируем schema-relative //evil.com/x (open-redirect).
        if next_url.startswith("/") and not next_url.startswith("//"):
            return redirect(next_url)
        if child_ids:
            return redirect(url_for("children.child_card", child_id=child_ids[0]))
        return redirect(url_for("children.incidents_registry"))

    selected_children = [link.child for link in IncidentChild.query.filter_by(incident_id=inc.id).all() if link.child]
    grouped = {}
    for ch in selected_children:
        cl = ch.current_class
        key = getattr(cl, "id", None) or f"child-{ch.id}"
        if key not in grouped:
            grouped[key] = {
                "grade": getattr(cl, "grade", None) or "",
                "class_id": getattr(cl, "id", None) or "",
                "child_ids": [],
            }
        grouped[key]["child_ids"].append(ch.id)

    selected_blocks = list(grouped.values()) or [{"grade": "", "class_id": "", "child_ids": []}]

    assignees = []
    if _can_change_status():
        from app.models_legacy import User as _User, UserRole as _UserRole, Role as _Role
        assignees = (
            _User.query
            .join(_UserRole, _UserRole.user_id == _User.id)
            .join(_Role, _Role.id == _UserRole.role_id)
            .filter(_Role.code.in_([
                "ADMIN", "DEPUTY_DIRECTOR", "PSYCHOLOGIST", "SOCIAL_PEDAGOG",
                "METHODIST", "CLASS_TEACHER", "TEACHER",
            ]))
            .distinct()
            .order_by(_User.last_name, _User.first_name)
            .all()
        )

    is_author = inc.author_id == getattr(current_user, "id", None)
    is_assignee = _uid_is_assignee(inc, getattr(current_user, "id", None))
    can_add_note = _can_change_status() or has_role("SOCIAL_PEDAGOG") or is_author or is_assignee
    # Отвечать (reply) могут: ADMIN/DEPUTY, SOCIAL_PEDAGOG, автор инцидента.
    # Исполнитель-assignee не в списке по договорённости с пользователем.
    can_reply = _can_change_status() or has_role("SOCIAL_PEDAGOG") or is_author

    # Группируем заметки в 2-уровневые треды: root + ответы.
    # inc.notes отсортированы по created_at ASC (order_by в relationship).
    note_threads = []
    _by_id = {}
    for _n in (inc.notes or []):
        _by_id[_n.id] = {"note": _n, "replies": []}
    for _n in (inc.notes or []):
        if _n.parent_id and _n.parent_id in _by_id:
            # Ответ на ответ — прикрепляем к корню треда, чтобы не было
            # глубокой вложенности.
            _root_id = _n.parent_id
            while _by_id[_root_id]["note"].parent_id and _by_id[_root_id]["note"].parent_id in _by_id:
                _root_id = _by_id[_root_id]["note"].parent_id
            _by_id[_root_id]["replies"].append(_by_id[_n.id])
        else:
            note_threads.append(_by_id[_n.id])

    # Связанные задачи (автосозданные через _auto_create_task_for_incident + ручные).
    # backref `inc.tasks` отсортирован по created_at desc; падать на старте всё равно нельзя.
    try:
        related_tasks = list(inc.tasks)
    except Exception:
        related_tasks = []

    selected_assignee_ids = sorted(_incident_assignee_ids(inc))
    return render_template(
        "incident_edit.html",
        incident=inc,
        categories=INCIDENT_CATEGORIES,
        selected_blocks=selected_blocks,
        can_change_status=_can_change_status(),
        can_edit_incident=_can_edit_incident(inc),
        assignees=assignees,
        selected_assignee_ids=selected_assignee_ids,
        notes=inc.notes,
        note_threads=note_threads,
        can_add_note=can_add_note,
        can_reply=can_reply,
        related_tasks=related_tasks,
    )


@children_bp.route("/incidents/<int:incident_id>/delete", methods=["POST"])
@login_required
def incident_delete(incident_id):
    inc = Incident.query.get_or_404(incident_id)
    if not _can_edit_incident(inc):
        abort(403)
    # METHODIST: edit-свои разрешено (s76), delete — нет.
    if has_role("METHODIST") and not _can_change_status() and not has_role("SOCIAL_PEDAGOG"):
        abort(403)

    child_id = None
    first_link = IncidentChild.query.filter_by(incident_id=inc.id).first()
    if first_link:
        child_id = first_link.child_id

    IncidentChild.query.filter_by(incident_id=inc.id).delete()

    from app.models_legacy import IncidentNotification
    from app.models.tasks import Task
    IncidentNotification.query.filter_by(incident_id=inc.id).delete()
    Task.query.filter_by(incident_id=inc.id).update({"incident_id": None})

    db.session.delete(inc)
    db.session.commit()
    flash("Инцидент удалён", "success")

    next_url = (request.form.get("next") or request.referrer or "").strip()
    if next_url.startswith("/") and not next_url.startswith("//"):
        return redirect(next_url)
    if child_id:
        return redirect(url_for("children.child_card", child_id=child_id))
    return redirect(url_for("children.incidents_registry"))


@children_bp.route("/incidents/<int:incident_id>/set-status", methods=["POST"])
@login_required
def incident_set_status(incident_id):
    if not _can_change_status():
        return jsonify({"error": "forbidden"}), 403
    new_status = (request.form.get("status") or "").strip()
    if new_status not in Incident.STATUS_LABELS:
        return jsonify({"error": "invalid"}), 400
    # Row-lock на PG (на SQLite молча no-op) — два менеджера не перезатирают друг друга.
    inc = Incident.query.filter_by(id=incident_id).with_for_update().first()
    if inc is None:
        abort(404)
    old_status = inc.status or "new"
    if old_status != new_status:
        inc.status = new_status
        _log_status_change(inc, old_status, new_status)
        _actor_id = getattr(current_user, "id", None)
        _notified = set()
        if inc.author_id and inc.author_id != _actor_id:
            _notify_user(
                inc.author_id, inc.id, "incident_status_change",
                f"Статус инцидента #{inc.id}: {Incident.STATUS_LABELS.get(new_status, new_status)}",
                f"{inc.category}",
                new_status=new_status,
            )
            _notified.add(inc.author_id)
        for _aid in _incident_assignee_ids(inc):
            if _aid in _notified or _aid == _actor_id:
                continue
            _notify_user(
                _aid, inc.id, "incident_status_change",
                f"Статус инцидента #{inc.id}: {Incident.STATUS_LABELS.get(new_status, new_status)}",
                f"{inc.category}",
                new_status=new_status,
            )
            _notified.add(_aid)
    db.session.commit()
    return jsonify({"ok": True, "status": inc.status, "label": inc.status_label})


@children_bp.route("/incidents/<int:incident_id>/set-assignee", methods=["POST"])
@login_required
def incident_set_assignee(incident_id):
    """Назначить исполнителя прямо из реестра / /incidents/my, без захода в карточку.
    Повторяет логику формы incident_edit: автопереход статусов и автосоздание Task.
    Принимает опциональное поле `note` — пояснение для исполнителя; сохраняется
    в IncidentAssignment.note, попадает в журнал работы (IncidentNote) и
    добавляется в preview уведомления."""
    if not _can_change_status():
        return jsonify({"error": "forbidden"}), 403
    # Row-lock на PG (на SQLite no-op) — два менеджера не перезатирают друг друга.
    inc = Incident.query.filter_by(id=incident_id).with_for_update().first()
    if inc is None:
        abort(404)

    raw_ids = request.form.getlist("assignee_ids")
    new_ids = set()
    for v in raw_ids:
        v = (v or "").strip()
        if v.isdigit():
            new_ids.add(int(v))
    if not raw_ids:
        # Совместимость: одно поле assignee_id (пустое = снять).
        legacy = (request.form.get("assignee_id") or "").strip()
        if legacy.isdigit():
            new_ids.add(int(legacy))
    note_text = (request.form.get("note") or "").strip() or None

    added, removed = _apply_assignees_change(inc, new_ids, note_text=note_text)
    db.session.commit()
    return jsonify({
        "ok": True,
        "status": inc.status,
        "assignee_id": inc.assignee_id,
        "assignee_ids": sorted(_incident_assignee_ids(inc)),
        "assignee_label": _assignees_short_label(inc),
    })


def _user_short_label(u):
    if not u:
        return None
    last = getattr(u, "last_name", None) or ""
    first = getattr(u, "first_name", None) or ""
    if last and first:
        return f"{last} {first[:1]}."
    return last or first or (getattr(u, "username", "") or "—")


def _assignee_short_label(inc):
    return _user_short_label(inc.assignee)


def _assignees_short_label(inc, max_show=2):
    """«Иванова И., Петров П.» или «Иванова И. + 2» если больше max_show."""
    users = list(inc.assignees or [])
    if not users and inc.assignee:
        users = [inc.assignee]
    if not users:
        return None
    labels = [_user_short_label(u) for u in users if u]
    labels = [x for x in labels if x]
    if len(labels) <= max_show:
        return ", ".join(labels)
    return f"{', '.join(labels[:max_show])} +{len(labels) - max_show}"


@children_bp.route("/incidents/<int:incident_id>/mark-resolved", methods=["POST"])
@login_required
def incident_mark_resolved(incident_id):
    """Исполнитель помечает инцидент как «Отработан» — финальное закрытие
    по-прежнему за управляющим.
    Требует непустой comment. Файлы (files[]) опциональны — сохраняются как
    вложения к той же заметке через тот же пайп, что у add_note."""
    inc = Incident.query.get_or_404(incident_id)
    if not _can_mark_resolved(inc):
        return jsonify({"error": "forbidden"}), 403
    if inc.status in ("resolved", "closed"):
        return jsonify({"ok": True, "status": inc.status, "label": inc.status_label})
    comment = (request.form.get("comment") or "").strip()
    if not comment:
        return jsonify({"error": "comment_required"}), 400

    old_status = inc.status or "new"
    inc.status = "resolved"
    _log_status_change(inc, old_status, "resolved", comment=comment)

    # Заметку с префиксом [Отработано] создаём всегда — чтобы описание работы
    # было видно на странице инцидента и в журнале.
    from app.models_legacy import IncidentNote
    note = IncidentNote(
        incident_id=inc.id,
        author_id=getattr(current_user, "id", None),
        text=f"[Отработано] {comment}",
    )
    db.session.add(note)
    db.session.flush()  # нужен note.id для путей вложений

    # Опциональные вложения (как в add_note)
    files = request.files.getlist("files[]") or request.files.getlist("files")
    saved_attachments = []
    if files and any(f and getattr(f, "filename", "") for f in files):
        try:
            saved_attachments = _save_incident_note_attachments(note, files)
        except ValueError as _ve:
            db.session.rollback()
            return jsonify({"error": str(_ve)}), 400
        except Exception:
            db.session.rollback()
            return jsonify({"error": "upload_failed"}), 500

    # уведомления управляющим и автору
    if inc.author_id and inc.author_id != getattr(current_user, "id", None):
        _notify_user(
            inc.author_id, inc.id, "incident_resolved",
            f"Инцидент #{inc.id} отработан",
            f"{inc.category}: {comment[:120]}",
        )
    # уведомления всем ADMIN и DEPUTY_DIRECTOR
    try:
        from app.models_legacy import User as _User, UserRole as _UserRole, Role as _Role
        managers = (
            _User.query
            .join(_UserRole, _UserRole.user_id == _User.id)
            .join(_Role, _Role.id == _UserRole.role_id)
            .filter(_Role.code.in_(["ADMIN", "DEPUTY_DIRECTOR"]))
            .all()
        )
        me_id = getattr(current_user, "id", None)
        for mgr in managers:
            if mgr.id == me_id:
                continue
            _notify_user(
                mgr.id, inc.id, "incident_resolved",
                f"Исполнитель отметил инцидент #{inc.id} как «Отработан»",
                f"{inc.category}: {comment[:120]}",
            )
    except Exception:
        pass

    db.session.commit()
    return jsonify({
        "ok": True,
        "status": inc.status,
        "label": inc.status_label,
        "attachments": len(saved_attachments),
    })


@children_bp.route("/incidents/<int:incident_id>/add-note", methods=["POST"])
@login_required
def incident_add_note(incident_id):
    inc = Incident.query.get_or_404(incident_id)
    # Писать заметки могут: ADMIN/DEPUTY, SOCIAL_PEDAGOG, автор инцидента, исполнитель.
    # Для ответа (reply) круг тот же — соответствует договорённости.
    is_author = inc.author_id == getattr(current_user, "id", None)
    is_assignee = _uid_is_assignee(inc, getattr(current_user, "id", None))
    if not (_can_change_status() or has_role("SOCIAL_PEDAGOG") or is_author or is_assignee):
        return jsonify({"error": "forbidden"}), 403
    text = (request.form.get("text") or "").strip()
    incoming_files = request.files.getlist("files[]") or request.files.getlist("files")
    has_files = any(f and getattr(f, "filename", "") for f in incoming_files)
    if not text and not has_files:
        return jsonify({"error": "empty"}), 400
    if not text and has_files:
        text = "(вложение)"

    parent_id = request.form.get("parent_id", type=int)
    parent = None
    if parent_id:
        parent = IncidentNote.query.get(parent_id)
        # parent должен принадлежать тому же инциденту — иначе игнорируем
        if not parent or parent.incident_id != inc.id:
            parent = None
            parent_id = None

    note = IncidentNote(
        incident_id=inc.id,
        author_id=current_user.id,
        text=text,
        parent_id=(parent.id if parent else None),
    )
    db.session.add(note)
    db.session.flush()  # получить note.id для пути вложений

    # Сохранение вложений (multipart/form-data, поле files[])
    saved_attachments = []
    try:
        files = request.files.getlist("files[]") or request.files.getlist("files")
        if files:
            saved_attachments = _save_incident_note_attachments(note, files)
    except ValueError as _ve:
        db.session.rollback()
        return jsonify({"error": str(_ve)}), 400
    except Exception:
        db.session.rollback()
        return jsonify({"error": "upload_failed"}), 500

    me_id = getattr(current_user, "id", None)
    preview = text[:120]

    if parent:
        # Целевое уведомление автору исходного комментария: «вам ответили»
        if parent.author_id and parent.author_id != me_id:
            _notify_user(
                parent.author_id, inc.id, "incident_note_reply",
                f"Вам ответили в инциденте #{inc.id}",
                preview,
            )
        # Плюс — общее уведомление участникам цепочки (автор/исполнитель),
        # кроме того, кому только что отправили целевое.
        targets = set()
        if inc.author_id and inc.author_id not in (me_id, parent.author_id):
            targets.add(inc.author_id)
        for _aid in _incident_assignee_ids(inc):
            if _aid not in (me_id, parent.author_id):
                targets.add(_aid)
        for tgt in targets:
            _notify_user(
                tgt, inc.id, "incident_note",
                f"Новая заметка в инциденте #{inc.id}",
                preview,
            )
    else:
        targets = set()
        if inc.author_id and inc.author_id != me_id:
            targets.add(inc.author_id)
        for _aid in _incident_assignee_ids(inc):
            if _aid != me_id:
                targets.add(_aid)
        for tgt in targets:
            _notify_user(
                tgt, inc.id, "incident_note",
                f"Новая заметка в инциденте #{inc.id}",
                preview,
            )
    db.session.commit()
    author = note.author
    author_name = " ".join(p for p in [author.last_name or "", author.first_name or ""] if p)
    return jsonify({
        "ok": True,
        "id": note.id,
        "text": note.text,
        "author": author_name or author.username,
        "created_at": note.created_at.strftime("%d.%m.%Y %H:%M"),
        "parent_id": note.parent_id,
        "attachments": [_inc_note_att_to_dict(a) for a in saved_attachments],
    })


@children_bp.route("/incidents/notifications/read/<int:notif_id>", methods=["POST"])
@login_required
def incident_notification_read(notif_id):
    from app.models_legacy import IncidentNotification
    n = IncidentNotification.query.get_or_404(notif_id)
    if n.user_id != current_user.id:
        return jsonify({"error": "forbidden"}), 403
    if not n.is_read:
        n.is_read = True
        n.read_at = datetime.utcnow()
        db.session.commit()
    incident_id = n.incident_id
    next_url = request.form.get("next") or url_for("children.incident_edit", incident_id=incident_id)
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True, "incident_id": incident_id})
    return redirect(next_url)


@children_bp.route("/incidents/notifications/mark-all-read", methods=["POST"])
@login_required
def incident_notifications_mark_all():
    from app.models_legacy import IncidentNotification
    now = datetime.utcnow()
    (
        IncidentNotification.query
        .filter_by(user_id=current_user.id, is_read=False)
        .update({"is_read": True, "read_at": now})
    )
    db.session.commit()
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True})
    return redirect(request.referrer or url_for("children.incidents_my"))


@children_bp.route("/incidents/<int:incident_id>/timeline")
@login_required
def incident_timeline(incident_id):
    """Объединённый таймлайн: заметки + смены статуса + смены исполнителя."""
    from app.models_legacy import IncidentNote, IncidentStatusHistory, IncidentAssignment
    inc = Incident.query.get_or_404(incident_id)
    # Права на просмотр таймлайна: автор, исполнитель или тот, кто видит реестр
    uid = getattr(current_user, "id", None)
    if not (
        _can_change_status()
        or inc.author_id == uid
        or _uid_is_assignee(inc, uid)
        or has_permission("incident_registry_view")
    ):
        abort(403)

    events = []
    events.append({
        "kind": "created",
        "ts": inc.created_at or inc.occurred_at,
        "actor": inc.author,
        "text": f"Инцидент создан: {inc.category}",
    })
    for note in (inc.notes or []):
        events.append({
            "kind": "note",
            "ts": note.created_at,
            "actor": note.author,
            "text": note.text,
        })
    for s in (inc.status_history or []):
        from_lbl = Incident.STATUS_LABELS.get(s.from_status, s.from_status or "—")
        to_lbl = Incident.STATUS_LABELS.get(s.to_status, s.to_status)
        events.append({
            "kind": "status",
            "ts": s.changed_at,
            "actor": s.changed_by,
            "text": f"Статус: {from_lbl} → {to_lbl}" + (f". {s.comment}" if s.comment else ""),
        })
    for a in (inc.assignments or []):
        from_name = (a.from_user.fio if a.from_user else "—")
        to_name = (a.to_user.fio if a.to_user else "—")
        events.append({
            "kind": "assignment",
            "ts": a.assigned_at,
            "actor": a.assigned_by,
            "text": f"Ведение: {from_name} → {to_name}" + (f". {a.note}" if a.note else ""),
        })
    events.sort(key=lambda e: e["ts"] or datetime.utcnow())

    return render_template(
        "incident_timeline.html",
        incident=inc,
        events=events,
    )


@children_bp.route("/incidents/registry")
@require_roles("ADMIN", "METHODIST", "PSYCHOLOGIST", "SOCIAL_PEDAGOG")
def incidents_registry():
    q_text = (request.args.get("q") or "").strip()
    grade = request.args.get("grade", type=int)
    class_id = request.args.get("class_id", type=int)
    category = (request.args.get("category") or "").strip()
    status_filter = (request.args.get("status") or "").strip()

    year = _get_current_year()
    year_id = year.id if year else None

    iq = (
        db.session.query(Incident)
        .join(IncidentChild, IncidentChild.incident_id == Incident.id)
        .join(Child, Child.id == IncidentChild.child_id)
        .options(
            joinedload(Incident.author),
            joinedload(Incident.assignee),
        )
    )

    if category:
        iq = iq.filter(Incident.category == category)

    if status_filter == "open":
        iq = iq.filter(Incident.status.in_(["new", "assigned", "in_progress"]))
    elif status_filter in Incident.STATUS_LABELS:
        iq = iq.filter(Incident.status == status_filter)

    if q_text:
        ql = q_text.lower()
        iq = iq.filter(
            func.lower(func.coalesce(Child.last_name, "")).like(f"%{ql}%") |
            func.lower(func.coalesce(Child.first_name, "")).like(f"%{ql}%") |
            func.lower(func.coalesce(Child.middle_name, "")).like(f"%{ql}%")
        )

    if year_id and (grade or class_id):
        iq = iq.join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        iq = iq.filter(
            ChildEnrollment.academic_year_id == year_id,
            ChildEnrollment.ended_at.is_(None)
        )
        if class_id:
            iq = iq.filter(ChildEnrollment.school_class_id == class_id)
        elif grade:
            iq = iq.join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            iq = iq.filter(SchoolClass.grade == grade)

    iq = iq.distinct(Incident.id) if False else iq  # joinedload не любит distinct, считаем через подзапрос ниже
    iq = iq.order_by(Incident.occurred_at.desc(), Incident.id.desc())

    # Серверная пагинация: 100 строк на страницу.
    _PER_PAGE_REG = 100
    try:
        page = max(1, int(request.args.get("page", 1)))
    except (TypeError, ValueError):
        page = 1
    total_count = iq.order_by(None).with_entities(func.count(func.distinct(Incident.id))).scalar() or 0
    total_pages = max(1, (total_count + _PER_PAGE_REG - 1) // _PER_PAGE_REG)
    if page > total_pages:
        page = total_pages
    # Подзапрос-окно: id-инциденты текущей страницы (без дублей из IncidentChild).
    page_ids_q = (
        iq.with_entities(Incident.id, Incident.occurred_at)
        .group_by(Incident.id, Incident.occurred_at)
        .order_by(Incident.occurred_at.desc(), Incident.id.desc())
        .limit(_PER_PAGE_REG)
        .offset((page - 1) * _PER_PAGE_REG)
    )
    page_ids = [row[0] for row in page_ids_q.all()]
    if page_ids:
        incidents = (
            db.session.query(Incident)
            .options(joinedload(Incident.author), joinedload(Incident.assignee))
            .filter(Incident.id.in_(page_ids))
            .order_by(Incident.occurred_at.desc(), Incident.id.desc())
            .all()
        )
    else:
        incidents = []

    rows = _build_incident_rows(incidents, include_author=True)

    f_group_by = (request.args.get("group_by") or "").strip()
    if f_group_by not in _GROUP_BY_KEYS:
        f_group_by = ""
    groups = None
    if f_group_by:
        groups = _group_incident_rows(rows, f_group_by, Incident.STATUS_LABELS)

    classes = (
        SchoolClass.query
        .filter(SchoolClass.academic_year_id == year.id)
        .order_by(
            SchoolClass.grade.asc().nullslast(),
            SchoolClass.letter.asc().nullslast(),
            SchoolClass.name.asc()
        )
        .all()
    )

    return render_template(
        "incidents_registry.html",
        title="Реестр инцидентов",
        rows=rows,
        groups=groups,
        f_group_by=f_group_by,
        q=q_text,
        grade=grade,
        class_id=class_id,
        category=category,
        status_filter=status_filter,
        categories=INCIDENT_CATEGORIES,
        classes=classes,
        can_change_status=_can_change_status(),
        is_admin=is_admin(current_user),
        is_social_pedagog=has_role("SOCIAL_PEDAGOG"),
        is_methodist=has_role("METHODIST"),
        page=page,
        total_pages=total_pages,
        total_count=total_count,
        per_page=_PER_PAGE_REG,
        export_url=url_for("children.incidents_registry_export", grade=grade, class_id=class_id, category=category, q=q_text, status=status_filter or None)
    )


@children_bp.route("/incidents/registry/export")
@require_roles("ADMIN", "METHODIST", "PSYCHOLOGIST", "SOCIAL_PEDAGOG")
def incidents_registry_export():
    q_text = (request.args.get("q") or "").strip()
    grade = request.args.get("grade", type=int)
    class_id = request.args.get("class_id", type=int)
    category = (request.args.get("category") or "").strip()
    status_filter = (request.args.get("status") or "").strip()
    # ?hide_cols=col-desc,col-author — CSV-список ключей колонок, которые НЕ выгружать.
    # JS в incidents_registry.html шлёт ключи с префиксом col-, поэтому здесь
    # принимаем оба варианта.
    hide_cols = set()
    for c in (request.args.get("hide_cols") or "").split(","):
        c = c.strip()
        if not c:
            continue
        hide_cols.add(c[4:] if c.startswith("col-") else c)

    year = _get_current_year()
    year_id = year.id if year else None

    iq = (
        db.session.query(Incident)
        .join(IncidentChild, IncidentChild.incident_id == Incident.id)
        .join(Child, Child.id == IncidentChild.child_id)
    )

    if category:
        iq = iq.filter(Incident.category == category)

    if status_filter == "open":
        iq = iq.filter(Incident.status.in_(["new", "assigned", "in_progress"]))
    elif status_filter in Incident.STATUS_LABELS:
        iq = iq.filter(Incident.status == status_filter)

    if q_text:
        ql = q_text.lower()
        iq = iq.filter(
            func.lower(func.coalesce(Child.last_name, "")).like(f"%{ql}%") |
            func.lower(func.coalesce(Child.first_name, "")).like(f"%{ql}%") |
            func.lower(func.coalesce(Child.middle_name, "")).like(f"%{ql}%")
        )

    if year_id and (grade or class_id):
        iq = iq.join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        iq = iq.filter(
            ChildEnrollment.academic_year_id == year_id,
            ChildEnrollment.ended_at.is_(None)
        )
        if class_id:
            iq = iq.filter(ChildEnrollment.school_class_id == class_id)
        elif grade:
            iq = iq.join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            iq = iq.filter(SchoolClass.grade == grade)

    incidents = iq.order_by(Incident.occurred_at.desc(), Incident.id.desc()).all()

    rows = _build_incident_rows(incidents, include_author=True)

    # Колонки экспорта. Ключи совпадают с col-picker в incidents_registry.html
    # (col-cat, col-students, col-class, col-desc, col-status, col-author).
    # ?hide_cols=desc,author — пропускает соответствующие столбцы.
    col_defs = [
        ("num",       "№",            lambda inc, r, idx: idx),
        ("date",      "Дата/время",   lambda inc, r, idx: inc.occurred_at.strftime("%d.%m.%Y %H:%M") if inc.occurred_at else ""),
        ("cat",       "Категория",    lambda inc, r, idx: inc.category or ""),
        ("students",  "Обучающиеся",  lambda inc, r, idx: "; ".join(k["fio"] for k in r["children"])),
        ("class",     "Классы",       lambda inc, r, idx: "; ".join(dict.fromkeys(k["class"] for k in r["children"]))),
        ("desc",      "Описание",     lambda inc, r, idx: inc.description or ""),
        ("status",    "Статус",       lambda inc, r, idx: Incident.STATUS_LABELS.get(inc.status or "new", inc.status or "")),
        ("author",    "Автор",        lambda inc, r, idx: r.get("author_label", "")),
    ]
    visible = [c for c in col_defs if c[0] not in hide_cols]

    wb = Workbook()
    ws = wb.active
    ws.title = "Инциденты"
    ws.append([c[1] for c in visible])

    for idx, r in enumerate(rows, start=1):
        inc = r["inc"]
        ws.append([fn(inc, r, idx) for _, _, fn in visible])

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="Реестр_инцидентов.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@children_bp.route("/incidents/my")
@login_required
def incidents_my():
    """
    Страница «Инциденты / Мои заявки». Два режима:
      • admin-view (ADMIN/DEPUTY_DIRECTOR): 3 вкладки по всем инцидентам
        — Входящие (new), В работе (assigned+in_progress), Завершённые (resolved+closed).
        На каждой вкладке — фильтры: категория, класс, поиск и сортировка.
      • user-view: «Мои заявки» — только инциденты, где текущий пользователь автор,
        с подгруженной последней заметкой исполнителя.
    """
    if not has_permission("incident_add"):
        abort(403)
    uid = current_user.id
    is_admin_view = _can_change_status()
    # Социальный педагог получает admin-view, но ограниченный его назначениями
    # и без picker-а исполнителя / смены статуса. Флаг отличает его от ADMIN/DEPUTY
    # там, где это важно (фильтр по assignee, tab-labels, доступ к колонке Кубок).
    is_social_view = (not is_admin_view) and has_role("SOCIAL_PEDAGOG")
    # METHODIST — read-only просмотр всех инцидентов по школе. Без picker-а исполнителя,
    # без смены статуса, без edit/delete в строках. 3 общие вкладки, без «Назначенные мне».
    is_methodist_view = (
        (not is_admin_view) and (not is_social_view) and has_role("METHODIST")
    )

    # Общие параметры фильтрации (действуют на всех вкладках админ-вида)
    default_tab = "incoming" if (is_admin_view or is_social_view) else "mine"
    active_tab = (request.args.get("tab") or default_tab).strip()
    f_category = (request.args.get("category") or "").strip()
    f_class_id = request.args.get("class_id", type=int)
    f_q = (request.args.get("q") or "").strip()
    f_sort = (request.args.get("sort") or "date_desc").strip()
    f_group_by = (request.args.get("group_by") or "").strip()
    if f_group_by not in _GROUP_BY_KEYS:
        f_group_by = ""

    year = _get_current_year()

    # ── Вариант для обычного пользователя ──
    # Две вкладки: «Мои заявки» (где я автор) + «Назначены мне» (где я исполнитель).
    # SOCIAL_PEDAGOG не попадает сюда — он использует admin-view, но с фильтром по assignee.
    if not is_admin_view and not is_social_view and not is_methodist_view:
        user_tab = active_tab if active_tab in ("mine", "assigned") else "mine"

        authored = (
            db.session.query(Incident)
            .filter(Incident.author_id == uid)
            .order_by(Incident.occurred_at.desc(), Incident.id.desc())
            .all()
        )
        assigned = (
            db.session.query(Incident)
            .filter(Incident.assignees.any(id=uid))
            .order_by(Incident.occurred_at.desc(), Incident.id.desc())
            .all()
        )
        authored_rows = _build_incident_rows(authored, include_author=True)
        assigned_rows = _build_incident_rows(assigned, include_author=True)

        from app.models_legacy import IncidentNote
        from sqlalchemy.orm import joinedload
        authored_ids = [r["inc"].id for r in authored_rows]
        last_note_map = {}
        if authored_ids:
            notes = (
                IncidentNote.query
                .options(joinedload(IncidentNote.author))
                .filter(IncidentNote.incident_id.in_(authored_ids))
                .order_by(IncidentNote.incident_id, IncidentNote.created_at.desc())
                .all()
            )
            for n in notes:
                if n.incident_id not in last_note_map:
                    last_note_map[n.incident_id] = n
        for r in authored_rows:
            r["last_note"] = last_note_map.get(r["inc"].id)

        user_counters = {
            "mine":     len(authored_rows),
            "assigned": len(assigned_rows),
        }

        return render_template(
            "incidents_my.html",
            is_admin_view=False,
            user_tab=user_tab,
            user_counters=user_counters,
            authored_rows=authored_rows,
            assigned_rows=assigned_rows,
            status_labels=Incident.STATUS_LABELS,
            current_user_id=uid,
        )

    # ── Вариант для ADMIN / DEPUTY_DIRECTOR и SOCIAL_PEDAGOG ──
    # ADMIN/DEPUTY: 3 вкладки по всем инцидентам.
    # SOCIAL_PEDAGOG: те же 3 вкладки (все инциденты по школе, read-only — без picker/смены статуса),
    # плюс 4-я вкладка «Назначенные мне» (только мои) — там разрешены edit/delete и «Я отработал».
    STATUS_BUCKETS = {
        "incoming":  ["new"],
        "in_work":   ["assigned", "in_progress"],
        "completed": ["resolved", "closed"],
    }

    view_mode_raw = (request.args.get("view") or "table").strip()
    is_kanban = (view_mode_raw == "kanban")
    is_calendar = (view_mode_raw == "calendar")

    # s78: parse ?month=YYYY-MM для календаря (default = current month по МСК).
    cal_year_num = cal_month_num = None
    if is_calendar:
        from datetime import date as _cal_date, datetime as _cal_dt, timezone as _cal_tz, timedelta as _cal_td
        try:
            from zoneinfo import ZoneInfo as _ZI
            _MSK_TZ = _ZI("Europe/Moscow")
        except Exception:
            _MSK_TZ = _cal_tz(_cal_td(hours=3))
        _today_msk = _cal_dt.now(_MSK_TZ).date()
        _month_raw = (request.args.get("month") or "").strip()
        try:
            _y, _m = _month_raw.split("-")
            cal_year_num, cal_month_num = int(_y), int(_m)
            _cal_date(cal_year_num, cal_month_num, 1)
        except Exception:
            cal_year_num, cal_month_num = _today_msk.year, _today_msk.month

    # Для social-view добавляем виртуальный bucket mine_all — все статусы кроме closed,
    # отфильтрованные по assignee_id=uid.
    if active_tab == "mine" and is_social_view and not is_kanban:
        iq = db.session.query(Incident).filter(
            Incident.assignees.any(id=uid),
            Incident.status != "closed",
        )
        statuses = None  # не используется ниже
    elif is_kanban:
        # Kanban игнорирует active_tab — показывает все 3 колонки сразу.
        # Для social-view с tab=mine — фильтр по assignee=я, иначе по всей школе.
        iq = db.session.query(Incident)
        if active_tab == "mine" and is_social_view:
            iq = iq.filter(Incident.assignees.any(id=uid))
        statuses = None
    elif view_mode_raw in ("table", "list", "calendar"):
        # s72: в табличном/списочном режиме показываем все статусы.
        # s78: calendar — отдельный месячный рендер, фильтр по occurred_at в пределах месяца.
        # Фильтрация по статусу — через popup-фильтр (?status=...). status-subtabs
        # остаются только в kanban (где статус — это сами колонки).
        statuses = None
        iq = db.session.query(Incident)
        f_status = (request.args.get("status") or "").strip()
        if f_status:
            iq = iq.filter(Incident.status == f_status)
        if is_calendar:
            from datetime import date as _cal_date
            _first = _cal_date(cal_year_num, cal_month_num, 1)
            if cal_month_num == 12:
                _next = _cal_date(cal_year_num + 1, 1, 1)
            else:
                _next = _cal_date(cal_year_num, cal_month_num + 1, 1)
            iq = iq.filter(Incident.occurred_at >= _first, Incident.occurred_at < _next)
    else:
        statuses = STATUS_BUCKETS.get(active_tab, STATUS_BUCKETS["incoming"])
        iq = db.session.query(Incident).filter(Incident.status.in_(statuses))

    if f_category:
        iq = iq.filter(Incident.category == f_category)

    if f_q:
        ql = f_q.lower()
        iq = (
            iq.join(IncidentChild, IncidentChild.incident_id == Incident.id, isouter=True)
              .join(Child, Child.id == IncidentChild.child_id, isouter=True)
              .filter(
                  func.lower(func.coalesce(Child.last_name, "")).like(f"%{ql}%") |
                  func.lower(func.coalesce(Child.first_name, "")).like(f"%{ql}%") |
                  func.lower(func.coalesce(Child.middle_name, "")).like(f"%{ql}%") |
                  func.lower(func.coalesce(Incident.description, "")).like(f"%{ql}%")
              )
              .distinct()
        )

    if f_class_id and year:
        iq = (
            iq.join(IncidentChild, IncidentChild.incident_id == Incident.id, isouter=True)
              .join(ChildEnrollment, ChildEnrollment.child_id == IncidentChild.child_id, isouter=True)
              .filter(
                  ChildEnrollment.academic_year_id == year.id,
                  ChildEnrollment.ended_at.is_(None),
                  ChildEnrollment.school_class_id == f_class_id,
              )
              .distinct()
        )

    # Сортировка
    if f_sort == "date_asc":
        iq = iq.order_by(Incident.occurred_at.asc().nullslast(), Incident.id.asc())
    elif f_sort == "category":
        iq = iq.order_by(Incident.category.asc().nullslast(), Incident.occurred_at.desc())
    elif f_sort == "status":
        # Сортировка по STATUS_ORDER (логика жизненного цикла), не по алфавиту кода.
        from sqlalchemy import case as _sa_case
        _status_rank = _sa_case(
            {s: i for i, s in enumerate(Incident.STATUS_ORDER)},
            value=Incident.status,
            else_=len(Incident.STATUS_ORDER),
        )
        iq = iq.order_by(_status_rank.asc(), Incident.occurred_at.desc())
    else:
        iq = iq.order_by(Incident.occurred_at.desc().nullslast(), Incident.id.desc())

    # Пагинация: 100 строк на страницу, ?page=N. Для админ-вида (admin/social/methodist).
    _PER_PAGE = 100
    try:
        page = max(1, int(request.args.get("page", 1)))
    except (TypeError, ValueError):
        page = 1

    kanban_groups = None
    task_counts_map = {}
    cal_ctx = None
    if is_calendar:
        # s78: загрузить все инциденты выбранного месяца (без пагинации), сгруппировать по дате.
        import calendar as _cal_mod
        _CAL_LIMIT = 2000
        incidents = iq.order_by(Incident.occurred_at.asc()).limit(_CAL_LIMIT).all()
        rows = _build_incident_rows(incidents, include_author=True)
        total_count = len(rows)
        total_pages = 1
        page = 1
        rows_limit_reached = (total_count >= _CAL_LIMIT)
        _LIMIT = _CAL_LIMIT
        _by_date = {}
        for _r in rows:
            _oa = _r["inc"].occurred_at
            if _oa:
                _by_date.setdefault(_oa.date(), []).append(_r)
        _weeks = []
        for _wk in _cal_mod.Calendar(firstweekday=0).monthdatescalendar(cal_year_num, cal_month_num):
            _weeks.append([
                {"date": _d, "in_month": (_d.month == cal_month_num), "rows": _by_date.get(_d, [])}
                for _d in _wk
            ])
        _MONTH_RU = ["", "Январь","Февраль","Март","Апрель","Май","Июнь","Июль","Август","Сентябрь","Октябрь","Ноябрь","Декабрь"]
        _prev_y, _prev_m = (cal_year_num - 1, 12) if cal_month_num == 1 else (cal_year_num, cal_month_num - 1)
        _next_y, _next_m = (cal_year_num + 1, 1) if cal_month_num == 12 else (cal_year_num, cal_month_num + 1)
        cal_ctx = {
            "weeks": _weeks,
            "month_label": "{} {}".format(_MONTH_RU[cal_month_num], cal_year_num),
            "month_value": "{}-{:02d}".format(cal_year_num, cal_month_num),
            "prev_month_value": "{}-{:02d}".format(_prev_y, _prev_m),
            "next_month_value": "{}-{:02d}".format(_next_y, _next_m),
            "today": _today_msk,
        }
    elif is_kanban:
        # Лимит 600 (по 200 на колонку макс). Группируем в Python по 3 бакетам.
        _KANBAN_LIMIT = 600
        incidents = iq.limit(_KANBAN_LIMIT).all()
        rows = _build_incident_rows(incidents, include_author=True)
        # Подгрузим счётчики задач одним запросом (Task.incident_id backref с-63b).
        from app.models.tasks import Task as _Task
        inc_ids = [r["inc"].id for r in rows]
        if inc_ids:
            _DONE_STATUSES = ("Выполнена", "Закрыта", "Отменена")
            for tid, st in db.session.query(_Task.incident_id, _Task.status).filter(_Task.incident_id.in_(inc_ids)).all():
                slot = task_counts_map.setdefault(tid, {"total": 0, "done": 0})
                slot["total"] += 1
                if st in _DONE_STATUSES:
                    slot["done"] += 1
        kanban_groups = {"incoming": [], "in_work": [], "completed": []}
        for r in rows:
            s = r["inc"].status or "new"
            if s == "new":
                bucket = "incoming"
            elif s in ("assigned", "in_progress"):
                bucket = "in_work"
            else:
                bucket = "completed"
            r["task_total"] = task_counts_map.get(r["inc"].id, {}).get("total", 0)
            r["task_done"]  = task_counts_map.get(r["inc"].id, {}).get("done", 0)
            kanban_groups[bucket].append(r)
        total_count = len(rows)
        total_pages = 1
        page = 1
        rows_limit_reached = (total_count >= _KANBAN_LIMIT)
        _LIMIT = _KANBAN_LIMIT
    else:
        total_count = iq.order_by(None).with_entities(func.count(func.distinct(Incident.id))).scalar() or 0
        total_pages = max(1, (total_count + _PER_PAGE - 1) // _PER_PAGE)
        if page > total_pages:
            page = total_pages
        incidents = iq.limit(_PER_PAGE).offset((page - 1) * _PER_PAGE).all()
        rows = _build_incident_rows(incidents, include_author=True)
        # Совместимость с шаблоном с-64 (rows_limit_reached) — теперь не нужно, оставляем False.
        rows_limit_reached = False
        _LIMIT = _PER_PAGE

    # Счётчики по вкладкам. Для social-view: 3 общие вкладки показывают кол-во по школе,
    # 4-я «Назначенные мне» — только мои активные.
    def _count(status_list):
        return db.session.query(Incident).filter(Incident.status.in_(status_list)).count()
    counters = {
        "incoming":  _count(STATUS_BUCKETS["incoming"]),
        "in_work":   _count(STATUS_BUCKETS["in_work"]),
        "completed": _count(STATUS_BUCKETS["completed"]),
    }
    if is_social_view:
        counters["mine"] = (
            db.session.query(Incident)
            .filter(Incident.assignees.any(id=uid), Incident.status != "closed")
            .count()
        )

    # Список классов для фильтра
    classes = []
    if year:
        classes = (
            SchoolClass.query
            .filter(SchoolClass.academic_year_id == year.id)
            .order_by(
                SchoolClass.grade.asc().nullslast(),
                SchoolClass.letter.asc().nullslast(),
                SchoolClass.name.asc(),
            )
            .all()
        )

    # Список исполнителей для picker-а в строке таблицы.
    # Двойная схема ролей: user.role (старая) ИЛИ user_role→role (новая).
    # На локали user_role часто пустая — полагаемся на user.role. На проде
    # наоборот — основная — user_role. Ищем по обеим сразу через OR.
    from app.models_legacy import User as _User, UserRole as _UserRole, Role as _Role
    _ROLE_CODES = [
        "ADMIN", "DEPUTY_DIRECTOR", "PSYCHOLOGIST", "SOCIAL_PEDAGOG",
        "METHODIST", "CLASS_TEACHER", "TEACHER",
    ]
    _users_via_ur = (
        db.session.query(_User.id)
        .join(_UserRole, _UserRole.user_id == _User.id)
        .join(_Role, _Role.id == _UserRole.role_id)
        .filter(_Role.code.in_(_ROLE_CODES))
    )
    assignees = (
        _User.query
        .filter(
            _User.role.in_(_ROLE_CODES) | _User.id.in_(_users_via_ur)
        )
        .order_by(_User.last_name.asc().nullslast(), _User.first_name.asc().nullslast())
        .all()
    )

    # Лейблы вкладок
    tab_labels = {
        "incoming":  "Входящие",
        "in_work":   "В работе",
        "completed": "Завершённые",
    }
    if is_social_view:
        tab_labels["mine"] = "Назначенные мне"

    # Для social-view: edit/delete в строках доступны только на «Назначенные мне».
    # На общих 3 вкладках соц.педагог в режиме read-only (как просил пользователь —
    # видит очередь по школе по аналогии с дашбордом).
    # METHODIST — может редактировать только свои (автор/исполнитель). На остальных
    # карандаш заменяется на глаз, корзина прячется. Per-row проверка в шаблоне.
    if is_methodist_view:
        can_edit_rows = True
    elif is_social_view:
        can_edit_rows = (active_tab == "mine")
    else:
        can_edit_rows = True

    # Группировка для table/list (kanban уже сгруппирован по статусу-колонкам).
    groups = None
    if f_group_by and not is_kanban:
        groups = _group_incident_rows(rows, f_group_by, Incident.STATUS_LABELS)

    return render_template(
        "incidents_my.html",
        is_admin_view=True,
        is_social_view=is_social_view,
        rows=rows,
        groups=groups,
        f_group_by=f_group_by,
        counters=counters,
        active_tab=active_tab,
        tab_labels=tab_labels,
        f_category=f_category,
        f_class_id=f_class_id,
        f_q=f_q,
        f_sort=f_sort,
        categories=INCIDENT_CATEGORIES,
        classes=classes,
        assignees=assignees,
        can_change_status=(not is_social_view) and (not is_methodist_view),
        can_edit_rows=can_edit_rows,
        status_labels=Incident.STATUS_LABELS,
        current_user_id=uid,
        is_admin=is_admin(current_user),
        view_mode=view_mode_raw,
        kanban_groups=kanban_groups,
        rows_limit_reached=rows_limit_reached,
        rows_limit=_LIMIT,
        page=page,
        total_pages=total_pages,
        total_count=total_count,
        per_page=_PER_PAGE,
        is_methodist_view=is_methodist_view,
        cal_ctx=cal_ctx,
    )


@children_bp.route("/incidents/dashboard")
@require_roles("ADMIN", "METHODIST", "SOCIAL_PEDAGOG", "PSYCHOLOGIST")
def incidents_dashboard():
    return redirect(url_for("children.incidents_dashboard_legacy", **request.args))


@children_bp.route("/incidents/dashboard-legacy")
@require_roles("ADMIN", "METHODIST", "SOCIAL_PEDAGOG", "PSYCHOLOGIST")
def incidents_dashboard_legacy():
    grade = request.args.get("grade", type=int)
    class_id = request.args.get("class_id", type=int)
    category = (request.args.get("category") or "").strip()
    status_filter = (request.args.get("status") or "").strip()

    year = _get_current_year()
    year_id = year.id if year else None

    base = (
        db.session.query(Incident)
        .join(IncidentChild, IncidentChild.incident_id == Incident.id)
        .join(Child, Child.id == IncidentChild.child_id)
    )

    if category:
        base = base.filter(Incident.category == category)
    if status_filter:
        base = base.filter(Incident.status == status_filter)

    if year_id and (grade or class_id):
        base = base.join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        base = base.filter(
            ChildEnrollment.academic_year_id == year_id,
            ChildEnrollment.ended_at.is_(None)
        )
        if class_id:
            base = base.filter(ChildEnrollment.school_class_id == class_id)
        elif grade:
            base = base.join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            base = base.filter(SchoolClass.grade == grade)

    base = base.distinct()

    now = datetime.utcnow()
    d7 = now - timedelta(days=7)
    d30 = now - timedelta(days=30)

    total_all = base.count()
    total_7 = base.filter(Incident.occurred_at >= d7).count()
    total_30 = base.filter(Incident.occurred_at >= d30).count()

    top_categories = (
        db.session.query(Incident.category, func.count(func.distinct(Incident.id)))
        .select_from(Incident)
        .join(IncidentChild, IncidentChild.incident_id == Incident.id)
        .join(Child, Child.id == IncidentChild.child_id)
        .filter(Incident.occurred_at >= d30)
    )

    if category:
        top_categories = top_categories.filter(Incident.category == category)

    if year_id and (grade or class_id):
        top_categories = top_categories.join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        top_categories = top_categories.filter(
            ChildEnrollment.academic_year_id == year_id,
            ChildEnrollment.ended_at.is_(None)
        )
        if class_id:
            top_categories = top_categories.filter(ChildEnrollment.school_class_id == class_id)
        elif grade:
            top_categories = top_categories.join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            top_categories = top_categories.filter(SchoolClass.grade == grade)

    top_categories = (
        top_categories
        .group_by(Incident.category)
        .order_by(func.count(func.distinct(Incident.id)).desc())
        .limit(10)
        .all()
    )

    top_classes = []
    if year_id:
        tc = (
            db.session.query(
                SchoolClass.name,
                func.count(func.distinct(Incident.id)),
                SchoolClass.id,
                Building.name,
            )
            .select_from(Incident)
            .join(IncidentChild, IncidentChild.incident_id == Incident.id)
            .join(Child, Child.id == IncidentChild.child_id)
            .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
            .join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            .outerjoin(Building, Building.id == SchoolClass.building_id)
            .filter(
                Incident.occurred_at >= d30,
                ChildEnrollment.academic_year_id == year_id,
                ChildEnrollment.ended_at.is_(None)
            )
        )

        if category:
            tc = tc.filter(Incident.category == category)
        if class_id:
            tc = tc.filter(SchoolClass.id == class_id)
        elif grade:
            tc = tc.filter(SchoolClass.grade == grade)

        top_classes = (
            tc.group_by(SchoolClass.name, SchoolClass.id, Building.name)
            .order_by(func.count(func.distinct(Incident.id)).desc())
            .limit(10)
            .all()
        )

    top_buildings = []
    if year_id:
        tb = (
            db.session.query(Building.name, func.count(func.distinct(Incident.id)))
            .select_from(Incident)
            .join(IncidentChild, IncidentChild.incident_id == Incident.id)
            .join(Child, Child.id == IncidentChild.child_id)
            .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
            .join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            .outerjoin(Building, Building.id == SchoolClass.building_id)
            .filter(
                Incident.occurred_at >= d30,
                ChildEnrollment.academic_year_id == year_id,
                ChildEnrollment.ended_at.is_(None)
            )
        )

        if category:
            tb = tb.filter(Incident.category == category)
        if class_id:
            tb = tb.filter(SchoolClass.id == class_id)
        elif grade:
            tb = tb.filter(SchoolClass.grade == grade)

        top_buildings = (
            tb.group_by(Building.name)
            .order_by(func.count(func.distinct(Incident.id)).desc())
            .limit(10)
            .all()
        )

    # daily counts — 1 query with GROUP BY instead of 7 separate COUNTs
    week_start = (now - timedelta(days=6)).replace(hour=0, minute=0, second=0, microsecond=0)
    prev_week_start = week_start - timedelta(days=7)
    _day_expr = func.date(Incident.occurred_at)
    daily_q = (
        base.with_entities(
            _day_expr.label("day"),
            func.count(func.distinct(Incident.id)),
        )
        .filter(Incident.occurred_at >= week_start)
        .group_by(_day_expr)
        .all()
    )
    daily_map = {str(d): c for d, c in daily_q}
    recent_daily = []
    daily_labels = []
    max_daily = 0
    for i in range(6, -1, -1):
        day = (now - timedelta(days=i)).date()
        cnt = daily_map.get(str(day), 0)
        recent_daily.append({"label": day.strftime("%d.%m"), "count": cnt})
        daily_labels.append(day.strftime("%d.%m"))
        max_daily = max(max_daily, cnt)

    # Сравнение с предыдущей 7-дневкой (для подписи «vs прошлая неделя»).
    week_total = sum(item["count"] for item in recent_daily)
    prev_week_total = (
        base.filter(
            Incident.occurred_at >= prev_week_start,
            Incident.occurred_at < week_start,
        ).count()
    )
    if prev_week_total > 0:
        week_delta_pct = round((week_total - prev_week_total) * 100 / prev_week_total)
    elif week_total > 0:
        week_delta_pct = None  # нет базы для сравнения
    else:
        week_delta_pct = 0

    max_category = max([cnt for _, cnt in top_categories], default=0)
    max_class = max([cnt for _, cnt, _c, _b in top_classes], default=0)
    max_building = max([cnt for _, cnt in top_buildings], default=0)

    _ua_q = (
        db.session.query(
            User.last_name,
            User.first_name,
            User.middle_name,
            User.username,
            func.count(func.distinct(Incident.id)).label("cnt")
        )
        .select_from(Incident)
        .join(User, User.id == Incident.author_id)
        .join(IncidentChild, IncidentChild.incident_id == Incident.id)
        .join(Child, Child.id == IncidentChild.child_id)
        .filter(Incident.occurred_at >= d30)
    )
    if year_id and (grade or class_id):
        _ua_q = _ua_q.join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        _ua_q = _ua_q.filter(
            ChildEnrollment.academic_year_id == year_id,
            ChildEnrollment.ended_at.is_(None)
        )
        if class_id:
            _ua_q = _ua_q.filter(ChildEnrollment.school_class_id == class_id)
        elif grade:
            _ua_q = _ua_q.join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
            _ua_q = _ua_q.filter(SchoolClass.grade == grade)
    if category:
        _ua_q = _ua_q.filter(Incident.category == category)
    _ua_rows = (
        _ua_q
        .group_by(User.id, User.last_name, User.first_name, User.middle_name, User.username)
        .order_by(func.count(func.distinct(Incident.id)).desc())
        .limit(10)
        .all()
    )
    user_activity = [
        {
            "name": " ".join(p for p in [r.last_name or "", r.first_name or "", r.middle_name or ""] if p.strip()) or r.username,
            "cnt": r.cnt,
        }
        for r in _ua_rows
    ]
    max_user_activity = max((r["cnt"] for r in user_activity), default=0)

    # status_breakdown — 1 query with GROUP BY instead of 3 separate COUNTs
    _sb_q = (
        base.with_entities(Incident.status, func.count(func.distinct(Incident.id)))
        .group_by(Incident.status)
        .all()
    )
    _sb_map = {s: c for s, c in _sb_q}
    status_breakdown = {
        "new": _sb_map.get("new", 0),
        "assigned": _sb_map.get("assigned", 0),
        "in_progress": _sb_map.get("in_progress", 0),
        "resolved": _sb_map.get("resolved", 0),
        "closed": _sb_map.get("closed", 0),
    }
    open_count = (
        status_breakdown["new"]
        + status_breakdown["assigned"]
        + status_breakdown["in_progress"]
    )

    recent = base.order_by(Incident.occurred_at.desc(), Incident.id.desc()).limit(20).all()
    rows = _build_incident_rows(recent, include_author=True)

    classes = (
        SchoolClass.query
        .filter(SchoolClass.academic_year_id == year.id)
        .order_by(
            SchoolClass.grade.asc().nullslast(),
            SchoolClass.letter.asc().nullslast(),
            SchoolClass.name.asc()
        )
        .all()
    )

    return render_template(
        "incidents_dashboard.html",
        title="Инциденты — дашборд",
        total_all=total_all,
        total_7=total_7,
        total_30=total_30,
        top_categories=top_categories,
        top_classes=top_classes,
        top_buildings=top_buildings,
        rows=rows,
        recent_daily=recent_daily,
        max_daily=max_daily,
        week_total=week_total,
        prev_week_total=prev_week_total,
        week_delta_pct=week_delta_pct,
        max_category=max_category,
        max_class=max_class,
        max_building=max_building,
        grade=grade,
        class_id=class_id,
        category=category,
        categories=INCIDENT_CATEGORIES,
        classes=classes,
        user_activity=user_activity,
        max_user_activity=max_user_activity,
        status_breakdown=status_breakdown,
        open_count=open_count,
        can_change_status=_can_change_status(),
        can_edit_any=(_can_change_status() or has_role("SOCIAL_PEDAGOG")),
        status_filter=status_filter,
    )


# =========================================================
# BUILDINGS
# =========================================================
@children_bp.route("/buildings")
@require_roles("ADMIN")
def buildings_registry():
    buildings = Building.query.order_by(Building.name.asc()).all()
    return render_template(
        "buildings_list.html",
        buildings=buildings,
        building_tone_choices=BUILDING_MATRIX_TONE_CHOICES,
    )


@children_bp.route("/buildings/new", methods=["POST"])
@require_roles("ADMIN")
def buildings_new():
    name = (request.form.get("name") or "").strip()
    address = (request.form.get("address") or "").strip() or None
    short_name = (request.form.get("short_name") or "").strip() or None
    matrix_tone = normalize_building_matrix_tone(
        request.form.get("matrix_tone")
    )

    if not name:
        flash("Укажите название здания", "danger")
        return redirect(url_for("children.buildings_registry"))

    db.session.add(Building(
        name=name,
        address=address,
        short_name=short_name,
        matrix_tone=matrix_tone,
    ))
    db.session.commit()
    flash("Здание добавлено", "success")
    return redirect(url_for("children.buildings_registry"))


@children_bp.route("/buildings/<int:building_id>/update", methods=["POST"])
@require_roles("ADMIN")
def buildings_update(building_id: int):
    b = Building.query.get_or_404(building_id)

    b.name = (request.form.get("name") or "").strip()
    b.short_name = (request.form.get("short_name") or "").strip() or None
    b.address = (request.form.get("address") or "").strip() or None
    b.matrix_tone = normalize_building_matrix_tone(
        request.form.get("matrix_tone")
    )

    if not b.name:
        flash("Название здания не может быть пустым", "danger")
        return redirect(url_for("children.buildings_registry"))

    db.session.commit()
    flash("Сохранено", "success")
    return redirect(url_for("children.buildings_registry"))


@children_bp.route("/buildings/<int:building_id>/delete", methods=["POST"])
@require_roles("ADMIN")
def buildings_delete(building_id: int):
    b = Building.query.get_or_404(building_id)

    SchoolClass.query.filter_by(building_id=b.id).update({"building_id": None})
    db.session.delete(b)
    db.session.commit()

    flash("Здание удалено", "success")
    return redirect(url_for("children.buildings_registry"))


# =========================================================
# SOCIAL PASSPORT REGISTRY / DASHBOARD
# =========================================================
@children_bp.route("/social-passport")
@login_required
def social_passport_registry():
    if not has_permission("social_passport_registry_view"):
        abort(403)

    year = _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    grade = parse_int(request.args.get("grade"))
    class_id = parse_int(request.args.get("class_id"))
    q_text = (request.args.get("q") or "").strip()

    # Response-кеш для тяжёлого реестра (~8 МБ HTML, ~1.3 с рендера).
    # Ключ: роли пользователя (для CLASS_TEACHER — плюс user_id, т.к. он видит только свои классы)
    #       + фильтры + year_id. TTL 60 с.
    is_class_teacher_user = has_role("CLASS_TEACHER")
    is_admin_user = has_role("ADMIN")
    is_methodist_user = has_role("METHODIST")
    is_social_pedagog_user = has_role("SOCIAL_PEDAGOG")
    scope_own_classes = (
        is_class_teacher_user and not (is_admin_user or is_methodist_user or is_social_pedagog_user)
    )
    role_bucket = "own" if scope_own_classes else "all"
    cache_owner = current_user.id if scope_own_classes else 0
    cache_key = make_key(
        "social_passport_registry", role_bucket, cache_owner, year.id, grade, class_id, q_text
    )
    cached_html = view_response_cache.get(cache_key)
    if cached_html is not None:
        from flask import Response
        return Response(cached_html, mimetype="text/html; charset=utf-8")

    q = (
        Child.query
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == year.id)
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
        .options(
            joinedload(Child.social),
            subqueryload(Child.parent_links).subqueryload(ChildParent.parent),
        )
        .add_columns(SchoolClass.name.label("_cls_name"))
    )

    if scope_own_classes:
        q = q.filter(SchoolClass.teacher_user_id == current_user.id)

    if grade is not None:
        q = q.filter(SchoolClass.grade == grade)
    if class_id:
        q = q.filter(SchoolClass.id == class_id)
    if q_text:
        like = f"%{q_text}%"
        q = q.filter(db.or_(
            Child.last_name.ilike(like),
            Child.first_name.ilike(like),
            Child.middle_name.ilike(like),
        ))

    classes_query = SchoolClass.query.filter_by(academic_year_id=year.id)
    if scope_own_classes:
        classes_query = classes_query.filter(SchoolClass.teacher_user_id == current_user.id)
    if grade is not None:
        classes_query = classes_query.filter(SchoolClass.grade == grade)
    classes = classes_query.order_by(SchoolClass.grade.asc().nullslast(), SchoolClass.name.asc()).all()
    grades = sorted({c.grade for c in classes if c.grade is not None})

    rows = q.order_by(SchoolClass.grade.asc().nullslast(), SchoolClass.name.asc(), Child.last_name.asc(), Child.first_name.asc()).all()

    # Кэшируем class_name из add_columns, чтобы шаблон не делал N+1
    children = []
    for row in rows:
        ch = row[0]
        ch._cls_name = row[1] or ""
        children.append(ch)

    selected_class = next((c for c in classes if c.id == class_id), None) if class_id else None

    html = render_template(
        "social_passport_registry.html",
        children=children,
        year=year,
        classes=classes,
        grades=grades,
        selected_grade=grade,
        selected_class_id=class_id,
        selected_class=selected_class,
        q_text=q_text,
        is_admin=is_admin_user,
        is_methodist=is_methodist_user,
        is_class_teacher=is_class_teacher_user,
        has_role_social_pedagog=is_social_pedagog_user,
    )
    view_response_cache.set(cache_key, html, timeout=60)
    return html


@children_bp.route("/comments/registry")
@login_required
def comments_registry():
    if not has_permission("children_registry_view"):
        abort(403)

    year = _get_current_year()
    q_text = (request.args.get("q") or "").strip()
    grade = parse_int(request.args.get("grade"))
    class_id = parse_int(request.args.get("class_id"))

    q = (
        ChildComment.query
        .join(Child, Child.id == ChildComment.child_id)
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == (year.id if year else 0))
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(User, User.id == ChildComment.author_id)
    )

    if has_role("CLASS_TEACHER") and not (has_role("ADMIN") or has_role("METHODIST")):
        q = q.filter(SchoolClass.teacher_user_id == current_user.id)

    if grade is not None:
        q = q.filter(SchoolClass.grade == grade)
    if class_id:
        q = q.filter(SchoolClass.id == class_id)
    if q_text:
        like = f"%{q_text}%"
        q = q.filter(db.or_(
            Child.last_name.ilike(like), Child.first_name.ilike(like), Child.middle_name.ilike(like),
            ChildComment.text.ilike(like), User.last_name.ilike(like), User.first_name.ilike(like)
        ))

    classes_q = SchoolClass.query
    if year:
        classes_q = classes_q.filter_by(academic_year_id=year.id)
    if has_role("CLASS_TEACHER") and not (has_role("ADMIN") or has_role("METHODIST")):
        classes_q = classes_q.filter(SchoolClass.teacher_user_id == current_user.id)
    if grade is not None:
        classes_q = classes_q.filter(SchoolClass.grade == grade)
    classes = classes_q.order_by(SchoolClass.grade.asc().nullslast(), SchoolClass.name.asc()).all()
    grades = sorted({c.grade for c in classes if c.grade is not None})
    comments = q.order_by(ChildComment.created_at.desc()).all()

    return render_template(
        "comments_registry.html",
        comments=comments,
        classes=classes,
        grades=grades,
        selected_grade=grade,
        selected_class_id=class_id,
        q_text=q_text,
    )


@children_bp.route("/social-passport/dashboard")
@require_roles("ADMIN", "METHODIST", "SOCIAL_PEDAGOG")
def social_passport_dashboard():
    year = _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    # Single query: Child + ChildSocial + Building.name via JOIN (no lazy loading)
    rows = (
        db.session.query(ChildSocial, Building.name.label("bname"))
        .select_from(Child)
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == year.id)
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(Building, Building.id == SchoolClass.building_id)
        .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
        .add_columns(Child.id)
        .distinct()
        .all()
    )

    totals = {
        "school_total": 0,
        "large_family": 0,
        "low_income": 0,
        "guardianship": 0,
        "orphan": 0,
        "parents_disability": 0,
        "socially_dangerous": 0,
        "hard_life": 0,
    }
    by_building = {}

    for social, bname, _child_id in rows:
        totals["school_total"] += 1
        totals["large_family"] += 1 if social and social.has_large_family else 0
        totals["low_income"] += 1 if social and social.has_low_income_family else 0
        totals["guardianship"] += 1 if social and social.has_guardianship else 0
        totals["orphan"] += 1 if social and social.has_orphan_status else 0
        totals["parents_disability"] += 1 if social and social.has_disability_parents else 0
        totals["socially_dangerous"] += 1 if social and social.is_socially_dangerous else 0
        totals["hard_life"] += 1 if social and social.is_hard_life else 0

        bkey = bname or "Без здания"
        if bkey not in by_building:
            by_building[bkey] = {
                "total": 0, "large_family": 0, "low_income": 0,
                "guardianship": 0, "orphan": 0, "parents_disability": 0,
                "socially_dangerous": 0, "hard_life": 0,
            }
        row = by_building[bkey]
        row["total"] += 1
        row["large_family"] += 1 if social and social.has_large_family else 0
        row["low_income"] += 1 if social and social.has_low_income_family else 0
        row["guardianship"] += 1 if social and social.has_guardianship else 0
        row["orphan"] += 1 if social and social.has_orphan_status else 0
        row["parents_disability"] += 1 if social and social.has_disability_parents else 0
        row["socially_dangerous"] += 1 if social and social.is_socially_dangerous else 0
        row["hard_life"] += 1 if social and social.is_hard_life else 0

    return render_template(
        "social_passport_dashboard.html",
        totals=totals,
        by_building=by_building
    )


# =========================================================
# SOCIAL PASSPORT — class report + school summary (s87)
# =========================================================
# 13 категорий по шаблону «Соц паспорт новый.docx».
# Ключи стабильные — используются и в фронте, и в backend-агрегациях.
SOCIAL_PASSPORT_CATEGORIES = [
    ("large_family",       "Многодетные"),
    ("incomplete",         "Неполная семья"),
    ("single_mother",      "Мать-одиночка"),
    ("single_father",      "Отец-одиночка"),
    ("parents_disability", "Родители-инвалиды"),
    ("child_disability",   "Дети-инвалиды"),
    ("guardianship",       "Опека"),
    ("repeat_year",        "Повторный курс"),
    ("ovz",                "ОВЗ"),
    ("low_income",         "Малообеспеченные"),
    ("vshu_kdn_pdn",       "ВШУ/КДН/ПДН"),
    ("orphan",             "Сирота"),
    ("svo_family",         "Семья СВО"),
]


def _passport_flags_for_child(child):
    """Возвращает dict с булями по 13 категориям соц.паспорта.
    `incomplete` определяется по тексту family_status (если содержит 'неполн').
    """
    s = child.social
    fs = ((s.family_status if s else "") or "").lower()
    return {
        "large_family":       bool(s and s.has_large_family),
        "incomplete":         "неполн" in fs,
        "single_mother":      bool(s and s.is_single_mother),
        "single_father":      bool(s and s.is_single_father),
        "parents_disability": bool(s and s.has_disability_parents),
        "child_disability":   bool(child.is_disabled),
        "guardianship":       bool(s and s.has_guardianship),
        "repeat_year":        bool(s and s.is_repeat_year),
        "ovz":                bool(child.is_ovz),
        "low_income":         bool(s and s.has_low_income_family),
        "vshu_kdn_pdn":       bool(s and (s.vshu_since or s.kdn_since or s.pdn_since)),
        "orphan":             bool(s and s.has_orphan_status),
        "svo_family":         bool(s and s.is_svo_family),
    }


def _empty_passport_totals():
    t = {k: 0 for k, _ in SOCIAL_PASSPORT_CATEGORIES}
    t["total"] = 0
    t["boys"] = 0
    t["girls"] = 0
    return t


def _accumulate_passport_totals(totals, child, flags):
    totals["total"] += 1
    g = (child.gender or "").lower()
    if g.startswith("м"):
        totals["boys"] += 1
    elif g.startswith("ж") or g.startswith("д"):
        totals["girls"] += 1
    for k, _ in SOCIAL_PASSPORT_CATEGORIES:
        if flags[k]:
            totals[k] += 1


def _can_view_class_passport(school_class):
    """ADMIN/METHODIST/SOCIAL_PEDAGOG — любой класс. CLASS_TEACHER — только свой."""
    if has_role("ADMIN") or has_role("METHODIST") or has_role("SOCIAL_PEDAGOG"):
        return True
    if has_role("CLASS_TEACHER"):
        return school_class.teacher_user_id == getattr(current_user, "id", None)
    return False


@children_bp.route("/social-passport/class/<int:class_id>")
@login_required
def social_passport_class(class_id: int):
    """Соц.паспорт класса — таблица по шаблону директора.
    Класс. рук видит свой класс, ADMIN/METHODIST/SOCIAL_PEDAGOG — любой."""
    if not has_permission("social_passport_registry_view"):
        abort(403)

    year = _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    school_class = SchoolClass.query.get_or_404(class_id)
    if not _can_view_class_passport(school_class):
        abort(403)

    children = (
        db.session.query(Child)
        .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.school_class_id == class_id,
        )
        .options(
            joinedload(Child.social),
            subqueryload(Child.parent_links).subqueryload(ChildParent.parent),
        )
        .order_by(Child.last_name.asc(), Child.first_name.asc())
        .all()
    )

    rows = []
    totals = _empty_passport_totals()
    for ch in children:
        flags = _passport_flags_for_child(ch)
        _accumulate_passport_totals(totals, ch, flags)
        rows.append({"child": ch, "flags": flags})

    teacher = school_class.teacher_user
    can_view_summary = (has_role("ADMIN") or has_role("METHODIST") or has_role("SOCIAL_PEDAGOG"))

    return render_template(
        "social_passport_class.html",
        school_class=school_class,
        year=year,
        rows=rows,
        totals=totals,
        teacher=teacher,
        categories=SOCIAL_PASSPORT_CATEGORIES,
        can_view_summary=can_view_summary,
    )


@children_bp.route("/social-passport/summary")
@login_required
def social_passport_summary():
    """Сводный соц.паспорт по школе. 4 разреза (вкладки):
    classes / parallels / buildings / school. ADMIN/METHODIST/SOCIAL_PEDAGOG.

    Перфоматирование (s87): один запрос колонками (без ORM-объектов)
    + response-кеш TTL 60с по (tab, year_id). Бенч до: 600ms, после: ~30ms.
    """
    if not (has_role("ADMIN") or has_role("METHODIST") or has_role("SOCIAL_PEDAGOG")):
        abort(403)

    year = _get_current_year()
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    tab = (request.args.get("tab") or "classes").strip()
    if tab not in ("classes", "parallels", "buildings", "school"):
        tab = "classes"

    cache_key = make_key("social_passport_summary", tab, year.id)
    cached_html = view_response_cache.get(cache_key)
    if cached_html is not None:
        from flask import Response
        return Response(cached_html, mimetype="text/html; charset=utf-8")

    # Один SELECT по колонкам — без загрузки 3400 ORM-объектов Child/Social.
    rows = (
        db.session.query(
            Child.gender, Child.is_disabled, Child.is_ovz,
            SchoolClass.id, SchoolClass.name, SchoolClass.grade,
            Building.name,
            ChildSocial.family_status,
            ChildSocial.has_large_family, ChildSocial.has_disability_parents,
            ChildSocial.has_low_income_family, ChildSocial.has_guardianship,
            ChildSocial.has_orphan_status, ChildSocial.is_socially_dangerous,
            ChildSocial.is_hard_life,
            ChildSocial.is_single_mother, ChildSocial.is_single_father,
            ChildSocial.is_repeat_year, ChildSocial.is_svo_family,
            ChildSocial.vshu_since, ChildSocial.kdn_since, ChildSocial.pdn_since,
        )
        .select_from(Child)
        .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        .join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(Building, Building.id == SchoolClass.building_id)
        .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None),
        )
        .all()
    )

    by_class = {}
    by_grade = {}
    by_building = {}
    school_totals = _empty_passport_totals()

    cat_keys = [k for k, _ in SOCIAL_PASSPORT_CATEGORIES]

    for (gender, is_disabled, is_ovz,
         sc_id, sc_name, sc_grade, bld_name,
         family_status,
         large_family, disability_parents, low_income, guardianship,
         orphan, socially_dangerous, hard_life,
         single_mother, single_father, repeat_year, svo_family,
         vshu_since, kdn_since, pdn_since) in rows:

        fs = (family_status or "").lower()
        flags = {
            "large_family":       bool(large_family),
            "incomplete":         "неполн" in fs,
            "single_mother":      bool(single_mother),
            "single_father":      bool(single_father),
            "parents_disability": bool(disability_parents),
            "child_disability":   bool(is_disabled),
            "guardianship":       bool(guardianship),
            "repeat_year":        bool(repeat_year),
            "ovz":                bool(is_ovz),
            "low_income":         bool(low_income),
            "vshu_kdn_pdn":       bool(vshu_since or kdn_since or pdn_since),
            "orphan":             bool(orphan),
            "svo_family":         bool(svo_family),
        }

        def _bump(t):
            t["total"] += 1
            g = (gender or "").lower()
            if g.startswith("м"):
                t["boys"] += 1
            elif g.startswith("ж") or g.startswith("д"):
                t["girls"] += 1
            for k in cat_keys:
                if flags[k]:
                    t[k] += 1

        _bump(school_totals)

        if sc_id not in by_class:
            by_class[sc_id] = {"label": sc_name, "grade": sc_grade, "id": sc_id,
                               "totals": _empty_passport_totals()}
        _bump(by_class[sc_id]["totals"])

        gk = sc_grade if sc_grade is not None else 0
        if gk not in by_grade:
            by_grade[gk] = {"label": (f"{gk} класс" if gk else "Без параллели"),
                            "totals": _empty_passport_totals()}
        _bump(by_grade[gk]["totals"])

        bk = bld_name or "Без корпуса"
        if bk not in by_building:
            by_building[bk] = {"label": bk, "totals": _empty_passport_totals()}
        _bump(by_building[bk]["totals"])

    classes_list = sorted(by_class.values(), key=lambda r: ((r["grade"] or 99), r["label"]))
    parallels_list = sorted(by_grade.values(), key=lambda r: r["label"])
    buildings_list = sorted(by_building.values(), key=lambda r: r["label"])

    html = render_template(
        "social_passport_summary.html",
        tab=tab,
        year=year,
        categories=SOCIAL_PASSPORT_CATEGORIES,
        classes_list=classes_list,
        parallels_list=parallels_list,
        buildings_list=buildings_list,
        school_totals=school_totals,
    )
    view_response_cache.set(cache_key, html, timeout=60)
    return html


def _xlsx_response(wb, filename):
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    from flask import send_file
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@children_bp.route("/social-passport/class/<int:class_id>/export.xlsx")
@login_required
def social_passport_class_export(class_id: int):
    """Excel-выгрузка соц.паспорта одного класса (повторяет docx-таблицу)."""
    if not has_permission("social_passport_registry_view"):
        abort(403)

    year = _get_current_year()
    if not year:
        abort(404)

    school_class = SchoolClass.query.get_or_404(class_id)
    if not _can_view_class_passport(school_class):
        abort(403)

    children = (
        db.session.query(Child)
        .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None),
            ChildEnrollment.school_class_id == class_id,
        )
        .options(
            joinedload(Child.social),
            subqueryload(Child.parent_links).subqueryload(ChildParent.parent),
        )
        .order_by(Child.last_name.asc(), Child.first_name.asc())
        .all()
    )

    wb = Workbook()
    ws = wb.active
    ws.title = (school_class.name or f"Класс {class_id}")[:31]

    headers = ["№", "ФИО", "Телефон", "Дата рождения", "Адрес рег.", "Адрес факт.",
               "Мать (ФИО, тел.)", "Отец (ФИО, тел.)"] + [lbl for _, lbl in SOCIAL_PASSPORT_CATEGORIES]
    ws.append(headers)

    totals = _empty_passport_totals()
    for i, ch in enumerate(children, start=1):
        flags = _passport_flags_for_child(ch)
        _accumulate_passport_totals(totals, ch, flags)
        row = [
            i,
            ch.fio,
            ch.phone or "",
            ch.birth_date.strftime("%d.%m.%Y") if ch.birth_date else "",
            ch.reg_address or "",
            ch.actual_address or "",
            (ch.mother_fio or "") + (f", {ch.mother_phone}" if ch.mother_phone else ""),
            (ch.father_fio or "") + (f", {ch.father_phone}" if ch.father_phone else ""),
        ]
        for k, _ in SOCIAL_PASSPORT_CATEGORIES:
            row.append("✓" if flags[k] else "")
        ws.append(row)

    # Строка итогов
    totals_row = [
        "", f"Итого: {totals['total']} (М: {totals['boys']}, Д: {totals['girls']})",
        "", "", "", "", "", "",
    ]
    for k, _ in SOCIAL_PASSPORT_CATEGORIES:
        totals_row.append(totals[k])
    ws.append(totals_row)

    teacher = school_class.teacher_user
    if teacher:
        teacher_line = teacher.fio
        if teacher.phone:
            teacher_line += f", {teacher.phone}"
        elif teacher.email:
            teacher_line += f", {teacher.email}"
        ws.append([])
        ws.append(["Классный руководитель:", teacher_line])

    ws.column_dimensions["B"].width = 32
    ws.column_dimensions["E"].width = 30
    ws.column_dimensions["F"].width = 30
    ws.column_dimensions["G"].width = 30
    ws.column_dimensions["H"].width = 30

    fname = f"social_passport_class_{school_class.name or class_id}.xlsx"
    return _xlsx_response(wb, fname)


@children_bp.route("/social-passport/summary/export.xlsx")
@login_required
def social_passport_summary_export():
    """Excel-выгрузка сводного соц.паспорта.
    `?tab=` определяет лист (по умолчанию выгружаем все 4 листа)."""
    if not (has_role("ADMIN") or has_role("METHODIST") or has_role("SOCIAL_PEDAGOG")):
        abort(403)

    year = _get_current_year()
    if not year:
        abort(404)

    rows = (
        db.session.query(
            Child.gender, Child.is_disabled, Child.is_ovz,
            SchoolClass.id, SchoolClass.name, SchoolClass.grade,
            Building.name,
            ChildSocial.family_status,
            ChildSocial.has_large_family, ChildSocial.has_disability_parents,
            ChildSocial.has_low_income_family, ChildSocial.has_guardianship,
            ChildSocial.has_orphan_status, ChildSocial.is_socially_dangerous,
            ChildSocial.is_hard_life,
            ChildSocial.is_single_mother, ChildSocial.is_single_father,
            ChildSocial.is_repeat_year, ChildSocial.is_svo_family,
            ChildSocial.vshu_since, ChildSocial.kdn_since, ChildSocial.pdn_since,
        )
        .select_from(Child)
        .join(ChildEnrollment, ChildEnrollment.child_id == Child.id)
        .join(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(Building, Building.id == SchoolClass.building_id)
        .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
        .filter(
            ChildEnrollment.academic_year_id == year.id,
            ChildEnrollment.ended_at.is_(None),
        )
        .all()
    )

    by_class, by_grade, by_building = {}, {}, {}
    school_totals = _empty_passport_totals()
    cat_keys = [k for k, _ in SOCIAL_PASSPORT_CATEGORIES]

    for (gender, is_disabled, is_ovz,
         sc_id, sc_name, sc_grade, bld_name,
         family_status,
         large_family, disability_parents, low_income, guardianship,
         orphan, socially_dangerous, hard_life,
         single_mother, single_father, repeat_year, svo_family,
         vshu_since, kdn_since, pdn_since) in rows:

        fs = (family_status or "").lower()
        flags = {
            "large_family": bool(large_family),
            "incomplete": "неполн" in fs,
            "single_mother": bool(single_mother),
            "single_father": bool(single_father),
            "parents_disability": bool(disability_parents),
            "child_disability": bool(is_disabled),
            "guardianship": bool(guardianship),
            "repeat_year": bool(repeat_year),
            "ovz": bool(is_ovz),
            "low_income": bool(low_income),
            "vshu_kdn_pdn": bool(vshu_since or kdn_since or pdn_since),
            "orphan": bool(orphan),
            "svo_family": bool(svo_family),
        }

        def _bump(t):
            t["total"] += 1
            g = (gender or "").lower()
            if g.startswith("м"):
                t["boys"] += 1
            elif g.startswith("ж") or g.startswith("д"):
                t["girls"] += 1
            for k in cat_keys:
                if flags[k]:
                    t[k] += 1

        _bump(school_totals)
        if sc_id not in by_class:
            by_class[sc_id] = {"label": sc_name, "grade": sc_grade,
                               "totals": _empty_passport_totals()}
        _bump(by_class[sc_id]["totals"])
        gk = sc_grade if sc_grade is not None else 0
        if gk not in by_grade:
            by_grade[gk] = {"label": (f"{gk} класс" if gk else "Без параллели"),
                            "totals": _empty_passport_totals()}
        _bump(by_grade[gk]["totals"])
        bk = bld_name or "Без корпуса"
        if bk not in by_building:
            by_building[bk] = {"label": bk, "totals": _empty_passport_totals()}
        _bump(by_building[bk]["totals"])

    wb = Workbook()
    wb.remove(wb.active)

    headers_base = ["", "Всего", "М", "Д"] + [lbl for _, lbl in SOCIAL_PASSPORT_CATEGORIES]

    def _write_sheet(name, label_col, items):
        ws = wb.create_sheet(name)
        ws.append([label_col] + headers_base[1:])
        for it in items:
            t = it["totals"]
            ws.append([it["label"], t["total"], t["boys"], t["girls"]] + [t[k] for k in cat_keys])
        ws.append(["Итого по школе", school_totals["total"],
                   school_totals["boys"], school_totals["girls"]] +
                  [school_totals[k] for k in cat_keys])
        ws.column_dimensions["A"].width = 28

    _write_sheet(
        "По классам", "Класс",
        sorted(by_class.values(), key=lambda r: ((r["grade"] or 99), r["label"])),
    )
    _write_sheet(
        "По параллелям", "Параллель",
        sorted(by_grade.values(), key=lambda r: r["label"]),
    )
    _write_sheet(
        "По корпусам", "Корпус",
        sorted(by_building.values(), key=lambda r: r["label"]),
    )

    ws = wb.create_sheet("Итого по школе")
    ws.append(["Школа"] + headers_base[1:])
    ws.append(["Вся школа", school_totals["total"],
               school_totals["boys"], school_totals["girls"]] +
              [school_totals[k] for k in cat_keys])
    ws.column_dimensions["A"].width = 28

    return _xlsx_response(wb, f"social_passport_summary_{year.name.replace('/', '-')}.xlsx")


# =========================================================
# SUBJECTS
# =========================================================
@children_bp.route("/subjects")
@require_roles("ADMIN")
def subjects_registry():
    return redirect(url_for(
        "workload.catalog",
        section="SUBJECTS",
        q=(request.args.get("q") or "").strip(),
    ))


@children_bp.route("/subjects/new", methods=["POST"])
@require_roles("ADMIN")
def subjects_new():
    flash("Предметы добавляются только через единый реестр.", "info")
    return redirect(url_for("children.subjects_registry"))


@children_bp.route("/subjects/<int:subject_id>/update", methods=["POST"])
@require_roles("ADMIN")
def subjects_update(subject_id: int):
    Subject.query.get_or_404(subject_id)
    flash("Предметы изменяются только через единый реестр.", "info")
    return redirect(url_for("children.subjects_registry"))


@children_bp.route("/subjects/<int:subject_id>/delete", methods=["POST"])
@require_roles("ADMIN")
def subjects_delete(subject_id: int):
    Subject.query.get_or_404(subject_id)
    flash(
        "Предмет нельзя удалить из старого раздела. "
        "Используйте единый реестр.",
        "info",
    )
    return redirect(url_for("children.subjects_registry"))

# =========================================================
# ACADEMIC YEARS
# =========================================================
@children_bp.route("/academic-years")
@require_roles("ADMIN")
def academic_years_registry():
    years = AcademicYear.query.order_by(AcademicYear.start_date.desc().nullslast(), AcademicYear.name.desc()).all()
    return render_template("academic_years_list.html", years=years)


@children_bp.route("/academic-years/new", methods=["POST"])
@require_roles("ADMIN")
def academic_years_new():
    name = (request.form.get("name") or "").strip()
    start_date = parse_date(request.form.get("start_date"))
    end_date = parse_date(request.form.get("end_date"))
    make_current = as_checkbox(request.form, "is_current")
    copy_structure = as_checkbox(request.form, "copy_structure")

    if not name:
        flash("Укажите название учебного года", "danger")
        return redirect(url_for("children.academic_years_registry"))

    exists = AcademicYear.query.filter_by(name=name).first()
    if exists:
        flash("Такой учебный год уже существует", "warning")
        return redirect(url_for("children.academic_years_registry"))

    source_year = _get_current_year() or (
        AcademicYear.query
        .order_by(
            AcademicYear.start_date.desc().nullslast(),
            AcademicYear.name.desc(),
        )
        .first()
    )
    if copy_structure and source_year and (not start_date or not end_date):
        flash(
            "Для копирования учебных планов укажите даты начала и "
            "окончания нового учебного года.",
            "danger",
        )
        return redirect(url_for("children.academic_years_registry"))
    if make_current:
        AcademicYear.query.update({"is_current": False})
    y = AcademicYear(
        name=name,
        start_date=start_date,
        end_date=end_date,
        is_current=make_current,
    )
    try:
        db.session.add(y)
        db.session.flush()
        rollover = None
        if source_year and copy_structure:
            from app.services.academic_year_rollover_service import (
                initialize_academic_year,
            )

            rollover = initialize_academic_year(
                source_year,
                y,
                user_id=current_user.id,
            )
        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception("Academic year creation failed")
        flash(
            "Не удалось создать учебный год и скопировать его структуру.",
            "danger",
        )
        return redirect(url_for("children.academic_years_registry"))

    if rollover:
        flash(
            f"Учебный год добавлен. Скопировано независимых комплектов УП: "
            f"{rollover.plans_created}. Классы, ученики, группы и нагрузка "
            "не копировались.",
            "success",
        )
    else:
        flash("Учебный год добавлен", "success")
    return redirect(url_for("children.academic_years_registry"))


@children_bp.route("/academic-years/create-next", methods=["POST"])
@require_roles("ADMIN")
def academic_year_create_next():
    current = _get_current_year()
    base = current or AcademicYear.query.order_by(AcademicYear.start_date.desc().nullslast(), AcademicYear.name.desc()).first()
    if not base:
        flash("Сначала создайте хотя бы один учебный год вручную", "warning")
        return redirect(url_for("children.academic_years_registry"))

    if base.start_date and base.end_date:
        start_date = base.start_date.replace(year=base.start_date.year + 1)
        end_date = base.end_date.replace(year=base.end_date.year + 1)
        name = f"{start_date.year}/{end_date.year}"
    else:
        parts = (base.name or "").replace('-', '/').split('/')
        if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
            name = f"{int(parts[0])+1}/{int(parts[1])+1}"
            start_date = date(int(parts[0])+1, 9, 1)
            end_date = date(int(parts[1])+1, 8, 31)
        else:
            flash("Не удалось автоматически определить следующий учебный год", "danger")
            return redirect(url_for("children.academic_years_registry"))

    if AcademicYear.query.filter_by(name=name).first():
        flash(f"Учебный год {name} уже существует", "warning")
        return redirect(url_for("children.academic_years_registry"))

    new_year = AcademicYear(
        name=name,
        start_date=start_date,
        end_date=end_date,
        is_current=False,
    )
    try:
        db.session.add(new_year)
        db.session.flush()
        from app.services.academic_year_rollover_service import (
            initialize_academic_year,
        )

        rollover = initialize_academic_year(
            base,
            new_year,
            user_id=current_user.id,
        )
        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception("Academic year rollover failed")
        flash(
            "Не удалось создать следующий учебный год. "
            "Изменения отменены полностью.",
            "danger",
        )
        return redirect(url_for("children.academic_years_registry"))
    flash(
        f"Создан учебный год {name}. Скопировано независимых комплектов УП: "
        f"{rollover.plans_created}. Классы, ученики, группы и нагрузка "
        "не копировались.",
        "success",
    )
    return redirect(url_for("children.academic_years_registry"))


@children_bp.route("/academic-years/<int:year_id>/toggle-closed", methods=["POST"])
@require_roles("ADMIN")
def academic_year_toggle_closed(year_id: int):
    year = AcademicYear.query.get_or_404(year_id)
    year.is_closed = not bool(getattr(year, 'is_closed', False))
    db.session.commit()
    flash("Статус учебного года обновлён", "success")
    return redirect(url_for("children.academic_years_registry"))


@children_bp.route("/academic-years/<int:year_id>/toggle-archive", methods=["POST"])
@require_roles("ADMIN")
def academic_year_toggle_archive(year_id: int):
    year = AcademicYear.query.get_or_404(year_id)
    year.is_archived = not bool(getattr(year, 'is_archived', False))
    db.session.commit()
    flash("Архивный статус учебного года обновлён", "success")
    return redirect(url_for("children.academic_years_registry"))


@children_bp.route("/academic-years/<int:year_id>/make-current", methods=["POST"])
@require_roles("ADMIN")
def academic_year_make_current(year_id: int):
    year = AcademicYear.query.get_or_404(year_id)

    AcademicYear.query.update({"is_current": False})
    year.is_current = True
    db.session.commit()

    flash(f"Текущий учебный год: {year.name}", "success")
    return redirect(url_for("children.academic_years_registry"))


@children_bp.route("/academic-years/<int:year_id>/update", methods=["POST"])
@require_roles("ADMIN")
def academic_year_update(year_id: int):
    year = AcademicYear.query.get_or_404(year_id)

    name = (request.form.get("name") or "").strip()
    start_date = parse_date(request.form.get("start_date"))
    end_date = parse_date(request.form.get("end_date"))

    if not name:
        flash("Название учебного года не может быть пустым", "danger")
        return redirect(url_for("children.academic_years_registry"))

    exists = (
        AcademicYear.query
        .filter(AcademicYear.name == name, AcademicYear.id != year.id)
        .first()
    )
    if exists:
        flash("Учебный год с таким названием уже существует", "warning")
        return redirect(url_for("children.academic_years_registry"))

    year.name = name
    year.start_date = start_date
    year.end_date = end_date

    db.session.commit()
    flash("Учебный год сохранён", "success")
    return redirect(url_for("children.academic_years_registry"))

@children_bp.route("/subjects/import", methods=["GET", "POST"])
@require_roles("ADMIN")
def subjects_import():
    if request.method == "POST":
        f = request.files.get("file")
        if not f or not f.filename:
            flash("Выберите Excel файл", "danger")
            return redirect(url_for("children.subjects_registry"))

        wb = load_workbook(f, data_only=True)
        ws = wb.active

        headers = [(str(cell.value).strip() if cell.value is not None else "") for cell in ws[1]]
        idx = {h: i for i, h in enumerate(headers)}

        if "name" not in idx:
            flash("В файле должна быть колонка: name", "danger")
            return redirect(url_for("children.subjects_registry"))

        created = 0
        skipped = 0

        for r in range(2, ws.max_row + 1):
            row = [ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)]

            name = str(row[idx["name"]] or "").strip()
            short_name = None
            if "short_name" in idx:
                short_name = str(row[idx["short_name"]] or "").strip() or None

            if not name:
                skipped += 1
                continue

            _activity, was_created = get_or_create_subject_activity(
                name,
                short_name=short_name,
                created_by_user_id=current_user.id,
            )
            if not was_created:
                skipped += 1
                continue
            created += 1

        db.session.commit()
        flash(f"Импорт завершён. Добавлено: {created}, пропущено: {skipped}", "success")
        return redirect(url_for("children.subjects_registry"))

    return render_template("subjects_import.html")

@children_bp.route("/children/import-parents", methods=["GET", "POST"])
@require_roles("ADMIN")
def parents_import():
    if request.method == "POST":
        f = request.files.get("file")
        if not f or not f.filename:
            flash("Выберите Excel файл", "danger")
            return redirect(url_for("children.parents_import"))

        wb = load_workbook(f, data_only=True)
        ws = wb.active

        headers = [(str(cell.value).strip() if cell.value is not None else "") for cell in ws[1]]
        idx = {h: i for i, h in enumerate(headers)}

        required = [
            "ФИО",
            "Дата рождения",
            "Тип представителя",
            "ФИО представителя",
            "Телефон представителя",
            "E-mail представителя",
        ]
        missing = [c for c in required if c not in idx]
        if missing:
            flash(f"Не хватает колонок: {', '.join(missing)}", "danger")
            return redirect(url_for("children.parents_import"))

        created_links = 0
        created_parents = 0
        skipped = 0
        not_found = 0

        def parse_birth(x):
            if not x:
                return None
            if isinstance(x, datetime):
                return x.date()
            if isinstance(x, date):
                return x
            s = str(x).strip()
            m = re.match(r"^(\d{2})\.(\d{2})\.(\d{4})$", s)
            if m:
                return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
            try:
                return datetime.strptime(s, "%Y-%m-%d").date()
            except Exception:
                return None

        def split_fio(full_fio: str):
            parts = [p.strip() for p in str(full_fio or "").split() if p.strip()]
            last_name = parts[0] if len(parts) > 0 else None
            first_name = parts[1] if len(parts) > 1 else None
            middle_name = parts[2] if len(parts) > 2 else None
            return last_name, first_name, middle_name

        def normalize_relation(value: str):
            s = (value or "").strip().lower()
            if s == "мать":
                return "mother"
            if s == "отец":
                return "father"
            if s == "опекун":
                return "guardian"
            return "other"

        def split_multi_values(raw: str):
            if not raw:
                return []
            parts = re.split(r"[,\n;]+", str(raw))
            return [p.strip() for p in parts if p and p.strip()]

        for r in range(2, ws.max_row + 1):
            row = [ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)]

            child_fio = str(row[idx["ФИО"]] or "").strip()
            child_birth_date = parse_birth(row[idx["Дата рождения"]])

            relation_raw = str(row[idx["Тип представителя"]] or "").strip()
            parent_fio = str(row[idx["ФИО представителя"]] or "").strip()

            raw_phone = str(row[idx["Телефон представителя"]] or "").strip()
            raw_email = str(row[idx["E-mail представителя"]] or "").strip()

            parent_phone = None
            parent_email = None
            notes_parts = []

            phone_list = split_multi_values(raw_phone)
            if phone_list:
                parent_phone = phone_list[0][:50]
                if len(phone_list) > 1:
                    notes_parts.append("Доп. телефоны: " + ", ".join(phone_list[1:]))

            email_list = split_multi_values(raw_email)
            if email_list:
                parent_email = email_list[0][:120]
                if len(email_list) > 1:
                    notes_parts.append("Доп. e-mail: " + ", ".join(email_list[1:]))

            parent_notes = "\n".join(notes_parts) if notes_parts else None

            if not child_fio or not parent_fio:
                skipped += 1
                continue

            child_last_name, child_first_name, child_middle_name = split_fio(child_fio)

            q = Child.query.filter(
                db.func.lower(Child.last_name) == (child_last_name or "").lower(),
                db.func.lower(Child.first_name) == (child_first_name or "").lower(),
            )

            if child_middle_name:
                q = q.filter(db.func.lower(Child.middle_name) == child_middle_name.lower())

            if child_birth_date:
                q = q.filter(Child.birth_date == child_birth_date)

            child = q.first()

            if not child:
                not_found += 1
                continue

            relation_type = normalize_relation(relation_raw)

            existing_parent = Parent.query.filter(
                db.func.lower(Parent.fio) == parent_fio.lower()
            ).first()

            if existing_parent:
                parent = existing_parent

                if parent_phone and not parent.phone:
                    parent.phone = parent_phone

                if parent_email and not parent.email:
                    parent.email = parent_email

                if parent_notes:
                    old_notes = (parent.notes or "").strip()
                    if old_notes:
                        if parent_notes not in old_notes:
                            parent.notes = old_notes + "\n" + parent_notes
                    else:
                        parent.notes = parent_notes
            else:
                parent = Parent(
                    fio=parent_fio,
                    phone=parent_phone,
                    email=parent_email,
                    notes=parent_notes,
                )
                db.session.add(parent)
                db.session.flush()
                created_parents += 1

            exists_link = ChildParent.query.filter_by(
                child_id=child.id,
                parent_id=parent.id,
                relation_type=relation_type
            ).first()

            if not exists_link:
                link = ChildParent(
                    child_id=child.id,
                    parent_id=parent.id,
                    relation_type=relation_type,
                    is_legal_representative=True,
                )
                db.session.add(link)
                created_links += 1
            else:
                skipped += 1

        db.session.commit()

        flash(
            f"Импорт родителей завершён. "
            f"Создано представителей: {created_parents}, "
            f"создано связей: {created_links}, "
            f"не найдено детей: {not_found}, "
            f"пропущено: {skipped}",
            "success"
        )
        return redirect(url_for("children.list_children"))

    return render_template("parents_import.html")

@children_bp.route("/classes/copy-from-year", methods=["POST"])
@require_roles("ADMIN")
def classes_copy_from_year():
    target_year_id = request.form.get("target_year_id", type=int)
    source_year_id = request.form.get("source_year_id", type=int)
    target_year = AcademicYear.query.get_or_404(target_year_id)
    source_year = AcademicYear.query.get_or_404(source_year_id)
    if (
        source_year.start_date
        and target_year.start_date
        and source_year.start_date >= target_year.start_date
    ):
        flash(
            "Структуру классов можно переносить только из более раннего "
            "учебного года.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=target_year.id,
        ))

    created = 0
    skipped_release = 0
    skipped_existing = 0
    created_ids = []
    source_classes = SchoolClass.query.filter_by(academic_year_id=source_year.id).order_by(SchoolClass.name.asc()).all()
    existing_names = {
        normalize_class_name(name)
        for name, in db.session.query(SchoolClass.name).filter_by(
            academic_year_id=target_year.id,
        )
    }
    for sc in source_classes:
        promoted = promoted_class_identity(sc)
        if promoted is None:
            skipped_release += 1
            continue
        target_name, target_grade, target_letter = promoted
        if target_name in existing_names:
            skipped_existing += 1
            continue
        clone = SchoolClass(
            academic_year_id=target_year.id,
            building_id=sc.building_id,
            name=target_name,
            grade=target_grade,
            letter=target_letter,
            max_students=sc.max_students,
            applications_count=0,
            teacher_user_id=sc.teacher_user_id,
            is_active=True,
        )
        db.session.add(clone)
        db.session.flush()
        created_ids.append(clone.id)
        existing_names.add(target_name)
        created += 1
    db.session.commit()
    if created_ids:
        session["last_class_copy"] = {
            "batch_id": str(uuid4()),
            "source_year_id": source_year.id,
            "source_year_name": source_year.name,
            "target_year_id": target_year.id,
            "class_ids": created_ids,
            "created_count": created,
        }
    else:
        session.pop("last_class_copy", None)
    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    flash(
        f"Создано классов следующей параллели: {created}. "
        f"Не перенесено выпускных и нераспознанных классов: "
        f"{skipped_release}. Уже существовало: {skipped_existing}. "
        "Первые классы создаются вручную.",
        "success",
    )
    return redirect(url_for("children.classes_registry", academic_year_id=target_year.id))


@children_bp.route("/classes/copy-from-year/undo", methods=["POST"])
@require_roles("ADMIN")
def classes_copy_from_year_undo():
    last_copy = session.get("last_class_copy") or {}
    target_year_id = request.form.get("target_year_id", type=int)
    if (
        not target_year_id
        or last_copy.get("target_year_id") != target_year_id
    ):
        flash("Последнее массовое копирование уже недоступно.", "warning")
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=target_year_id,
        ))

    class_ids = [
        int(class_id)
        for class_id in last_copy.get("class_ids", [])
        if str(class_id).isdigit()
    ]
    copied_classes = (
        SchoolClass.query
        .filter(
            SchoolClass.academic_year_id == target_year_id,
            SchoolClass.id.in_(class_ids),
        )
        .all()
        if class_ids else []
    )
    copied_class_ids = [school_class.id for school_class in copied_classes]
    has_children = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.school_class_id.in_(copied_class_ids),
            ChildEnrollment.ended_at.is_(None),
        )
        .first()
        is not None
        if copied_class_ids else False
    )
    if has_children:
        flash(
            "Отмена невозможна: в одном или нескольких созданных классах "
            "уже есть ученики.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=target_year_id,
        ))

    teacher_ids = {
        school_class.teacher_user_id
        for school_class in copied_classes
        if school_class.teacher_user_id
    }
    try:
        for school_class in copied_classes:
            db.session.delete(school_class)
        db.session.flush()
        for teacher_id in teacher_ids:
            _sync_class_teacher_role(teacher_id)
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        flash(
            "Отмена не выполнена: некоторые классы уже используются в "
            "учебных планах, группах или других данных. Сначала удалите "
            "эти связи.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=target_year_id,
        ))

    session.pop("last_class_copy", None)
    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    flash(
        f"Отменено последнее массовое копирование. Удалено классов: "
        f"{len(copied_classes)}.",
        "success",
    )
    return redirect(url_for(
        "children.classes_registry",
        academic_year_id=target_year_id,
    ))


@children_bp.route("/classes/delete-selected", methods=["POST"])
@require_roles("ADMIN")
def classes_delete_selected():
    academic_year_id = request.form.get("academic_year_id", type=int)
    class_ids = {
        int(class_id)
        for class_id in request.form.getlist("class_ids")
        if str(class_id).isdigit()
    }
    if not academic_year_id or not class_ids:
        flash("Выберите хотя бы один класс для удаления.", "warning")
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=academic_year_id,
        ))

    selected_classes = (
        SchoolClass.query
        .filter(
            SchoolClass.academic_year_id == academic_year_id,
            SchoolClass.id.in_(class_ids),
        )
        .order_by(SchoolClass.name.asc())
        .all()
    )
    active_class_ids = {
        class_id
        for class_id, in (
            db.session.query(ChildEnrollment.school_class_id)
            .filter(
                ChildEnrollment.school_class_id.in_(
                    [school_class.id for school_class in selected_classes]
                ),
                ChildEnrollment.ended_at.is_(None),
            )
            .distinct()
            .all()
        )
    }
    blocked_with_children = []
    blocked_with_links = []
    deleted_ids = []
    deleted_names = []
    deleted_records = []
    teacher_ids = set()

    for school_class in selected_classes:
        if school_class.id in active_class_ids:
            blocked_with_children.append(school_class.name)
            continue
        try:
            with db.session.begin_nested():
                teacher_user_id = school_class.teacher_user_id
                class_id = school_class.id
                class_name = school_class.name
                restore_record = {
                    "id": school_class.id,
                    "academic_year_id": school_class.academic_year_id,
                    "building_id": school_class.building_id,
                    "name": school_class.name,
                    "grade": school_class.grade,
                    "letter": school_class.letter,
                    "max_students": school_class.max_students,
                    "applications_count": school_class.applications_count,
                    "teacher_user_id": school_class.teacher_user_id,
                    "is_active": school_class.is_active,
                    "is_archived": school_class.is_archived,
                    "created_at": school_class.created_at,
                }
                db.session.delete(school_class)
                db.session.flush()
            deleted_ids.append(class_id)
            deleted_names.append(class_name)
            deleted_records.append(restore_record)
            if teacher_user_id:
                teacher_ids.add(teacher_user_id)
        except IntegrityError:
            blocked_with_links.append(school_class.name)

    for teacher_id in teacher_ids:
        _sync_class_teacher_role(teacher_id)
    db.session.commit()
    if deleted_records:
        delete_token = str(uuid4())
        cache.set(
            make_key("deleted_classes", delete_token),
            deleted_records,
            timeout=900,
        )
        session["last_class_delete"] = {
            "token": delete_token,
            "academic_year_id": academic_year_id,
            "deleted_count": len(deleted_records),
        }

    last_copy = session.get("last_class_copy") or {}
    if last_copy.get("target_year_id") == academic_year_id:
        remaining_ids = [
            class_id
            for class_id in last_copy.get("class_ids", [])
            if class_id not in deleted_ids
        ]
        if remaining_ids:
            last_copy["class_ids"] = remaining_ids
            last_copy["created_count"] = len(remaining_ids)
            session["last_class_copy"] = last_copy
        else:
            session.pop("last_class_copy", None)

    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    if deleted_names:
        flash(
            f"Удалено классов: {len(deleted_names)}.",
            "success",
        )
    if blocked_with_children:
        flash(
            "Не удалены классы с учениками: "
            + ", ".join(blocked_with_children)
            + ".",
            "warning",
        )
    if blocked_with_links:
        flash(
            "Не удалены классы, используемые в учебных планах, группах "
            "или других данных: "
            + ", ".join(blocked_with_links)
            + ".",
            "warning",
        )

    return redirect(url_for(
        "children.classes_registry",
        academic_year_id=academic_year_id,
    ))


@children_bp.route("/classes/delete/undo", methods=["POST"])
@require_roles("ADMIN")
def classes_delete_undo():
    academic_year_id = request.form.get("academic_year_id", type=int)
    last_delete = session.get("last_class_delete") or {}
    delete_token = last_delete.get("token")
    if (
        not academic_year_id
        or last_delete.get("academic_year_id") != academic_year_id
        or not delete_token
    ):
        flash("Отмена удаления уже недоступна.", "warning")
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=academic_year_id,
        ))

    cache_key = make_key("deleted_classes", delete_token)
    deleted_records = cache.get(cache_key) or []
    if not deleted_records:
        session.pop("last_class_delete", None)
        flash("Срок отмены удаления истёк.", "warning")
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=academic_year_id,
        ))

    restored_names = []
    skipped_names = []
    teacher_ids = set()
    for record in deleted_records:
        record_name = record["name"]
        id_exists = db.session.get(SchoolClass, record["id"]) is not None
        if id_exists or _class_name_exists(academic_year_id, record_name):
            skipped_names.append(record_name)
            continue
        restored = SchoolClass(**record)
        db.session.add(restored)
        restored_names.append(record_name)
        if restored.teacher_user_id:
            teacher_ids.add(restored.teacher_user_id)

    try:
        db.session.flush()
        for teacher_id in teacher_ids:
            _sync_class_teacher_role(teacher_id)
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        flash(
            "Не удалось отменить удаление: один из классов уже используется "
            "или его название занято.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=academic_year_id,
        ))

    cache.delete(cache_key)
    session.pop("last_class_delete", None)
    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    if restored_names:
        flash(
            f"Восстановлено классов: {len(restored_names)}.",
            "success",
        )
    if skipped_names:
        flash(
            "Не восстановлены классы с уже занятыми названиями: "
            + ", ".join(skipped_names)
            + ".",
            "warning",
        )
    return redirect(url_for(
        "children.classes_registry",
        academic_year_id=academic_year_id,
    ))


@children_bp.route("/classes/<int:class_id>/delete", methods=["POST"])
@require_roles("ADMIN")
def classes_delete(class_id: int):
    c = SchoolClass.query.get_or_404(class_id)
    academic_year_id = c.academic_year_id

    teacher_user_id = c.teacher_user_id

    has_children = (
        ChildEnrollment.query
        .filter(
            ChildEnrollment.school_class_id == c.id,
            ChildEnrollment.ended_at.is_(None)
        )
        .first()
        is not None
    )

    if has_children:
        flash("Нельзя удалить класс: в нём есть активные дети", "danger")
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=academic_year_id,
        ))

    restore_record = {
        "id": c.id,
        "academic_year_id": c.academic_year_id,
        "building_id": c.building_id,
        "name": c.name,
        "grade": c.grade,
        "letter": c.letter,
        "max_students": c.max_students,
        "applications_count": c.applications_count,
        "teacher_user_id": c.teacher_user_id,
        "is_active": c.is_active,
        "is_archived": c.is_archived,
        "created_at": c.created_at,
    }

    try:
        db.session.delete(c)
        db.session.flush()
        _sync_class_teacher_role(teacher_user_id)
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        flash(
            "Нельзя удалить класс: он используется в учебных планах, "
            "группах или других данных.",
            "danger",
        )
        return redirect(url_for(
            "children.classes_registry",
            academic_year_id=academic_year_id,
        ))

    delete_token = str(uuid4())
    cache.set(
        make_key("deleted_classes", delete_token),
        [restore_record],
        timeout=900,
    )
    session["last_class_delete"] = {
        "token": delete_token,
        "academic_year_id": academic_year_id,
        "deleted_count": 1,
    }
    view_response_cache.delete_prefix("classes_registry")
    view_response_cache.delete_prefix("social_passport_registry")
    flash("Класс удалён", "success")
    return redirect(url_for(
        "children.classes_registry",
        academic_year_id=academic_year_id,
    ))


@children_bp.route(
    "/classes/<int:class_id>/applications",
    methods=["POST"],
)
@require_roles(
    "ADMIN",
    "DIRECTOR",
    "DEPUTY_DIRECTOR",
    "SECRETARY",
    "SECRETARY_ACADEMIC",
)
def class_applications_update(class_id: int):
    school_class = SchoolClass.query.get_or_404(class_id)
    applications_count = request.form.get("applications_count", type=int)
    if applications_count is None or applications_count < 0:
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify({
                "ok": False,
                "error": (
                    "Количество заявлений должно быть целым "
                    "неотрицательным числом."
                ),
            }), 400
        flash(
            "Количество заявлений должно быть целым неотрицательным числом.",
            "danger",
        )
    else:
        school_class.applications_count = applications_count
        db.session.commit()
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify({
                "ok": True,
                "class_id": school_class.id,
                "applications_count": school_class.applications_count,
            })
        flash(
            f"Количество заявлений для {school_class.name} сохранено.",
            "success",
        )

    return redirect(url_for(
        "children.contingent",
        year_id=school_class.academic_year_id,
        building_id=request.form.get("building_id", type=int),
    ))

@children_bp.route("/registry/kdn")
@login_required
def registry_kdn():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    q = (
        Child.query
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == year.id)
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
        .filter(ChildSocial.kdn_since.isnot(None))
    )

    if filters["selected_grade"] is not None:
        q = q.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        q = q.filter(SchoolClass.id == filters["selected_class_id"])

    children = q.order_by(SchoolClass.name.asc(), Child.last_name.asc(), Child.first_name.asc()).all()

    if filters["q_text"]:
        children = [ch for ch in children if _match_fio_query(ch, filters["q_text"])]

    return render_template(
        "registry_children.html",
        title="Реестр КДН",
        children=children,
        q_text=filters["q_text"],
        classes=filters["classes"],
        grades=filters["grades"],
        selected_grade=filters["selected_grade"],
        selected_class_id=filters["selected_class_id"],
        export_url=url_for("children.registry_kdn_export", grade=filters["selected_grade_raw"], class_id=filters["selected_class_id"], q=filters["q_text"])
    )

@children_bp.route("/registry/kdn/export")
@login_required
def registry_kdn_export():
    year = _get_current_year()
    filters = _registry_filter_state(year, allow_only_own_class=should_limit_children_to_own_class())
    if not year:
        flash("Не найден текущий учебный год", "danger")
        return redirect(url_for("children.home"))

    q = (
        Child.query
        .outerjoin(
            ChildEnrollment,
            (ChildEnrollment.child_id == Child.id)
            & (ChildEnrollment.academic_year_id == year.id)
            & (ChildEnrollment.ended_at.is_(None))
        )
        .outerjoin(SchoolClass, SchoolClass.id == ChildEnrollment.school_class_id)
        .outerjoin(ChildSocial, ChildSocial.child_id == Child.id)
        .filter(ChildSocial.kdn_since.isnot(None))
    )

    if filters["selected_grade"] is not None:
        q = q.filter(SchoolClass.grade == filters["selected_grade"])
    if filters["selected_class_id"]:
        q = q.filter(SchoolClass.id == filters["selected_class_id"])

    children = q.order_by(Child.last_name.asc(), Child.first_name.asc()).all()

    if filters["q_text"]:
        children = [ch for ch in children if _match_fio_query(ch, filters["q_text"])]

    return _export_children_xlsx("Реестр_КДН", children)

@children_bp.route("/admin/roles", methods=["GET", "POST"])
@require_roles("ADMIN")
def roles_admin():
    if request.method == "POST":
        user_id = request.form.get("user_id", type=int)
        selected_role_codes = request.form.getlist("roles")

        user = User.query.get_or_404(user_id)

        editable_roles = Role.query.filter(Role.code != "CLASS_TEACHER").all()
        editable_role_ids = [r.id for r in editable_roles]

        UserRole.query.filter(
            UserRole.user_id == user.id,
            UserRole.role_id.in_(editable_role_ids)
        ).delete(synchronize_session=False)

        for code in selected_role_codes:
            if code == "CLASS_TEACHER":
                continue

            role = Role.query.filter_by(code=code).first()
            if not role:
                role = Role(code=code, name=code)
                db.session.add(role)
                db.session.flush()
            db.session.add(UserRole(user_id=user.id, role_id=role.id))

        # Keep legacy user.role in sync so old single-role checks still work
        _ROLE_PRIORITY = ["ADMIN", "DEPUTY_DIRECTOR", "METHODIST", "PSYCHOLOGIST",
                          "SOCIAL_PEDAGOG", "LOGOPEDIST",
                          "DEFECTOLOGIST", "OLIGOPHRENOPEDAGOG", "CLASS_TEACHER",
                          "TEACHER", "EDUCATOR", "SENIOR_EDUCATOR", "SPECIALIST",
                          "TUTOR", "KPP", "VIEWER"]
        non_ct = [c for c in selected_role_codes if c != "CLASS_TEACHER"]
        if non_ct:
            user.role = min(non_ct, key=lambda c: _ROLE_PRIORITY.index(c) if c in _ROLE_PRIORITY else 99)
        elif "CLASS_TEACHER" in selected_role_codes:
            user.role = "CLASS_TEACHER"
        else:
            user.role = "VIEWER"

        db.session.commit()
        flash("Роли пользователя сохранены", "success")
        return redirect(url_for("children.roles_admin", q=request.args.get("q", "")))

    q = (request.args.get("q") or "").strip().lower()

    users = User.query.order_by(
        User.last_name.asc(),
        User.first_name.asc(),
        User.middle_name.asc()
    ).all()

    if q:
        def match_user(u):
            text = " ".join([
                u.fio or "",
                u.username or "",
                u.phone or "",
                u.email or "",
            ]).lower()
            return q in text

        users = [u for u in users if match_user(u)]

    roles = Role.query.order_by(Role.name.asc()).all()

    rows = []
    for u in users:
        rows.append({
            "user": u,
            "role_codes": set(u.role_codes),
        })

    return render_template(
        "roles_admin.html",
        rows=rows,
        roles=roles,
        q=q,
    )

