import hashlib
import os
import re
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from uuid import uuid4

from flask import current_app
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sqlalchemy import func

from app.core.extensions import db
from app.models import TariffDocumentArtifact
from app.services.tariff_calculation_service import latest_successful_run
from app.services.tariff_workflow_service import (
    TariffWorkflowError,
    latest_review_cycle,
    latest_validation_run,
)


TEMPLATE_VERSION = "ALT-DOCX-1.0"
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)
DOCUMENT_LABELS = {
    "SUMMARY_TARIFF": "Сводный тарификационный список",
    "PERSONAL_TARIFF": "Персональный тарификационный лист",
    "REVIEW_PROTOCOL": "Протокол проверки и согласования",
    "APPROVAL_SHEET": "Лист утверждения",
    "ORDER_DRAFT": "Проект приказа об установлении нагрузки",
    "CHANGE_ORDER_DRAFT": "Проект приказа об изменении нагрузки",
}
FINANCIAL_DOCUMENT_TYPES = {"SUMMARY_TARIFF", "PERSONAL_TARIFF"}
OFFICIAL_VERSION_STATUSES = {
    "APPROVED",
    "EFFECTIVE",
    "SUPERSEDED",
    "ARCHIVED",
}


def _safe_part(value):
    text = re.sub(r"[^A-Za-zА-Яа-яЁё0-9._-]+", "_", str(value or ""))
    return text.strip("._")[:80] or "document"


def _money(value):
    return f"{Decimal(value or 0):,.2f}".replace(",", " ").replace(".", ",")


def _number(value):
    normalized = Decimal(value or 0)
    return format(normalized.normalize(), "f").replace(".", ",")


def _person(user):
    if user is None:
        return "Вакансия"
    return user.fio or user.username or f"Сотрудник {user.id}"


def _organization_name(version):
    organization = version.tariff_cycle.organization
    return (
        organization.display_name
        if organization is not None
        else "Образовательная организация"
    )


def _set_document_style(document):
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.2)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = "Arial"
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.size = Pt(11)


def _add_header(document, version, title, status):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_organization_name(version))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"{version.tariff_cycle.name}. Версия № {version.version_no}. "
        f"Статус документа: {status}."
    )
    if status == "ПРОЕКТ":
        project = document.add_paragraph()
        project.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project.add_run("ПРОЕКТ. НЕ ЯВЛЯЕТСЯ УТВЕРЖДЁННЫМ ДОКУМЕНТОМ")
        run.bold = True


def _table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(value)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(8)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value if value is not None else "")
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    return table


def _tariff_rows(run, employee_user_id=None, department_id=None):
    lines = sorted(
        run.lines,
        key=lambda row: (
            _person(row.employee),
            row.education_activity.name,
            row.teaching_group.name if row.teaching_group else "",
        ),
    )
    if employee_user_id is not None:
        lines = [
            row for row in lines if row.employee_user_id == employee_user_id
        ]
    if department_id is not None:
        lines = [row for row in lines if row.department_id == department_id]
    return lines


def _build_summary(
    document,
    run,
    employee_user_id=None,
    department_id=None,
    *,
    include_amount=True,
):
    lines = _tariff_rows(run, employee_user_id, department_id)
    rows = []
    for line in lines:
        values = [
            _person(line.employee),
            line.education_activity.name,
            line.teaching_group.name if line.teaching_group else "Без группы",
            _number(line.weekly_hours),
            _number(line.annual_hours),
            _number(line.fte_value),
        ]
        if include_amount:
            values.append(_money(line.total_amount))
        rows.append(values)
    headers = [
        "Педагог",
        "Дисциплина",
        "Класс / группа",
        "Часов в неделю",
        "Часов за период",
        "Ставок",
    ]
    if include_amount:
        headers.append("Сумма, руб.")
    _table(
        document,
        headers,
        rows,
    )
    totals_text = (
        "Итого: "
        f"{_number(sum((row.weekly_hours for row in lines), Decimal('0')))} "
        "ч. в неделю; "
        f"{_number(sum((row.fte_value for row in lines), Decimal('0')))} "
        "ставок"
    )
    if include_amount:
        totals_text += (
            "; "
            f"{_money(sum((row.total_amount for row in lines), Decimal('0')))} "
            "руб."
        )
    document.add_paragraph(totals_text + ".")


def _build_review_protocol(document, version):
    validation = latest_validation_run(version.id)
    document.add_heading("Автоматическая проверка", level=1)
    if validation is None:
        document.add_paragraph("Проверка не выполнялась.")
    else:
        counts = validation.summary_data.get("counts", {})
        document.add_paragraph(
            f"Запуск № {validation.run_no}: {validation.status}. "
            f"Блокирующих: {counts.get('BLOCKER', 0)}, "
            f"ошибок: {counts.get('ERROR', 0)}, "
            f"предупреждений: {counts.get('WARNING', 0)}."
        )
        _table(
            document,
            ("Код", "Уровень", "Сообщение", "Как исправить"),
            (
                (
                    item.rule_code,
                    item.severity,
                    item.message,
                    item.remediation or "",
                )
                for item in validation.issues
            ),
        )

    cycle = latest_review_cycle(version.id)
    document.add_heading("Заключения", level=1)
    if cycle is None:
        document.add_paragraph("Цикл согласования не запускался.")
        return
    _table(
        document,
        ("Этап", "Решение", "Проверяющий", "Дата", "Комментарий"),
        (
            (
                item.review_stage,
                item.decision,
                _person(item.decided_by),
                item.decided_at.strftime("%d.%m.%Y %H:%M"),
                item.comment or "",
            )
            for item in cycle.decisions
        ),
    )
    document.add_heading("Замечания", level=1)
    _table(
        document,
        ("Этап", "Вид", "Статус", "Замечание", "Ответ"),
        (
            (
                item.review_stage,
                item.comment_kind,
                item.status,
                item.text,
                item.response_text or "",
            )
            for item in cycle.comments
        ),
    )


def _build_approval_sheet(document, version):
    document.add_heading("Реквизиты версии", level=1)
    _table(
        document,
        ("Показатель", "Значение"),
        (
            ("Тип версии", version.version_type),
            ("Статус", version.status),
            ("Начало действия", version.effective_from or ""),
            ("Окончание действия", version.effective_to or ""),
            ("Основание", version.reason_text or ""),
            ("Контрольная сумма", version.checksum or "Не утверждена"),
        ),
    )
    cycle = latest_review_cycle(version.id)
    document.add_heading("Согласование", level=1)
    _table(
        document,
        ("Этап", "Решение", "Проверяющий", "Дата"),
        (
            (
                item.review_stage,
                item.decision,
                _person(item.decided_by),
                item.decided_at.strftime("%d.%m.%Y %H:%M"),
            )
            for item in (cycle.decisions if cycle else [])
        ),
    )
    document.add_heading("Утверждение", level=1)
    _table(
        document,
        ("Решение", "Руководитель", "Дата", "Комментарий"),
        (
            (
                item.decision,
                _person(item.decided_by),
                item.decided_at.strftime("%d.%m.%Y %H:%M"),
                item.comment or "",
            )
            for item in version.approval_decisions
        ),
    )


def _build_order(document, version, run, change=False):
    heading = "Об изменении педагогической нагрузки" if change else (
        "Об установлении педагогической нагрузки"
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = paragraph.add_run(heading)
    run_title.bold = True
    document.add_paragraph(
        "Номер и дата приказа присваиваются в модуле «Приказы» после "
        "проверки и регистрации документа."
    )
    document.add_paragraph(
        "Основание: "
        + (
            version.reason_text
            or f"утверждённая тарификация на {version.tariff_cycle.name}"
        )
        + "."
    )
    document.add_paragraph(
        f"Установить нагрузку с {version.effective_from:%d.%m.%Y} "
        "согласно приложению."
    )
    _build_summary(document, run, include_amount=False)


def _build_document(
    version,
    document_type,
    run,
    employee_user_id=None,
    department_id=None,
):
    document = Document()
    _set_document_style(document)
    if document_type in {"SUMMARY_TARIFF", "PERSONAL_TARIFF"}:
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
    status_label = (
        "ОФИЦИАЛЬНЫЙ"
        if version.status in OFFICIAL_VERSION_STATUSES
        else "ПРОЕКТ"
    )
    _add_header(
        document,
        version,
        DOCUMENT_LABELS[document_type],
        status_label,
    )
    if document_type == "SUMMARY_TARIFF":
        _build_summary(document, run, department_id=department_id)
    elif document_type == "PERSONAL_TARIFF":
        _build_summary(document, run, employee_user_id=employee_user_id)
    elif document_type == "REVIEW_PROTOCOL":
        _build_review_protocol(document, version)
    elif document_type == "APPROVAL_SHEET":
        _build_approval_sheet(document, version)
    elif document_type == "ORDER_DRAFT":
        _build_order(document, version, run)
    elif document_type == "CHANGE_ORDER_DRAFT":
        _build_order(document, version, run, change=True)
    return document


def generate_tariff_document(
    version,
    *,
    document_type,
    user_id,
    employee_user_id=None,
    department_id=None,
):
    if document_type not in DOCUMENT_LABELS:
        raise TariffWorkflowError("Неизвестный тип документа.")
    if document_type == "PERSONAL_TARIFF" and not employee_user_id:
        raise TariffWorkflowError("Выберите сотрудника.")
    if document_type != "PERSONAL_TARIFF":
        employee_user_id = None
    scope_key = (
        f"EMPLOYEE:{employee_user_id}"
        if employee_user_id is not None
        else (
            f"DEPARTMENT:{department_id}"
            if department_id is not None
            else "ALL"
        )
    )
    if version.status in OFFICIAL_VERSION_STATUSES and not version.checksum:
        raise TariffWorkflowError(
            "У утверждённой версии отсутствует контрольная сумма."
        )
    run = latest_successful_run(version.id)
    if run is None:
        raise TariffWorkflowError(
            "Для формирования документа требуется успешный расчёт."
        )
    revision_no = (
        db.session.query(func.max(TariffDocumentArtifact.revision_no))
        .filter(
            TariffDocumentArtifact.tariff_version_id == version.id,
            TariffDocumentArtifact.document_type == document_type,
            TariffDocumentArtifact.scope_key == scope_key,
        )
        .scalar()
        or 0
    ) + 1
    document = _build_document(
        version,
        document_type,
        run,
        employee_user_id=employee_user_id,
        department_id=department_id,
    )
    root = Path(current_app.config["UPLOAD_FOLDER"]).resolve()
    folder = root / "workload_documents" / str(version.id)
    folder.mkdir(parents=True, exist_ok=True)
    filename = (
        f"{_safe_part(document_type.lower())}_v{version.version_no}"
        f"_r{revision_no}.docx"
    )
    stored_name = f"{uuid4().hex}_{filename}"
    full_path = folder / stored_name
    document.save(full_path)
    payload = full_path.read_bytes()
    checksum = hashlib.sha256(payload).hexdigest()
    relative_path = full_path.relative_to(root)
    artifact = TariffDocumentArtifact(
        tariff_version_id=version.id,
        calculation_run_id=run.id,
        document_type=document_type,
        status=(
            "OFFICIAL"
            if version.status in OFFICIAL_VERSION_STATUSES
            else "PROJECT"
        ),
        revision_no=revision_no,
        employee_user_id=employee_user_id,
        department_id=department_id,
        scope_key=scope_key,
        filename=filename,
        storage_path=str(relative_path),
        mime_type=DOCX_MIME,
        file_size=len(payload),
        checksum_sha256=checksum,
        version_checksum=version.checksum,
        template_version=TEMPLATE_VERSION,
        generation_parameters={
            "version_status": version.status,
            "generated_at": datetime.utcnow().isoformat(),
        },
        created_by_user_id=user_id,
    )
    db.session.add(artifact)
    return artifact


def resolve_artifact_path(artifact):
    root = Path(current_app.config["UPLOAD_FOLDER"]).resolve()
    path = (root / artifact.storage_path).resolve()
    if root not in path.parents:
        raise TariffWorkflowError("Недопустимый путь к документу.")
    if not path.is_file():
        raise TariffWorkflowError("Файл документа не найден.")
    checksum = hashlib.sha256(path.read_bytes()).hexdigest()
    if not os.path.isfile(path) or checksum != artifact.checksum_sha256:
        raise TariffWorkflowError(
            "Контрольная сумма документа не совпадает."
        )
    return path


__all__ = [
    "DOCUMENT_LABELS",
    "FINANCIAL_DOCUMENT_TYPES",
    "TEMPLATE_VERSION",
    "generate_tariff_document",
    "resolve_artifact_path",
]
