from datetime import datetime
from io import BytesIO
from collections import defaultdict
import json
from pathlib import Path

from flask import Blueprint, abort, flash, redirect, render_template, request, url_for, send_file
from flask_login import current_user, login_required
from openpyxl import load_workbook
from openpyxl import Workbook as XWorkbook
from sqlalchemy import or_
from sqlalchemy.orm import joinedload

from app.core.extensions import db
from app.services.org_settings_service import get_active_organization_settings, get_organization_signature_block
from app.models import AcademicYear, Building, Child, SchoolClass, User
from app.models.iom import (
    IOM_STATUS_CHOICES,
    IOM_TYPE_CHOICES,
    IomCard,
    IomCyclegramLink,
    IomHistory,
    IomImportSessionSchedule,
    IomMonitoringEntry,
    IomMonitoringTemplate,
    IomExportLog,
    IomScheduleCorrection,
    IomScheduleLesson,
    IomSectionData,
    IomSpecialistPlan,
)
from app.models.service_staff import (
    ServiceActivityType,
    ServiceAssignment,
    ServiceCyclegram,
    ServiceCyclegramEntry,
    ServiceResponsible,
    ServiceSpecialist,
)
from app.permissions import has_any_role, has_role, is_admin, is_class_teacher_of_child

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None
    canvas = None
    simpleSplit = None
    pdfmetrics = None
    TTFont = None


iom_bp = Blueprint("iom", __name__, url_prefix="/iom")


EDUCATION_LEVEL_CHOICES = [
    ("DO", "Дошкольное образование"),
    ("NOO", "НОО"),
    ("OOO", "ООО"),
    ("SOO", "СОО"),
    ("UO", "УО"),
]
EDUCATION_LEVEL_LABELS = dict(EDUCATION_LEVEL_CHOICES)


AOP_VARIANT_CHOICES = {
    "DO": [
        ("DO_TNR", "АООП ДО для обучающихся с ТНР"),
        ("DO_ZPR", "АООП ДО для обучающихся с ЗПР"),
        ("DO_RAS", "АООП ДО для обучающихся с РАС"),
        ("DO_NODA", "АООП ДО для обучающихся с НОДА"),
        ("DO_HEARING", "АООП ДО для обучающихся с нарушением слуха"),
        ("DO_VISION", "АООП ДО для обучающихся с нарушением зрения"),
        ("DO_INTELLECT", "АООП ДО для обучающихся с интеллектуальными нарушениями"),
        ("DO_TMNR", "АООП ДО для обучающихся с ТМНР"),
    ],
    "NOO": [
        ("NOO_1_1", "АООП НОО вариант 1.1"),
        ("NOO_1_2", "АООП НОО вариант 1.2"),
        ("NOO_1_3", "АООП НОО вариант 1.3"),
        ("NOO_1_4", "АООП НОО вариант 1.4"),
        ("NOO_2_1", "АООП НОО вариант 2.1"),
        ("NOO_2_2", "АООП НОО вариант 2.2"),
        ("NOO_2_3", "АООП НОО вариант 2.3"),
        ("NOO_5_1", "АООП НОО вариант 5.1"),
        ("NOO_5_2", "АООП НОО вариант 5.2"),
        ("NOO_6_1", "АООП НОО вариант 6.1"),
        ("NOO_6_2", "АООП НОО вариант 6.2"),
        ("NOO_6_3", "АООП НОО вариант 6.3"),
        ("NOO_6_4", "АООП НОО вариант 6.4"),
        ("NOO_7_1", "АООП НОО вариант 7.1"),
        ("NOO_7_2", "АООП НОО вариант 7.2"),
        ("NOO_8_1", "АООП НОО вариант 8.1"),
        ("NOO_8_2", "АООП НОО вариант 8.2"),
        ("NOO_8_3", "АООП НОО вариант 8.3"),
        ("NOO_8_4", "АООП НОО вариант 8.4"),
    ],
    "OOO": [
        ("OOO_1", "АООП ООО вариант 1"),
        ("OOO_2", "АООП ООО вариант 2"),
        ("OOO_RAS", "АООП ООО для обучающихся с РАС"),
        ("OOO_NODA", "АООП ООО для обучающихся с НОДА"),
        ("OOO_HEARING", "АООП ООО для обучающихся с нарушением слуха"),
        ("OOO_VISION", "АООП ООО для обучающихся с нарушением зрения"),
        ("OOO_TNR", "АООП ООО для обучающихся с ТНР"),
    ],
    "SOO": [
        ("SOO_1", "АООП СОО вариант 1"),
        ("SOO_2", "АООП СОО вариант 2"),
        ("SOO_RAS", "АООП СОО для обучающихся с РАС"),
        ("SOO_NODA", "АООП СОО для обучающихся с НОДА"),
        ("SOO_HEARING", "АООП СОО для обучающихся с нарушением слуха"),
        ("SOO_VISION", "АООП СОО для обучающихся с нарушением зрения"),
        ("SOO_TNR", "АООП СОО для обучающихся с ТНР"),
    ],
    "UO": [
        ("UO_1", "АООП для обучающихся с умственной отсталостью, вариант 1"),
        ("UO_2", "АООП для обучающихся с умственной отсталостью, вариант 2"),
    ],
}
AOP_VARIANT_LABELS = {value: label for items in AOP_VARIANT_CHOICES.values() for value, label in items}

WEEKDAY_ORDER = ["monday", "tuesday", "wednesday", "thursday", "friday"]
WEEKDAY_LABELS = {
    "monday": "Понедельник",
    "tuesday": "Вторник",
    "wednesday": "Среда",
    "thursday": "Четверг",
    "friday": "Пятница",
    "saturday": "Суббота",
}
WEEKDAY_RU_TO_KEY = {
    "понедельник": "monday",
    "вторник": "tuesday",
    "среда": "wednesday",
    "четверг": "thursday",
    "пятница": "friday",
    "суббота": "saturday",
}



def _find_pdf_font_candidates():
    candidates = [
        '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        '/System/Library/Fonts/Supplemental/Arial Unicode MS.ttf',
        '/System/Library/Fonts/Supplemental/Helvetica.ttc',
        '/System/Library/Fonts/Supplemental/Times New Roman.ttf',
        '/System/Library/Fonts/Supplemental/Arial.ttf',
        '/Library/Fonts/Arial Unicode.ttf',
        '/Library/Fonts/Arial Unicode MS.ttf',
        '/Library/Fonts/Arial.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        'C:/Windows/Fonts/arial.ttf',
        'C:/Windows/Fonts/arialbd.ttf',
        'C:/Windows/Fonts/times.ttf',
        'C:/Windows/Fonts/timesbd.ttf',
    ]
    regular = None
    bold = None
    for path in candidates:
        if regular is None and Path(path).exists() and ('Bold' not in path and 'bd' not in path.lower()):
            regular = path
        if bold is None and Path(path).exists() and ('Bold' in path or 'bd' in path.lower()):
            bold = path
    if regular and not bold:
        bold = regular
    return regular, bold


def _ensure_pdf_fonts():
    if pdfmetrics is None or TTFont is None:
        return 'Helvetica', 'Helvetica-Bold'
    try:
        pdfmetrics.getFont('IOMSans')
        return 'IOMSans', 'IOMSans-Bold'
    except Exception:
        pass
    regular, bold = _find_pdf_font_candidates()
    if not regular:
        return 'Helvetica', 'Helvetica-Bold'
    pdfmetrics.registerFont(TTFont('IOMSans', regular))
    pdfmetrics.registerFont(TTFont('IOMSans-Bold', bold or regular))
    return 'IOMSans', 'IOMSans-Bold'


def _pdf_wrap_lines(text, font_name, font_size, max_width):
    if simpleSplit is None:
        return [str(text or '')]
    lines = []
    for raw in str(text or '').splitlines() or ['']:
        wrapped = simpleSplit(raw, font_name, font_size, max_width)
        lines.extend(wrapped or [''])
    return lines or ['']

ROLE_PRIORITY = [
    "ADMIN",
    "METHODIST",
    "SOCIAL_PEDAGOG",
    "SOCIAL_PEDAGOGUE",
    "PSYCHOLOGIST",
    "LOGOPEDIST",
    "DEFECTOLOGIST",
    "TUTOR",
    "ASSISTANT",
    "CLASS_TEACHER",
]



MONITORING_PERIOD_CHOICES = [("INTERIM", "Промежуточный мониторинг"), ("FINAL", "Итоговый мониторинг")]
MONITORING_PERIOD_LABELS = dict(MONITORING_PERIOD_CHOICES)
MONITORING_SCALE_OPTIONS = {
    "SCORE_0_3": ["", "0", "1", "2", "3"],
    "DYNAMICS": ["", "положительная", "фрагментарная", "без динамики"],
    "TEXT": [],
}

DEFAULT_MONITORING_TEMPLATES = {
    "NOO": [
        ("INTERIM", "results", "Личностные результаты", "personal", "Личностные результаты", "SCORE_0_3"),
        ("INTERIM", "life", "Жизненные компетенции", "life", "Жизненные компетенции", "SCORE_0_3"),
        ("INTERIM", "specialists", "Динамика по направлениям работы специалистов", "dynamic", "Динамика по направлениям", "DYNAMICS"),
        ("FINAL", "results", "Личностные результаты", "personal", "Личностные результаты", "SCORE_0_3"),
        ("FINAL", "life", "Жизненные компетенции", "life", "Жизненные компетенции", "SCORE_0_3"),
        ("FINAL", "specialists", "Результаты коррекционно-развивающей работы", "dynamic", "Результаты коррекционно-развивающей работы", "DYNAMICS"),
    ],
    "OOO_SOO": [
        ("INTERIM", "results", "Личностные результаты", "personal", "Личностные результаты", "SCORE_0_3"),
        ("INTERIM", "specialists", "Динамика по направлениям работы специалистов", "dynamic", "Динамика по направлениям", "DYNAMICS"),
        ("FINAL", "results", "Личностные результаты", "personal", "Личностные результаты", "SCORE_0_3"),
        ("FINAL", "specialists", "Результаты коррекционно-развивающей работы", "dynamic", "Результаты коррекционно-развивающей работы", "DYNAMICS"),
    ],
    "UO": [
        ("INTERIM", "results", "Личностные результаты", "personal", "Личностные результаты", "SCORE_0_3"),
        ("INTERIM", "life", "Жизненные компетенции", "life", "Жизненные компетенции", "SCORE_0_3"),
        ("FINAL", "results", "Личностные результаты", "personal", "Личностные результаты", "SCORE_0_3"),
        ("FINAL", "life", "Жизненные компетенции", "life", "Жизненные компетенции", "SCORE_0_3"),
    ],
    "DO": [
        ("INTERIM", "preschool", "Мониторинг дошкольных направлений", "preschool", "Развитие по направлениям ДО", "DYNAMICS"),
        ("FINAL", "preschool", "Мониторинг дошкольных направлений", "preschool", "Развитие по направлениям ДО", "DYNAMICS"),
    ],
}

def _fio(user):
    if not user:
        return ""
    return getattr(user, "fio", None) or " ".join(
        x.strip() for x in [user.last_name or "", user.first_name or "", user.middle_name or ""] if x and x.strip()
    )


def _parse_date(value):
    value = (value or "").strip()
    if not value:
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        return None


def _parse_int(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _normalize_time_text(value):
    value = (value or "").strip()
    if not value:
        return ""
    value = value.replace('.', ':')
    parts = [x for x in value.split(':') if x != ""]
    if len(parts) >= 2 and all(part.isdigit() for part in parts[:2]):
        return f"{int(parts[0]):02d}:{int(parts[1]):02d}"
    return value


def _lesson_sort_key(row):
    time_text = _normalize_time_text(getattr(row, 'start_time', '') or '')
    pieces = time_text.split(':')
    if len(pieces) >= 2 and pieces[0].isdigit() and pieces[1].isdigit():
        return (int(pieces[0]), int(pieces[1]), getattr(row, 'sort_order', 0) or 0)
    return (99, 99, getattr(row, 'sort_order', 0) or 0)


def _safe_text(value):
    if value is None:
        return ""
    return str(value).replace('\xa0', ' ').strip()


def _is_placeholder_subject(value):
    text = _safe_text(value)
    lowered = text.lower()
    if not text:
        return True
    if lowered in {'none', 'nan', '-'}:
        return True
    if '${short_name}' in text:
        return True
    if lowered.startswith('обучающийся:'):
        return True
    return False


def _education_level_for_child(child):
    current_class = child.current_class if child else None
    grade = getattr(current_class, "grade", None)
    if grade is None:
        name = (getattr(current_class, "name", "") or "").strip()
        digits = "".join(ch for ch in name if ch.isdigit())
        grade = int(digits) if digits else None
    if grade is None:
        return "DO" if not current_class else "NOO"
    if grade <= 4:
        return "NOO"
    if grade <= 9:
        return "OOO"
    return "SOO"


def _parallel_for_child(child):
    current_class = child.current_class if child else None
    if not current_class:
        return ""
    if getattr(current_class, "grade", None) is not None:
        return str(current_class.grade)
    digits = "".join(ch for ch in (current_class.name or "") if ch.isdigit())
    return digits or ""


def _parent_info_for_child(child):
    if not child:
        return ""
    rows = []
    for link in getattr(child, "parent_links", []) or []:
        if not link.parent:
            continue
        relation = {
            "mother": "Мать",
            "father": "Отец",
            "guardian": "Законный представитель",
        }.get((link.relation_type or "").lower(), "Представитель")
        tail = []
        if getattr(link.parent, "phone", None):
            tail.append(link.parent.phone)
        if getattr(link.parent, "email", None):
            tail.append(link.parent.email)
        line = f"{relation}: {link.parent.fio}"
        if tail:
            line += f" ({', '.join(tail)})"
        rows.append(line)
    return "\n".join(rows)


def _support_staff_summary(child):
    if not child:
        return ""
    items = []
    assignments = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.specialist))
        .filter_by(child_id=child.id, status="ACTIVE")
        .order_by(ServiceAssignment.created_at.asc())
        .all()
    )
    for row in assignments:
        spec = row.specialist
        if not spec:
            continue
        title = row.role_title or spec.position_title or "Специалист"
        items.append(f"{title}: {spec.fio}")
    return "\n".join(dict.fromkeys(items))


def _specialist_access_child_ids():
    spec = ServiceSpecialist.query.filter_by(user_id=current_user.id, is_active=True).first()
    if not spec:
        return set()
    ids = {
        row.child_id
        for row in ServiceAssignment.query.filter_by(specialist_id=spec.id, status="ACTIVE").all()
        if row.child_id
    }
    return ids


def _is_sppiss_head(user=None):
    user = user or current_user
    if is_admin(user):
        return True
    if has_role("METHODIST", user=user):
        return True
    spec = ServiceSpecialist.query.filter_by(user_id=getattr(user, "id", None), is_active=True).first()
    if not spec:
        return False
    return ServiceResponsible.query.filter_by(specialist_id=spec.id, is_active=True).first() is not None


def _can_open_module(user=None):
    user = user or current_user
    if not getattr(user, "is_authenticated", False):
        return False
    return has_any_role(*ROLE_PRIORITY, user=user) or _is_sppiss_head(user)


def _can_view_card(card, user=None):
    user = user or current_user
    if is_admin(user) or _is_sppiss_head(user):
        return True
    if has_role("CLASS_TEACHER", user=user) and card.child and is_class_teacher_of_child(card.child, user=user):
        return True
    child_ids = _specialist_access_child_ids() if getattr(user, 'id', None) == getattr(current_user, 'id', None) else set()
    if card.child_id in child_ids:
        return True
    return False


def _can_edit_card(card=None, user=None):
    user = user or current_user
    return is_admin(user) or _is_sppiss_head(user)


def _can_delete_card(card, user=None):
    user = user or current_user
    return is_admin(user) and (card.status or "").upper() == "DRAFT"


def _ensure_module_access():
    if not _can_open_module():
        abort(403)


def _append_history(card, action, comment=""):
    db.session.add(IomHistory(
        iom_card=card,
        action=action,
        comment=comment,
        created_by_user_id=getattr(current_user, "id", None),
    ))



def _time_to_minutes(value):
    text = _normalize_time_text(value)
    parts = text.split(':')
    if len(parts) >= 2 and parts[0].isdigit() and parts[1].isdigit():
        return int(parts[0]) * 60 + int(parts[1])
    return None


def _minutes_to_hhmm(minutes):
    if minutes is None:
        return ''
    hours = max(minutes, 0) // 60
    mins = max(minutes, 0) % 60
    return f"{hours:02d}:{mins:02d}"


def _find_specialist_for_user(user_id):
    if not user_id:
        return None
    return ServiceSpecialist.query.filter_by(user_id=user_id, is_active=True).first()


def _assigned_specialists(card):
    rows = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.specialist))
        .filter_by(child_id=card.child_id, status='ACTIVE')
        .order_by(ServiceAssignment.created_at.asc())
        .all()
    )
    unique = []
    seen = set()
    for row in rows:
        spec = row.specialist
        if not spec or spec.id in seen:
            continue
        seen.add(spec.id)
        unique.append({
            'specialist': spec,
            'role_title': row.role_title or spec.position_title or 'Специалист',
            'display_name': f"{spec.fio} — {row.role_title or spec.position_title or 'Специалист'}",
        })
    return unique


def _can_manage_corrections(card, user=None):
    user = user or current_user
    if is_admin(user) or _is_sppiss_head(user):
        return True
    spec = _find_specialist_for_user(getattr(user, 'id', None))
    if not spec:
        return False
    return ServiceAssignment.query.filter_by(child_id=card.child_id, specialist_id=spec.id, status='ACTIVE').first() is not None


def _overlap_exists(card, weekday, start_time, end_time, exclude_correction_id=None):
    start_minutes = _time_to_minutes(start_time)
    end_minutes = _time_to_minutes(end_time)
    if start_minutes is None or end_minutes is None or end_minutes <= start_minutes:
        return 'Укажите корректный интервал времени занятия.'

    # Школьные уроки: считаем занятость от начала урока до начала следующего урока, для последнего — 45 минут.
    day_lessons = [x for x in (card.schedule_lessons or []) if x.weekday == weekday]
    day_lessons = sorted(day_lessons, key=_lesson_sort_key)
    lesson_ranges = []
    for index, lesson in enumerate(day_lessons):
        lesson_start = _time_to_minutes(lesson.start_time)
        if lesson_start is None:
            continue
        if index + 1 < len(day_lessons):
            next_start = _time_to_minutes(day_lessons[index + 1].start_time)
            lesson_end = next_start if next_start and next_start > lesson_start else lesson_start + 45
        else:
            lesson_end = lesson_start + 45
        lesson_ranges.append((lesson_start, lesson_end, lesson.subject_name))

    for lesson_start, lesson_end, lesson_name in lesson_ranges:
        if start_minutes < lesson_end and end_minutes > lesson_start:
            return f'Время пересекается со школьным уроком: {lesson_name}.'

    for row in card.schedule_corrections or []:
        if row.weekday != weekday:
            continue
        if exclude_correction_id and row.id == exclude_correction_id:
            continue
        row_start = _time_to_minutes(row.start_time)
        row_end = _time_to_minutes(row.end_time)
        if row_start is None or row_end is None:
            continue
        if start_minutes < row_end and end_minutes > row_start:
            spec_name = row.specialist.fio if row.specialist else 'другим занятием'
            return f'Время пересекается с коррекционным занятием специалиста: {spec_name}.'
    return ''


def _correction_groups(card):
    grouped = {day: [] for day in WEEKDAY_ORDER}
    for row in sorted(card.schedule_corrections or [], key=lambda x: (WEEKDAY_ORDER.index(x.weekday) if x.weekday in WEEKDAY_ORDER else 99, _time_to_minutes(x.start_time) or 9999, _time_to_minutes(x.end_time) or 9999)):
        if row.weekday in grouped:
            grouped[row.weekday].append(row)
    return grouped


def _lesson_with_end_ranges(card):
    result = {day: [] for day in WEEKDAY_ORDER}
    for day in WEEKDAY_ORDER:
        rows = sorted([x for x in (card.schedule_lessons or []) if x.weekday == day], key=_lesson_sort_key)
        for index, row in enumerate(rows):
            start_minutes = _time_to_minutes(row.start_time)
            end_minutes = None
            if start_minutes is not None:
                if index + 1 < len(rows):
                    next_start = _time_to_minutes(rows[index + 1].start_time)
                    end_minutes = next_start if next_start and next_start > start_minutes else start_minutes + 45
                else:
                    end_minutes = start_minutes + 45
            result[day].append({
                'start_time': row.start_time,
                'end_time': _minutes_to_hhmm(end_minutes),
                'subject_name': row.subject_name,
                'source_type': row.source_type,
            })
    return result

def _registry_base_query():
    q = IomCard.query.options(
        joinedload(IomCard.child),
        joinedload(IomCard.academic_year),
        joinedload(IomCard.building),
        joinedload(IomCard.school_class),
    )
    if is_admin() or _is_sppiss_head():
        return q
    if has_role("CLASS_TEACHER"):
        return q.join(Child, IomCard.child_id == Child.id)
    child_ids = _specialist_access_child_ids()
    if child_ids:
        return q.filter(IomCard.child_id.in_(child_ids))
    return q.filter(IomCard.id == -1)


def _apply_filters(query):
    academic_year_id = _parse_int(request.args.get("academic_year_id"))
    iom_type = (request.args.get("iom_type") or "").strip().upper()
    education_level = (request.args.get("education_level") or "").strip().upper()
    building_id = _parse_int(request.args.get("building_id"))
    parallel = (request.args.get("parallel") or "").strip()
    school_class_id = _parse_int(request.args.get("school_class_id"))
    status = (request.args.get("status") or "").strip().upper()
    q = (request.args.get("q") or "").strip()

    if academic_year_id:
        query = query.filter(IomCard.academic_year_id == academic_year_id)
    if iom_type:
        query = query.filter(IomCard.iom_type == iom_type)
    if education_level:
        query = query.filter(IomCard.education_level == education_level)
    if building_id:
        query = query.filter(IomCard.building_id == building_id)
    if parallel:
        query = query.filter(IomCard.parallel == parallel)
    if school_class_id:
        query = query.filter(IomCard.school_class_id == school_class_id)
    if status:
        query = query.filter(IomCard.status == status)
    if q:
        like = f"%{q}%"
        query = query.filter(or_(IomCard.student_fio.ilike(like), IomCard.class_name.ilike(like), IomCard.notes.ilike(like)))
    if has_role("CLASS_TEACHER") and not (is_admin() or _is_sppiss_head()):
        query = query.filter(IomCard.curator_user_id == current_user.id)
    return query


def _seed_defaults(card, child):
    current_class = child.current_class if child else None
    card.student_fio = child.fio if child else card.student_fio
    card.birth_date = child.birth_date if child else card.birth_date
    card.birth_year = child.birth_date.year if getattr(child, 'birth_date', None) else card.birth_year
    card.education_level = _education_level_for_child(child)
    if current_class:
        card.school_class_id = current_class.id
        card.class_name = current_class.name
        card.parallel = _parallel_for_child(child)
        if current_class.building:
            card.building_id = current_class.building.id
            card.building_name = current_class.building.short_name or current_class.building.name
        if current_class.teacher_user:
            card.curator_user_id = current_class.teacher_user.id
            card.curator_name = current_class.teacher_user.fio
            card.class_teacher_name = current_class.teacher_user.fio
    card.ovz_status = "ОВЗ" if getattr(child, "is_ovz", False) else ""
    card.nosology = getattr(child, "ovz_nosology", None) or ""
    variant = getattr(child, "ovz_variant", None)
    if variant:
        card.aop_variant = f"Вариант {variant}"
    elif getattr(child, "social", None) and getattr(child.social, "aoop_variant_text", None):
        card.aop_variant = child.social.aoop_variant_text
    card.parent_info = _parent_info_for_child(child)
    card.support_staff_summary = _support_staff_summary(child)
    if not card.sppiss_head_name:
        card.sppiss_head_name = "Руководитель СППиСС"
    if not card.director_name:
        settings = get_active_organization_settings()
        default_director = (getattr(settings, 'director_name', None) or '').strip()
        default_position = (getattr(settings, 'director_position', None) or '').strip()
        card.director_name = f'{default_position} {default_director}'.strip() or get_organization_signature_block()




def _iom_readiness(card):
    checks = []
    def add(key, label, ok, hint=''):
        checks.append({'key': key, 'label': label, 'ok': bool(ok), 'hint': hint})

    add('student', 'Выбран ребенок из реестра ОВЗ', bool(card.child_id and card.student_fio), 'Карточка должна быть привязана к ребенку.')
    add('year', 'Указан учебный год и тип ИОМ', bool(card.academic_year_id and card.iom_type), 'Проверьте учебный год и тип ИОМ.')
    add('period', 'Заполнен срок реализации', bool(card.start_date and card.end_date), 'Нужно указать даты начала и окончания ИОМ.')
    add('aop', 'Выбрана программа АООП', bool((card.aop_variant or '').strip()), 'Выберите АООП из списка по уровню образования.')
    add('parents', 'Есть сведения о родителях / законных представителях', bool((card.parent_info or '').strip()), 'Проверьте карточку ребенка.')
    add('support', 'Назначены специалисты сопровождения', len(_assigned_specialists(card)) > 0, 'Сначала назначьте специалистов в сопровождении.')
    add('schedule', 'Загружено школьное расписание', len(card.schedule_lessons or []) > 0, 'Загрузите или заполните расписание ребенка.')
    add('corrections', 'Добавлены коррекционные занятия', len(card.schedule_corrections or []) > 0, 'Добавьте хотя бы одно коррекционное занятие.')
    add('head', 'Указан руководитель СППиСС', bool((card.sppiss_head_name or '').strip()), 'Заполните ФИО руководителя СППиСС.')
    add('director', 'Указан директор', bool((card.director_name or '').strip()), 'Заполните ФИО директора.')
    add('consent', 'Заполнена отметка об ознакомлении / согласии', bool((card.consent_mark or '').strip()), 'Поле нужно перед выпуском итогового ИОМ.')

    ready_on_approval = all(item['ok'] for item in checks[:8])
    ready_head = ready_on_approval and checks[8]['ok']
    ready_approved = ready_head and checks[9]['ok'] and checks[10]['ok']
    completed = sum(1 for item in checks if item['ok'])
    total = len(checks)
    return {
        'checks': checks,
        'completed': completed,
        'total': total,
        'percent': int(round((completed / total) * 100)) if total else 0,
        'ready_for_on_approval': ready_on_approval,
        'ready_for_head': ready_head,
        'ready_for_approved': ready_approved,
    }


def _validate_status_transition(card):
    readiness = _iom_readiness(card)
    status = (card.status or 'DRAFT').upper()
    if status == 'ON_APPROVAL' and not readiness['ready_for_on_approval']:
        missing = [item['label'] for item in readiness['checks'][:8] if not item['ok']]
        raise ValueError('Для статуса «На согласовании» заполните: ' + '; '.join(missing) + '.')
    if status == 'APPROVED_BY_HEAD' and not readiness['ready_for_head']:
        missing = [item['label'] for item in readiness['checks'][:9] if not item['ok']]
        raise ValueError('Для статуса «Согласован руководителем СППиСС» заполните: ' + '; '.join(missing) + '.')
    if status == 'APPROVED' and not readiness['ready_for_approved']:
        missing = [item['label'] for item in readiness['checks'] if not item['ok']]
        raise ValueError('Для статуса «Утвержден» заполните: ' + '; '.join(missing) + '.')
    return readiness

def _update_card_from_form(card):
    original = {
        "status": card.status,
        "student_fio": card.student_fio,
        "education_level": card.education_level,
        "class_name": card.class_name,
        "parallel": card.parallel,
        "ovz_status": card.ovz_status,
        "nosology": card.nosology,
        "aop_variant": card.aop_variant,
        "curator_name": card.curator_name,
        "sppiss_head_name": card.sppiss_head_name,
        "director_name": card.director_name,
        "parent_info": card.parent_info,
        "notes": card.notes,
    }

    child_id = _parse_int(request.form.get("child_id"))
    academic_year_id = _parse_int(request.form.get("academic_year_id"))
    iom_type = (request.form.get("iom_type") or "").strip().upper()
    status = (request.form.get("status") or "DRAFT").strip().upper()
    start_date = _parse_date(request.form.get("start_date"))
    end_date = _parse_date(request.form.get("end_date"))
    if not child_id or not academic_year_id or not iom_type or not start_date or not end_date:
        raise ValueError("Заполните обязательные поля: ребенок, учебный год, тип ИОМ и срок реализации.")

    child = Child.query.get_or_404(child_id)
    if not getattr(child, "is_ovz", False):
        raise ValueError("Для этапа 1 ИОМ можно создавать только для детей из реестра ОВЗ.")

    card.child_id = child.id
    card.academic_year_id = academic_year_id
    card.iom_type = iom_type
    card.status = status
    card.start_date = start_date
    card.end_date = end_date
    _seed_defaults(card, child)

    card.student_fio = (request.form.get("student_fio") or card.student_fio or "").strip()
    card.birth_date = _parse_date(request.form.get("birth_date")) or card.birth_date
    birth_year = (request.form.get("birth_year") or "").strip()
    card.birth_year = int(birth_year) if birth_year.isdigit() else card.birth_year
    card.education_level = (request.form.get("education_level") or card.education_level or "").strip().upper()
    card.class_name = (request.form.get("class_name") or card.class_name or "").strip()
    card.parallel = (request.form.get("parallel") or card.parallel or "").strip()
    building_id = _parse_int(request.form.get("building_id"))
    if building_id:
        building = Building.query.get(building_id)
        if building:
            card.building_id = building.id
            card.building_name = building.short_name or building.name
    else:
        card.building_id = None
        card.building_name = (request.form.get("building_name") or card.building_name or "").strip()
    class_id = _parse_int(request.form.get("school_class_id"))
    if class_id:
        school_class = SchoolClass.query.get(class_id)
        if school_class:
            card.school_class_id = school_class.id
            card.class_name = school_class.name
    else:
        card.school_class_id = None
    card.ovz_status = (request.form.get("ovz_status") or card.ovz_status or "").strip()
    card.nosology = (request.form.get("nosology") or card.nosology or "").strip()
    selected_aop_variant = (request.form.get("aop_variant") or "").strip()
    card.aop_variant = AOP_VARIANT_LABELS.get(selected_aop_variant, selected_aop_variant or card.aop_variant or "")
    card.parent_info = (request.form.get("parent_info") or card.parent_info or "").strip()
    card.curator_name = (request.form.get("curator_name") or card.curator_name or "").strip()
    curator_user_id = _parse_int(request.form.get("curator_user_id"))
    card.curator_user_id = curator_user_id
    if curator_user_id:
        curator = User.query.get(curator_user_id)
        if curator:
            card.curator_name = curator.fio
    card.sppiss_head_name = (request.form.get("sppiss_head_name") or card.sppiss_head_name or "").strip()
    card.director_name = (request.form.get("director_name") or card.director_name or "").strip()
    card.consent_mark = (request.form.get("consent_mark") or "").strip()
    card.agreed_at = _parse_date(request.form.get("agreed_at"))
    card.approved_at = _parse_date(request.form.get("approved_at"))
    card.notes = (request.form.get("notes") or "").strip()
    card.updated_by_user_id = current_user.id

    _validate_status_transition(card)

    changes = []
    for key, old in original.items():
        new = getattr(card, key)
        if old != new:
            changes.append(key)
    if changes:
        _append_history(card, "update", "Изменены поля: " + ", ".join(changes))
    if original["status"] != card.status:
        if card.status == "APPROVED_BY_HEAD":
            card.agreed_by_user_id = current_user.id
        if card.status == "APPROVED":
            card.approved_by_user_id = current_user.id
        _append_history(card, "status", f"Статус изменен на: {card.status_label}")


def _extract_schedule_rows_from_workbook(file_storage):
    filename = (getattr(file_storage, 'filename', None) or '').strip() or 'schedule.xlsx'
    payload = file_storage.read()
    if not payload:
        raise ValueError('Файл расписания пустой.')
    try:
        workbook = load_workbook(BytesIO(payload), data_only=True)
    except Exception as exc:
        raise ValueError('Не удалось прочитать Excel-файл расписания. Проверьте, что это корректная выгрузка .xlsx.') from exc

    worksheet = workbook[workbook.sheetnames[0]]
    rows = []
    sort_order = 1
    day_blocks = [
        (1, 2, 3),
        (5, 6, 7),
        (9, 10, 11),
        (13, 14, 15),
        (17, 18, 19),
    ]
    for lesson_row in range(7, worksheet.max_row + 1):
        for number_col, time_col, subject_col in day_blocks:
            weekday_raw = _safe_text(worksheet.cell(5, subject_col).value).lower()
            weekday_key = WEEKDAY_RU_TO_KEY.get(weekday_raw)
            if weekday_key not in WEEKDAY_ORDER:
                continue
            subject_text = _safe_text(worksheet.cell(lesson_row, subject_col).value)
            time_text = _normalize_time_text(_safe_text(worksheet.cell(lesson_row, time_col).value))
            if _is_placeholder_subject(subject_text):
                continue
            if not time_text:
                continue
            rows.append({
                'weekday': weekday_key,
                'start_time': time_text,
                'subject_name': subject_text,
                'sort_order': sort_order,
            })
            sort_order += 1
    if not rows:
        raise ValueError('В файле не найдено уроков для понедельника–пятницы. Проверьте структуру выгрузки МЭШ.')
    return filename, rows


def _replace_schedule_lessons(card, rows, filename='schedule.xlsx'):
    deleted = IomScheduleLesson.query.filter_by(iom_card_id=card.id).delete()
    loaded = 0
    for item in rows:
        lesson = IomScheduleLesson(
            iom_card_id=card.id,
            weekday=item['weekday'],
            start_time=item['start_time'],
            subject_name=item['subject_name'],
            source_type=item.get('source_type') or 'IMPORTED',
            sort_order=item.get('sort_order') or (loaded + 1),
        )
        db.session.add(lesson)
        loaded += 1
    db.session.add(IomImportSessionSchedule(
        iom_card_id=card.id,
        filename=filename,
        rows_loaded=loaded,
        imported_by_user_id=getattr(current_user, 'id', None),
        comment='Импорт школьного расписания МЭШ',
    ))
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'schedule_import', f'Загружено школьное расписание: {loaded} строк из файла {filename}.')
    return deleted, loaded


def _schedule_groups(card):
    grouped = {day: [] for day in WEEKDAY_ORDER}
    for row in sorted(card.schedule_lessons or [], key=lambda x: (WEEKDAY_ORDER.index(x.weekday) if x.weekday in WEEKDAY_ORDER else 99, _lesson_sort_key(x))):
        if row.weekday in grouped:
            grouped[row.weekday].append(row)
    return grouped


def _schedule_stats(card):
    lessons = card.schedule_lessons or []
    imported_rows = sum(1 for row in lessons if (row.source_type or '').upper() == 'IMPORTED')
    manual_rows = sum(1 for row in lessons if (row.source_type or '').upper() == 'MANUAL')
    return {
        'total': len(lessons),
        'imported': imported_rows,
        'manual': manual_rows,
        'last_import': card.schedule_import_sessions[0] if getattr(card, 'schedule_import_sessions', None) else None,
    }


def _schedule_context(card):
    return {
        'weekday_order': WEEKDAY_ORDER,
        'weekday_labels': WEEKDAY_LABELS,
        'schedule_by_day': _schedule_groups(card),
        'schedule_with_ranges': _lesson_with_end_ranges(card),
        'schedule_stats': _schedule_stats(card),
        'corrections_by_day': _correction_groups(card),
        'assigned_specialists': _assigned_specialists(card),
    }


def _section_row(card, code, title=''):
    row = next((x for x in getattr(card, 'section_rows', []) if x.section_code == code), None)
    if row:
        return row
    row = IomSectionData(iom_card_id=card.id, section_code=code, section_title=title or code)
    db.session.add(row)
    return row


def _section_payload(card, code, defaults=None):
    defaults = defaults or {}
    row = next((x for x in getattr(card, 'section_rows', []) if x.section_code == code), None)
    if not row or not row.payload_json:
        return dict(defaults)
    try:
        data = json.loads(row.payload_json)
        if isinstance(data, dict):
            result = dict(defaults)
            result.update(data)
            return result
    except Exception:
        pass
    return dict(defaults)


def _save_section(card, code, title, field_names):
    row = _section_row(card, code, title)
    payload = {}
    for field in field_names:
        payload[field] = (request.form.get(field) or '').strip()
    row.section_title = title
    row.payload_json = json.dumps(payload, ensure_ascii=False)
    row.updated_by_user_id = getattr(current_user, 'id', None)
    if not row.created_by_user_id:
        row.created_by_user_id = getattr(current_user, 'id', None)
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, f'section_{code}', f'Сохранен раздел «{title}».')


def _specialist_role_choices(card):
    choices = []
    seen = set()
    for item in _assigned_specialists(card):
        role_title = item['role_title'] or 'Специалист'
        if role_title not in seen:
            seen.add(role_title)
            choices.append(role_title)
    fallback = ['Педагог-психолог', 'Учитель-логопед', 'Учитель-дефектолог', 'Социальный педагог', 'Тьютор', 'Ассистент']
    for value in fallback:
        if value not in seen:
            choices.append(value)
    return choices


def _sync_specialist_plans_from_assignments(card):
    existing_keys = {(row.specialist_id, (row.role_title or '').strip().lower()) for row in getattr(card, 'specialist_plans', [])}
    created = 0
    for idx, row in enumerate((
        ServiceAssignment.query.options(joinedload(ServiceAssignment.specialist))
        .filter_by(child_id=card.child_id, status='ACTIVE')
        .order_by(ServiceAssignment.created_at.asc())
        .all()
    ), start=1):
        spec = row.specialist
        if not spec:
            continue
        role_title = (row.role_title or spec.position_title or 'Специалист').strip()
        key = (spec.id, role_title.lower())
        if key in existing_keys:
            continue
        db.session.add(IomSpecialistPlan(
            iom_card_id=card.id,
            role_title=role_title,
            specialist_id=spec.id,
            assignment_id=row.id,
            sort_order=idx * 10,
            created_by_user_id=getattr(current_user, 'id', None),
            updated_by_user_id=getattr(current_user, 'id', None),
        ))
        existing_keys.add(key)
        created += 1
    return created


def _related_cyclegrams(card):
    specialist_ids = [item['specialist'].id for item in _assigned_specialists(card)]
    if not specialist_ids:
        return []
    academic_year_name = card.academic_year.name if card.academic_year else ''
    rows = ServiceCyclegram.query.options(joinedload(ServiceCyclegram.specialist)).filter(ServiceCyclegram.specialist_id.in_(specialist_ids))
    if academic_year_name:
        rows = rows.filter(ServiceCyclegram.academic_year == academic_year_name)
    return rows.order_by(ServiceCyclegram.updated_at.desc()).all()


def _find_or_create_iom_activity_type():
    row = ServiceActivityType.query.filter_by(code='IOM_CORRECTION').first()
    if row:
        return row
    row = ServiceActivityType(
        code='IOM_CORRECTION',
        name='Коррекционно-развивающее занятие (ИОМ)',
        work_category='PRACTICAL',
        specialist_scope='IOM',
        template_text='ИОМ',
        requires_child=False,
        requires_group=True,
        is_group_activity=True,
        is_active=True,
        sort_order=25,
    )
    db.session.add(row)
    db.session.flush()
    return row


def _cyclegram_sync_key(card, correction):
    academic_year = card.academic_year.name if card.academic_year else ''
    return f"{academic_year}|{correction.specialist_id}|{correction.weekday}|{correction.start_time}|{correction.end_time}|{(correction.course_name or '').strip().lower()}"


def _sync_card_corrections_to_cyclegrams(card):
    activity_type = _find_or_create_iom_activity_type()
    academic_year_name = card.academic_year.name if card.academic_year else ''
    created_count = 0
    updated_count = 0
    synced_specs = set()
    current_correction_ids = {row.id for row in (card.schedule_corrections or [])}
    for orphan in IomCyclegramLink.query.join(IomScheduleCorrection, IomCyclegramLink.correction_id == IomScheduleCorrection.id).filter(IomScheduleCorrection.iom_card_id == card.id).all():
        if orphan.correction_id not in current_correction_ids:
            db.session.delete(orphan)
    groups = {}
    for correction in card.schedule_corrections or []:
        key = _cyclegram_sync_key(card, correction)
        groups.setdefault(key, []).append(correction)
    for key, rows in groups.items():
        first = rows[0]
        cyclegram = ServiceCyclegram.query.filter_by(specialist_id=first.specialist_id, academic_year=academic_year_name).order_by(ServiceCyclegram.id.desc()).first()
        if not cyclegram:
            continue
        synced_specs.add(first.specialist_id)
        all_matches = IomScheduleCorrection.query.join(IomCard, IomScheduleCorrection.iom_card_id == IomCard.id).filter(
            IomCard.academic_year_id == card.academic_year_id,
            IomScheduleCorrection.specialist_id == first.specialist_id,
            IomScheduleCorrection.weekday == first.weekday,
            IomScheduleCorrection.start_time == first.start_time,
            IomScheduleCorrection.end_time == first.end_time,
            IomScheduleCorrection.course_name == first.course_name,
        ).all()
        child_names = []
        for item in all_matches:
            if item.iom_card and item.iom_card.student_fio:
                child_names.append(item.iom_card.student_fio)
        child_names = list(dict.fromkeys(child_names))
        group_text = ', '.join(child_names)
        description = first.course_name
        minutes = max((_time_to_minutes(first.end_time) or 0) - (_time_to_minutes(first.start_time) or 0), 0)
        existing_link = IomCyclegramLink.query.filter_by(sync_key=key).first()
        entry = existing_link and ServiceCyclegramEntry.query.get(existing_link.cyclegram_entry_id) or None
        if entry and entry.cyclegram_id != cyclegram.id:
            entry = None
        if not entry:
            entry = ServiceCyclegramEntry(
                cyclegram_id=cyclegram.id,
                weekday=WEEKDAY_ORDER.index(first.weekday) + 1 if first.weekday in WEEKDAY_ORDER else 1,
                start_time=first.start_time,
                end_time=first.end_time,
                activity_type_id=activity_type.id,
                description=description,
                group_text=group_text,
                work_category='PRACTICAL',
                minutes=minutes,
                comment='Синхронизировано из ИОМ',
                sort_order=100,
            )
            db.session.add(entry)
            db.session.flush()
            created_count += 1
        else:
            entry.weekday = WEEKDAY_ORDER.index(first.weekday) + 1 if first.weekday in WEEKDAY_ORDER else entry.weekday
            entry.start_time = first.start_time
            entry.end_time = first.end_time
            entry.description = description
            entry.group_text = group_text
            entry.minutes = minutes
            entry.comment = 'Синхронизировано из ИОМ'
            updated_count += 1
        existing_links = IomCyclegramLink.query.filter_by(sync_key=key).all()
        keep_ids = {x.id for x in all_matches}
        for link in existing_links:
            if link.correction_id not in keep_ids:
                db.session.delete(link)
        linked_correction_ids = {x.correction_id for x in existing_links}
        for item in all_matches:
            if item.id in linked_correction_ids:
                continue
            db.session.add(IomCyclegramLink(correction_id=item.id, cyclegram_entry_id=entry.id, sync_key=key, synced_by_user_id=getattr(current_user, 'id', None)))
    return created_count, updated_count, synced_specs


def _section_context(card):
    analysis_defaults = {
        'conclusion_number': '', 'conclusion_date': '', 'valid_until': '', 'pmpk_recommendations': '', 'ipra_recommendations': '',
        'special_conditions': '', 'required_conditions': '', 'implementation_mark': '', 'comment': '',
    }
    goals_defaults = {
        'main_goal': '', 'support_tasks': '', 'learning_tasks': '', 'correction_tasks': '', 'socialization_tasks': '', 'parent_interaction_forms': '', 'extra_comment': '', 'technologies': '',
    }
    ppk_defaults = {
        'general_info': '', 'reason_text': '', 'ppk_conclusion': '', 'correction_directions': '', 'teacher_recommendations': '', 'parent_recommendations': '', 'individual_recommendations': '', 'repeat_review_term': '',
    }
    teacher_defaults = {
        'program_features': '', 'meta_results': '', 'personal_results': '', 'uud_bud': '', 'difficulty_reasons': '', 'overcoming_forms': '', 'teacher_recommendations': '', 'support_measures': '', 'class_teacher_comment': '',
    }
    return {
        'analysis_section': _section_payload(card, 'analysis', analysis_defaults),
        'goals_section': _section_payload(card, 'goals', goals_defaults),
        'ppk_section': _section_payload(card, 'ppk', ppk_defaults),
        'teacher_section': _section_payload(card, 'teacher', teacher_defaults),
        'parent_work_rows': _section_payload(card, 'parent_work', {'rows': []}).get('rows', []),
        'specialist_role_choices': _specialist_role_choices(card),
        'related_cyclegrams': _related_cyclegrams(card),
    }


def _seed_monitoring_templates():
    if IomMonitoringTemplate.query.first():
        return
    for iom_type, rows in DEFAULT_MONITORING_TEMPLATES.items():
        for idx, (period, block_code, block_title, line_code, line_title, scale_type) in enumerate(rows, start=1):
            db.session.add(IomMonitoringTemplate(
                iom_type=iom_type,
                period=period,
                block_code=block_code,
                block_title=block_title,
                line_code=line_code,
                line_title=line_title,
                scale_type=scale_type,
                sort_order=idx * 10,
                is_enabled=True,
            ))
    db.session.flush()


def _monitoring_templates(card):
    _seed_monitoring_templates()
    rows = IomMonitoringTemplate.query.filter_by(iom_type=card.iom_type, is_enabled=True).order_by(
        IomMonitoringTemplate.period.asc(), IomMonitoringTemplate.sort_order.asc(), IomMonitoringTemplate.id.asc()
    ).all()
    grouped = defaultdict(lambda: defaultdict(list))
    for row in rows:
        grouped[row.period][row.block_code].append(row)
    return grouped


def _monitoring_payload(card, period, block_code):
    row = IomMonitoringEntry.query.filter_by(iom_card_id=card.id, period=period, block_code=block_code).first()
    data = {"values": {}, "comments": {}, "conclusion": ""}
    if row and row.payload_json:
        try:
            parsed = json.loads(row.payload_json)
            if isinstance(parsed, dict):
                data.update(parsed)
        except Exception:
            pass
    return data, row


def _monitoring_context(card):
    templates = _monitoring_templates(card)
    payloads = {}
    for period, blocks in templates.items():
        payloads[period] = {}
        for block_code in blocks.keys():
            payloads[period][block_code], _ = _monitoring_payload(card, period, block_code)
    return {
        'monitoring_templates': templates,
        'monitoring_payloads': payloads,
        'monitoring_period_choices': MONITORING_PERIOD_CHOICES,
        'monitoring_scale_options': MONITORING_SCALE_OPTIONS,
    }


def _can_manage_archive(card=None, user=None):
    user = user or current_user
    return is_admin(user) or _is_sppiss_head(user)


def _can_export_card(card, user=None):
    user = user or current_user
    return _can_view_card(card, user=user)


def _monitoring_required_done(card):
    templates = _monitoring_templates(card)
    for period, blocks in templates.items():
        for block_code in blocks.keys():
            data, _ = _monitoring_payload(card, period, block_code)
            vals = data.get('values') or {}
            if not any((str(v).strip() for v in vals.values())) and not (data.get('conclusion') or '').strip():
                return False
    return True if templates else True


def _card_export_context(card):
    ctx = {}
    ctx.update(_schedule_context(card))
    ctx.update(_section_context(card))
    ctx.update(_monitoring_context(card))
    return ctx


def _add_export_log(card, fmt):
    db.session.add(IomExportLog(
        iom_card_id=card.id,
        export_format=fmt,
        status_snapshot=card.status,
        exported_by_user_id=getattr(current_user, 'id', None),
    ))
    _append_history(card, 'export', f'Выгружен документ {fmt}.')


def _build_docx(card):
    if Document is None:
        raise ValueError('Библиотека python-docx недоступна.')
    doc = Document()
    doc.add_heading('Индивидуальный образовательный маршрут', level=1)
    doc.add_paragraph(f'ФИО: {card.student_fio or ""}')
    doc.add_paragraph(f'Учебный год: {card.academic_year.name if card.academic_year else ""}')
    doc.add_paragraph(f'Тип ИОМ: {card.iom_type_label}')
    doc.add_paragraph(f'Статус: {card.status_label}')
    doc.add_paragraph(f'АООП: {card.aop_variant or ""}')
    doc.add_paragraph(f'Класс / группа: {card.class_name or ""}')
    doc.add_heading('Школьное расписание', level=2)
    for day in WEEKDAY_ORDER:
        items = [x for x in card.schedule_lessons if x.weekday == day]
        if not items:
            continue
        doc.add_paragraph(WEEKDAY_LABELS.get(day, day), style='List Bullet')
        for row in items:
            doc.add_paragraph(f'{row.start_time} — {row.subject_name}', style='List Bullet 2')
    mctx = _monitoring_context(card)
    if mctx['monitoring_templates']:
        doc.add_heading('Мониторинг', level=2)
        for period, blocks in mctx['monitoring_templates'].items():
            doc.add_heading(MONITORING_PERIOD_LABELS.get(period, period), level=3)
            for block_code, lines in blocks.items():
                doc.add_paragraph(lines[0].block_title)
                data = mctx['monitoring_payloads'][period][block_code]
                tbl = doc.add_table(rows=1, cols=3)
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Показатель'
                hdr[1].text = 'Значение'
                hdr[2].text = 'Комментарий'
                for line in lines:
                    vals = data.get('values') or {}
                    comments = data.get('comments') or {}
                    cells = tbl.add_row().cells
                    cells[0].text = line.line_title
                    cells[1].text = str(vals.get(line.line_code, '') or '')
                    cells[2].text = str(comments.get(line.line_code, '') or '')
                if data.get('conclusion'):
                    doc.add_paragraph('Заключение: ' + data['conclusion'])
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _build_pdf(card):
    if canvas is None or A4 is None:
        raise ValueError('Библиотека ReportLab недоступна.')
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    left = 40
    top = height - 40
    bottom = 50
    max_width = width - left * 2
    y = top
    base_font, bold_font = _ensure_pdf_fonts()

    def line(txt, size=11, bold=False, gap=4):
        nonlocal y
        font_name = bold_font if bold else base_font
        c.setFont(font_name, size)
        for chunk in _pdf_wrap_lines(txt, font_name, size, max_width):
            if y < bottom:
                c.showPage()
                y = top
                c.setFont(font_name, size)
            c.drawString(left, y, chunk)
            y -= size + gap

    line('Индивидуальный образовательный маршрут', 14, bold=True, gap=6)
    line(f'ФИО: {card.student_fio or ""}')
    line(f'Учебный год: {card.academic_year.name if card.academic_year else ""}')
    line(f'Тип ИОМ: {card.iom_type_label}')
    line(f'Статус: {card.status_label}')
    line(f'АООП: {card.aop_variant or ""}')
    line(f'Класс / группа: {card.class_name or ""}')
    line('Школьное расписание', 12, bold=True, gap=5)
    has_schedule = False
    for day in WEEKDAY_ORDER:
        items = [x for x in card.schedule_lessons if x.weekday == day]
        if not items:
            continue
        has_schedule = True
        line(WEEKDAY_LABELS.get(day, day), 11, bold=True)
        for row in items:
            line(f'{row.start_time} — {row.subject_name}', 10)
    if not has_schedule:
        line('Школьное расписание не заполнено.', 10)

    mctx = _monitoring_context(card)
    if mctx['monitoring_templates']:
        line('Мониторинг', 12, bold=True, gap=5)
        for period, blocks in mctx['monitoring_templates'].items():
            line(MONITORING_PERIOD_LABELS.get(period, period), 11, bold=True)
            for block_code, lines in blocks.items():
                line(lines[0].block_title, 10, bold=True)
                data = mctx['monitoring_payloads'][period][block_code]
                vals = data.get('values') or {}
                comments = data.get('comments') or {}
                for t in lines:
                    value = vals.get(t.line_code, '') or ''
                    comment = comments.get(t.line_code, '') or ''
                    suffix = f' — {comment}' if comment else ''
                    line(f'{t.line_title}: {value}{suffix}', 9)
                if data.get('conclusion'):
                    line('Заключение: ' + data['conclusion'], 9)

    c.save()
    bio.seek(0)
    return bio


def _duplicate_card(card, target_year_id, copy_mode):
    new_card = IomCard(
        child_id=card.child_id, academic_year_id=target_year_id, iom_type=card.iom_type, status='DRAFT',
        student_fio=card.student_fio, birth_date=card.birth_date, birth_year=card.birth_year,
        education_level=card.education_level, school_class_id=card.school_class_id, class_name=card.class_name,
        parallel=card.parallel, building_id=card.building_id, building_name=card.building_name,
        ovz_status=card.ovz_status, nosology=card.nosology, aop_variant=card.aop_variant,
        parent_info=card.parent_info, class_teacher_name=card.class_teacher_name, support_staff_summary=card.support_staff_summary,
        sppiss_head_name=card.sppiss_head_name, director_name=card.director_name, start_date=card.start_date, end_date=card.end_date,
        notes=card.notes, created_by_user_id=getattr(current_user,'id',None), updated_by_user_id=getattr(current_user,'id',None),
        previous_card_id=card.id,
    )
    db.session.add(new_card)
    db.session.flush()
    for row in card.section_rows or []:
        db.session.add(IomSectionData(iom_card_id=new_card.id, section_code=row.section_code, section_title=row.section_title, payload_json=row.payload_json, created_by_user_id=getattr(current_user,'id',None), updated_by_user_id=getattr(current_user,'id',None)))
    for row in card.specialist_plans or []:
        db.session.add(IomSpecialistPlan(iom_card_id=new_card.id, role_title=row.role_title, specialist_id=row.specialist_id, assignment_id=row.assignment_id, recommendation_text=row.recommendation_text, deficits_text=row.deficits_text, resources_text=row.resources_text, tasks_text=row.tasks_text, work_form=row.work_form, sessions_per_week=row.sessions_per_week, course_name=row.course_name, frequency=row.frequency, expected_result=row.expected_result, monitoring_terms=row.monitoring_terms, comment=row.comment, sort_order=row.sort_order, created_by_user_id=getattr(current_user,'id',None), updated_by_user_id=getattr(current_user,'id',None)))
    if copy_mode == 'WITH_SCHEDULE':
        for row in card.schedule_lessons or []:
            db.session.add(IomScheduleLesson(iom_card_id=new_card.id, weekday=row.weekday, start_time=row.start_time, subject_name=row.subject_name, source_type=row.source_type, sort_order=row.sort_order))
        for row in card.schedule_corrections or []:
            db.session.add(IomScheduleCorrection(iom_card_id=new_card.id, specialist_id=row.specialist_id, weekday=row.weekday, start_time=row.start_time, end_time=row.end_time, course_name=row.course_name, notes=row.notes, created_by_user_id=getattr(current_user,'id',None), updated_by_user_id=getattr(current_user,'id',None)))
    _append_history(new_card, 'duplicate_create', f'Карточка создана копированием из ИОМ #{card.id}.')
    return new_card


@iom_bp.route("/")
@login_required
def registry():
    _ensure_module_access()
    page = max(_parse_int(request.args.get("page")) or 1, 1)
    per_page = 20
    query = _apply_filters(_registry_base_query()).order_by(IomCard.updated_at.desc(), IomCard.id.desc())
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    return render_template(
        "iom/registry.html",
        cards=pagination.items,
        pagination=pagination,
        years=AcademicYear.query.order_by(AcademicYear.name.desc()).all(),
        buildings=Building.query.order_by(Building.name.asc()).all(),
        classes=SchoolClass.query.order_by(SchoolClass.name.asc()).all(),
        iom_type_choices=IOM_TYPE_CHOICES,
        education_level_choices=EDUCATION_LEVEL_CHOICES,
        status_choices=IOM_STATUS_CHOICES,
    )


@iom_bp.route("/new", methods=["GET", "POST"])
@login_required
def create_card():
    _ensure_module_access()
    if not _can_edit_card():
        abort(403)
    card = IomCard(status="DRAFT", created_by_user_id=current_user.id, updated_by_user_id=current_user.id)
    if request.method == "POST":
        try:
            _update_card_from_form(card)
            db.session.add(card)
            db.session.flush()
            _append_history(card, "create", "Карточка ИОМ создана")
            db.session.commit()
            flash("Карточка ИОМ создана.", "success")
            if request.form.get("save_and_continue"):
                return redirect(url_for("iom.edit_card", card_id=card.id))
            return redirect(url_for("iom.registry"))
        except ValueError as exc:
            db.session.rollback()
            flash(str(exc), "danger")
    selected_child_id = _parse_int(request.values.get("child_id"))
    if selected_child_id:
        child = Child.query.get(selected_child_id)
        if child:
            card.child_id = child.id
            _seed_defaults(card, child)
    return render_template("iom/form.html", **_form_context(card, is_new=True))


@iom_bp.route("/<int:card_id>")
@login_required
def card_view(card_id):
    _ensure_module_access()
    card = IomCard.query.options(
        joinedload(IomCard.child),
        joinedload(IomCard.history_entries).joinedload(IomHistory.created_by),
        joinedload(IomCard.schedule_lessons),
        joinedload(IomCard.schedule_import_sessions).joinedload(IomImportSessionSchedule.imported_by),
        joinedload(IomCard.schedule_corrections).joinedload(IomScheduleCorrection.specialist),
        joinedload(IomCard.section_rows),
        joinedload(IomCard.specialist_plans).joinedload(IomSpecialistPlan.specialist),
    ).get_or_404(card_id)
    if not _can_view_card(card):
        abort(403)
    readiness = _iom_readiness(card)
    return render_template(
        "iom/card.html",
        card=card,
        can_edit=_can_edit_card(card),
        can_delete=_can_delete_card(card),
        can_manage_corrections=_can_manage_corrections(card),
        readiness=readiness,
        **_schedule_context(card),
        **_section_context(card),
        **_monitoring_context(card),
        export_logs=card.export_logs[:10] if getattr(card, 'export_logs', None) else [],
        can_archive=_can_manage_archive(card),
        can_export=_can_export_card(card),
        years=AcademicYear.query.order_by(AcademicYear.name.desc()).all(),
    )


@iom_bp.route("/<int:card_id>/edit", methods=["GET", "POST"])
@login_required
def edit_card(card_id):
    _ensure_module_access()
    card = IomCard.query.get_or_404(card_id)
    if not _can_view_card(card):
        abort(403)
    if not _can_edit_card(card):
        abort(403)
    if request.method == "POST":
        try:
            _update_card_from_form(card)
            db.session.commit()
            flash("Карточка ИОМ сохранена.", "success")
            if request.form.get("save_and_continue"):
                return redirect(url_for("iom.edit_card", card_id=card.id))
            return redirect(url_for("iom.card_view", card_id=card.id))
        except ValueError as exc:
            db.session.rollback()
            flash(str(exc), "danger")
    return render_template("iom/form.html", **_form_context(card, is_new=False))


@iom_bp.route('/<int:card_id>/schedule/import', methods=['POST'])
@login_required
def import_schedule(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.schedule_lessons)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    file_storage = request.files.get('schedule_file')
    if not file_storage or not (file_storage.filename or '').strip():
        flash('Выберите файл расписания .xlsx.', 'danger')
        return redirect(url_for('iom.card_view', card_id=card.id))
    try:
        filename, rows = _extract_schedule_rows_from_workbook(file_storage)
        _replace_schedule_lessons(card, rows, filename=filename)
        db.session.commit()
        flash(f'Расписание загружено: {len(rows)} строк.', 'success')
    except ValueError as exc:
        db.session.rollback()
        flash(str(exc), 'danger')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/schedule/save', methods=['POST'])
@login_required
def save_schedule(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.schedule_lessons)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    updated = 0
    lesson_ids = request.form.getlist('lesson_id')
    weekdays = request.form.getlist('lesson_weekday')
    times = request.form.getlist('lesson_time')
    subjects = request.form.getlist('lesson_subject')

    rows_to_keep = []
    for lesson_id_raw, weekday_raw, time_raw, subject_raw in zip(lesson_ids, weekdays, times, subjects):
        lesson_id = _parse_int(lesson_id_raw)
        lesson = next((x for x in card.schedule_lessons if x.id == lesson_id), None)
        if not lesson:
            continue
        weekday_key = (weekday_raw or '').strip().lower()
        if weekday_key not in WEEKDAY_ORDER:
            continue
        time_text = _normalize_time_text(time_raw)
        subject_text = _safe_text(subject_raw)
        if not time_text or _is_placeholder_subject(subject_text):
            continue
        lesson.weekday = weekday_key
        lesson.start_time = time_text
        lesson.subject_name = subject_text
        lesson.sort_order = updated + 1
        rows_to_keep.append(lesson.id)
        updated += 1

    for lesson in list(card.schedule_lessons):
        if lesson.id not in rows_to_keep:
            db.session.delete(lesson)

    new_day = (request.form.get('new_weekday') or '').strip().lower()
    new_time = _normalize_time_text(request.form.get('new_time'))
    new_subject = _safe_text(request.form.get('new_subject'))
    if new_day and new_day in WEEKDAY_ORDER and new_time and not _is_placeholder_subject(new_subject):
        db.session.add(IomScheduleLesson(
            iom_card_id=card.id,
            weekday=new_day,
            start_time=new_time,
            subject_name=new_subject,
            source_type='MANUAL',
            sort_order=updated + 1,
        ))
        updated += 1

    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'schedule_update', 'Сетка школьного расписания обновлена вручную.')
    db.session.commit()
    flash('Сетка расписания сохранена.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/schedule/clear', methods=['POST'])
@login_required
def clear_schedule(card_id):
    _ensure_module_access()
    card = IomCard.query.get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    removed = IomScheduleLesson.query.filter_by(iom_card_id=card.id).delete()
    db.session.add(IomImportSessionSchedule(
        iom_card_id=card.id,
        filename='manual_clear',
        rows_loaded=0,
        imported_by_user_id=getattr(current_user, 'id', None),
        comment='Школьное расписание очищено вручную',
    ))
    _append_history(card, 'schedule_clear', f'Расписание очищено, удалено строк: {removed}.')
    db.session.commit()
    flash('Школьное расписание очищено.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/corrections/save', methods=['POST'])
@login_required
def save_corrections(card_id):
    _ensure_module_access()
    card = IomCard.query.options(
        joinedload(IomCard.schedule_lessons),
        joinedload(IomCard.schedule_corrections).joinedload(IomScheduleCorrection.specialist),
    ).get_or_404(card_id)
    if not _can_manage_corrections(card):
        abort(403)

    specialist_ids = request.form.getlist('correction_specialist_id')
    weekdays = request.form.getlist('correction_weekday')
    starts = request.form.getlist('correction_start_time')
    ends = request.form.getlist('correction_end_time')
    names = request.form.getlist('correction_course_name')
    notes_list = request.form.getlist('correction_notes')
    correction_ids = request.form.getlist('correction_id')

    allowed_specialist_ids = {str(item['specialist'].id) for item in _assigned_specialists(card)}
    user_spec = _find_specialist_for_user(getattr(current_user, 'id', None))
    rows_to_keep = []
    updated = 0

    for correction_id_raw, specialist_id_raw, weekday_raw, start_raw, end_raw, name_raw, note_raw in zip(correction_ids, specialist_ids, weekdays, starts, ends, names, notes_list):
        correction_id = _parse_int(correction_id_raw)
        row = next((x for x in card.schedule_corrections if x.id == correction_id), None)
        if not row:
            continue
        specialist_id = str(_parse_int(specialist_id_raw) or '')
        weekday = (weekday_raw or '').strip().lower()
        start_time = _normalize_time_text(start_raw)
        end_time = _normalize_time_text(end_raw)
        course_name = _safe_text(name_raw)
        notes = _safe_text(note_raw)
        if not specialist_id or specialist_id not in allowed_specialist_ids or weekday not in WEEKDAY_ORDER or not start_time or not end_time or not course_name:
            continue
        if user_spec and not (is_admin() or _is_sppiss_head()) and str(row.specialist_id) != str(user_spec.id):
            continue
        error_text = _overlap_exists(card, weekday, start_time, end_time, exclude_correction_id=row.id)
        if error_text:
            flash(error_text, 'danger')
            return redirect(url_for('iom.card_view', card_id=card.id))
        row.specialist_id = int(specialist_id)
        row.weekday = weekday
        row.start_time = start_time
        row.end_time = end_time
        row.course_name = course_name
        row.notes = notes
        row.updated_by_user_id = getattr(current_user, 'id', None)
        rows_to_keep.append(row.id)
        updated += 1

    for row in list(card.schedule_corrections):
        if row.id not in rows_to_keep:
            if user_spec and not (is_admin() or _is_sppiss_head()) and row.specialist_id != user_spec.id:
                continue
            db.session.delete(row)

    new_specialist_id = str(_parse_int(request.form.get('new_correction_specialist_id')) or '')
    new_weekday = (request.form.get('new_correction_weekday') or '').strip().lower()
    new_start = _normalize_time_text(request.form.get('new_correction_start_time'))
    new_end = _normalize_time_text(request.form.get('new_correction_end_time'))
    new_name = _safe_text(request.form.get('new_correction_course_name'))
    new_notes = _safe_text(request.form.get('new_correction_notes'))
    if new_specialist_id and new_weekday and new_start and new_end and new_name:
        if new_specialist_id not in allowed_specialist_ids:
            flash('Можно выбрать только специалиста, назначенного по сопровождению.', 'danger')
            return redirect(url_for('iom.card_view', card_id=card.id))
        if user_spec and not (is_admin() or _is_sppiss_head()) and new_specialist_id != str(user_spec.id):
            flash('Вы можете добавлять занятия только от своего имени.', 'danger')
            return redirect(url_for('iom.card_view', card_id=card.id))
        error_text = _overlap_exists(card, new_weekday, new_start, new_end)
        if error_text:
            flash(error_text, 'danger')
            return redirect(url_for('iom.card_view', card_id=card.id))
        db.session.add(IomScheduleCorrection(
            iom_card_id=card.id,
            specialist_id=int(new_specialist_id),
            weekday=new_weekday,
            start_time=new_start,
            end_time=new_end,
            course_name=new_name,
            notes=new_notes,
            created_by_user_id=getattr(current_user, 'id', None),
            updated_by_user_id=getattr(current_user, 'id', None),
        ))
        updated += 1

    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'correction_schedule_update', 'Сетка коррекционных занятий обновлена.')
    db.session.commit()
    flash('Коррекционные занятия сохранены.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/sections/analysis', methods=['POST'])
@login_required
def save_analysis_section(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.section_rows)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    _save_section(card, 'analysis', 'Анализ рекомендаций ЦПМПК / ИПРА', [
        'conclusion_number', 'conclusion_date', 'valid_until', 'pmpk_recommendations', 'ipra_recommendations',
        'special_conditions', 'required_conditions', 'implementation_mark', 'comment'
    ])
    db.session.commit()
    flash('Раздел с рекомендациями сохранен.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/sections/goals', methods=['POST'])
@login_required
def save_goals_section(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.section_rows)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    _save_section(card, 'goals', 'Цель и задачи ИОМ', [
        'main_goal', 'support_tasks', 'learning_tasks', 'correction_tasks', 'socialization_tasks', 'parent_interaction_forms', 'extra_comment', 'technologies'
    ])
    db.session.commit()
    flash('Раздел с целью и задачами сохранен.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/sections/ppk', methods=['POST'])
@login_required
def save_ppk_section(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.section_rows)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    _save_section(card, 'ppk', 'Коллегиальное заключение ППк', [
        'general_info', 'reason_text', 'ppk_conclusion', 'correction_directions', 'teacher_recommendations', 'parent_recommendations', 'individual_recommendations', 'repeat_review_term'
    ])
    db.session.commit()
    flash('Раздел ППк сохранен.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/sections/teacher', methods=['POST'])
@login_required
def save_teacher_section(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.section_rows)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    _save_section(card, 'teacher', 'Раздел классного руководителя / педагогов', [
        'program_features', 'meta_results', 'personal_results', 'uud_bud', 'difficulty_reasons', 'overcoming_forms', 'teacher_recommendations', 'support_measures', 'class_teacher_comment'
    ])
    db.session.commit()
    flash('Раздел классного руководителя сохранен.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/sections/parents', methods=['POST'])
@login_required
def save_parent_work_section(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.section_rows)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    directions = request.form.getlist('direction')
    results = request.form.getlist('planned_result')
    forms = request.form.getlist('work_form')
    frequencies = request.form.getlist('frequency')
    responsibles = request.form.getlist('responsible')
    comments = request.form.getlist('parent_comment')
    rows = []
    for direction, result, form_text, frequency, responsible, comment in zip(directions, results, forms, frequencies, responsibles, comments):
        packed = {
            'direction': (direction or '').strip(),
            'planned_result': (result or '').strip(),
            'work_form': (form_text or '').strip(),
            'frequency': (frequency or '').strip(),
            'responsible': (responsible or '').strip(),
            'comment': (comment or '').strip(),
        }
        if any(packed.values()):
            rows.append(packed)
    row = _section_row(card, 'parent_work', 'Работа с родителями')
    row.section_title = 'Работа с родителями'
    row.payload_json = json.dumps({'rows': rows}, ensure_ascii=False)
    row.updated_by_user_id = getattr(current_user, 'id', None)
    if not row.created_by_user_id:
        row.created_by_user_id = getattr(current_user, 'id', None)
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'section_parent_work', 'Раздел «Работа с родителями» сохранен.')
    db.session.commit()
    flash('Работа с родителями сохранена.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/specialists/sync', methods=['POST'])
@login_required
def sync_specialists(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.specialist_plans)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    created = _sync_specialist_plans_from_assignments(card)
    _append_history(card, 'specialists_sync', f'Из сопровождения подтянуто специалистов: {created}.')
    db.session.commit()
    flash(f'Из сопровождения подтянуто специалистов: {created}.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/specialists/save', methods=['POST'])
@login_required
def save_specialist_plans(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.specialist_plans)).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    plan_ids = request.form.getlist('plan_id')
    roles = request.form.getlist('plan_role_title')
    specialists = request.form.getlist('plan_specialist_id')
    recommendations = request.form.getlist('plan_recommendation_text')
    deficits = request.form.getlist('plan_deficits_text')
    resources = request.form.getlist('plan_resources_text')
    tasks = request.form.getlist('plan_tasks_text')
    work_forms = request.form.getlist('plan_work_form')
    sessions = request.form.getlist('plan_sessions_per_week')
    courses = request.form.getlist('plan_course_name')
    frequencies = request.form.getlist('plan_frequency')
    expecteds = request.form.getlist('plan_expected_result')
    monitorings = request.form.getlist('plan_monitoring_terms')
    comments = request.form.getlist('plan_comment')
    keep_ids = []
    for idx, values in enumerate(zip(plan_ids, roles, specialists, recommendations, deficits, resources, tasks, work_forms, sessions, courses, frequencies, expecteds, monitorings, comments), start=1):
        plan_id_raw, role_title, specialist_id_raw, recommendation_text, deficits_text, resources_text, tasks_text, work_form, sessions_per_week, course_name, frequency, expected_result, monitoring_terms, comment = values
        plan_id = _parse_int(plan_id_raw)
        row = next((x for x in card.specialist_plans if x.id == plan_id), None)
        if not row:
            continue
        role_title = (role_title or '').strip()
        specialist_id = _parse_int(specialist_id_raw)
        if not role_title and not specialist_id and not any((recommendation_text, deficits_text, resources_text, tasks_text, work_form, sessions_per_week, course_name, frequency, expected_result, monitoring_terms, comment)):
            continue
        row.role_title = role_title or row.role_title or 'Специалист'
        row.specialist_id = specialist_id
        row.recommendation_text = (recommendation_text or '').strip()
        row.deficits_text = (deficits_text or '').strip()
        row.resources_text = (resources_text or '').strip()
        row.tasks_text = (tasks_text or '').strip()
        row.work_form = (work_form or '').strip()
        row.sessions_per_week = (sessions_per_week or '').strip()
        row.course_name = (course_name or '').strip()
        row.frequency = (frequency or '').strip()
        row.expected_result = (expected_result or '').strip()
        row.monitoring_terms = (monitoring_terms or '').strip()
        row.comment = (comment or '').strip()
        row.sort_order = idx * 10
        row.updated_by_user_id = getattr(current_user, 'id', None)
        keep_ids.append(row.id)
    for row in list(card.specialist_plans):
        if row.id not in keep_ids:
            db.session.delete(row)
    new_role = (request.form.get('new_plan_role_title') or '').strip()
    new_specialist_id = _parse_int(request.form.get('new_plan_specialist_id'))
    new_course = (request.form.get('new_plan_course_name') or '').strip()
    if new_role or new_specialist_id or new_course:
        db.session.add(IomSpecialistPlan(
            iom_card_id=card.id,
            role_title=new_role or 'Специалист',
            specialist_id=new_specialist_id,
            course_name=new_course,
            work_form=(request.form.get('new_plan_work_form') or '').strip(),
            sessions_per_week=(request.form.get('new_plan_sessions_per_week') or '').strip(),
            frequency=(request.form.get('new_plan_frequency') or '').strip(),
            recommendation_text=(request.form.get('new_plan_recommendation_text') or '').strip(),
            deficits_text=(request.form.get('new_plan_deficits_text') or '').strip(),
            resources_text=(request.form.get('new_plan_resources_text') or '').strip(),
            tasks_text=(request.form.get('new_plan_tasks_text') or '').strip(),
            expected_result=(request.form.get('new_plan_expected_result') or '').strip(),
            monitoring_terms=(request.form.get('new_plan_monitoring_terms') or '').strip(),
            comment=(request.form.get('new_plan_comment') or '').strip(),
            sort_order=(len(keep_ids) + 1) * 10,
            created_by_user_id=getattr(current_user, 'id', None),
            updated_by_user_id=getattr(current_user, 'id', None),
        ))
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'specialists_save', 'Раздел специалистов ИОМ обновлен.')
    db.session.commit()
    flash('Карточки специалистов сохранены.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/cyclegram/sync', methods=['POST'])
@login_required
def sync_cyclegram(card_id):
    _ensure_module_access()
    card = IomCard.query.options(
        joinedload(IomCard.academic_year),
        joinedload(IomCard.schedule_corrections).joinedload(IomScheduleCorrection.iom_card),
    ).get_or_404(card_id)
    if not _can_edit_card(card):
        abort(403)
    created_count, updated_count, synced_specs = _sync_card_corrections_to_cyclegrams(card)
    if synced_specs:
        _append_history(card, 'cyclegram_sync', f'Синхронизировано с циклограммами: создано {created_count}, обновлено {updated_count}.')
        db.session.commit()
        flash(f'Синхронизировано с циклограммами: создано {created_count}, обновлено {updated_count}.', 'success')
    else:
        db.session.rollback()
        flash('Для назначенных специалистов не найдены циклограммы на этот учебный год.', 'warning')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/monitoring/save', methods=['POST'])
@login_required
def save_monitoring(card_id):
    _ensure_module_access()
    card = IomCard.query.get_or_404(card_id)
    if not _can_view_card(card):
        abort(403)
    period = (request.form.get('period') or '').strip().upper()
    block_code = (request.form.get('block_code') or '').strip()
    if period not in dict(MONITORING_PERIOD_CHOICES) or not block_code:
        flash('Некорректный блок мониторинга.', 'danger')
        return redirect(url_for('iom.card_view', card_id=card.id))
    templates = _monitoring_templates(card)
    if block_code not in templates.get(period, {}):
        flash('Блок мониторинга не найден для этого типа ИОМ.', 'danger')
        return redirect(url_for('iom.card_view', card_id=card.id))
    values, comments = {}, {}
    for row in templates[period][block_code]:
        values[row.line_code] = (request.form.get(f'value_{row.line_code}') or '').strip()
        comments[row.line_code] = (request.form.get(f'comment_{row.line_code}') or '').strip()
    payload = {'values': values, 'comments': comments, 'conclusion': (request.form.get('conclusion') or '').strip()}
    entry = IomMonitoringEntry.query.filter_by(iom_card_id=card.id, period=period, block_code=block_code).first()
    if not entry:
        entry = IomMonitoringEntry(iom_card_id=card.id, period=period, block_code=block_code)
        db.session.add(entry)
    entry.payload_json = json.dumps(payload, ensure_ascii=False)
    entry.updated_by_user_id = getattr(current_user, 'id', None)
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'monitoring_save', f'Сохранен мониторинг: {MONITORING_PERIOD_LABELS.get(period, period)} / {block_code}.')
    db.session.commit()
    flash('Мониторинг сохранен.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/export/docx')
@login_required
def export_docx(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.academic_year), joinedload(IomCard.schedule_lessons)).get_or_404(card_id)
    if not _can_export_card(card):
        abort(403)
    try:
        bio = _build_docx(card)
        _add_export_log(card, 'DOCX')
        db.session.commit()
        filename = f'iom_{card.id}.docx'
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as exc:
        db.session.rollback()
        flash(f'Не удалось сформировать DOCX: {exc}', 'danger')
        return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/export/pdf')
@login_required
def export_pdf(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.academic_year), joinedload(IomCard.schedule_lessons)).get_or_404(card_id)
    if not _can_export_card(card):
        abort(403)
    try:
        bio = _build_pdf(card)
        _add_export_log(card, 'PDF')
        db.session.commit()
        filename = f'iom_{card.id}.pdf'
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/pdf')
    except Exception as exc:
        db.session.rollback()
        flash(f'Не удалось сформировать PDF: {exc}', 'danger')
        return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/duplicate', methods=['POST'])
@login_required
def duplicate_card(card_id):
    _ensure_module_access()
    card = IomCard.query.options(joinedload(IomCard.section_rows), joinedload(IomCard.specialist_plans), joinedload(IomCard.schedule_lessons), joinedload(IomCard.schedule_corrections)).get_or_404(card_id)
    if not _can_manage_archive(card):
        abort(403)
    target_year_id = _parse_int(request.form.get('target_year_id'))
    copy_mode = (request.form.get('copy_mode') or 'STRUCTURE').strip().upper()
    if not target_year_id:
        flash('Выберите новый учебный год.', 'danger')
        return redirect(url_for('iom.card_view', card_id=card.id))
    new_card = _duplicate_card(card, target_year_id, 'WITH_SCHEDULE' if copy_mode == 'WITH_SCHEDULE' else 'STRUCTURE')
    _append_history(card, 'duplicate_source', f'Создана копия ИОМ на новый учебный год: #{new_card.id}.')
    db.session.commit()
    flash('Создана копия ИОМ на новый учебный год.', 'success')
    return redirect(url_for('iom.card_view', card_id=new_card.id))


@iom_bp.route('/<int:card_id>/archive', methods=['POST'])
@login_required
def archive_card(card_id):
    _ensure_module_access()
    card = IomCard.query.get_or_404(card_id)
    if not _can_manage_archive(card):
        abort(403)
    card.status = 'ARCHIVED'
    card.archived_at = datetime.utcnow()
    card.archived_by_user_id = getattr(current_user, 'id', None)
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'archive', 'Карточка переведена в архив.')
    db.session.commit()
    flash('Карточка переведена в архив.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route('/<int:card_id>/restore', methods=['POST'])
@login_required
def restore_card(card_id):
    _ensure_module_access()
    card = IomCard.query.get_or_404(card_id)
    if not _can_manage_archive(card):
        abort(403)
    card.status = 'DRAFT'
    card.archived_at = None
    card.archived_by_user_id = None
    card.updated_by_user_id = getattr(current_user, 'id', None)
    _append_history(card, 'restore', 'Карточка восстановлена из архива.')
    db.session.commit()
    flash('Карточка восстановлена из архива.', 'success')
    return redirect(url_for('iom.card_view', card_id=card.id))


@iom_bp.route("/<int:card_id>/delete", methods=["POST"])
@login_required
def delete_card(card_id):
    _ensure_module_access()
    card = IomCard.query.get_or_404(card_id)
    if not _can_delete_card(card):
        abort(403)
    _append_history(card, "delete", "Черновик удален")
    db.session.flush()
    db.session.delete(card)
    db.session.commit()
    flash("Черновик ИОМ удален.", "success")
    return redirect(url_for("iom.registry"))


def _form_context(card, is_new=False):
    selected_child = Child.query.get(card.child_id) if card.child_id else None
    readiness = _iom_readiness(card) if getattr(card, "id", None) else {
        "checks": [],
        "completed": 0,
        "total": 0,
        "percent": 0,
        "ready_for_on_approval": False,
        "ready_for_head": False,
        "ready_for_approved": False,
    }
    years = AcademicYear.query.order_by(AcademicYear.name.desc()).all()
    current_year = AcademicYear.query.filter_by(is_current=True).first()
    if is_new and not card.academic_year_id and current_year:
        card.academic_year_id = current_year.id
    children = (
        Child.query.filter_by(is_ovz=True)
        .order_by(Child.last_name.asc(), Child.first_name.asc(), Child.middle_name.asc())
        .all()
    )
    child_form_data = {}
    for child in children:
        level_code = _education_level_for_child(child)
        child_form_data[child.id] = {
            "education_level": level_code,
            "education_level_label": EDUCATION_LEVEL_LABELS.get(level_code, level_code or ""),
            "class_name": getattr(child.current_class, "name", "") if getattr(child, "current_class", None) else "",
            "birth_date": child.birth_date.strftime("%d.%m.%Y") if getattr(child, "birth_date", None) else "",
            "aop_variant": getattr(child, "ovz_variant", None) and f"Вариант {child.ovz_variant}" or (getattr(getattr(child, "social", None), "aoop_variant_text", None) or ""),
        }
    curators = User.query.order_by(User.last_name.asc(), User.first_name.asc()).all()
    classes = SchoolClass.query.order_by(SchoolClass.name.asc()).all()
    buildings = Building.query.order_by(Building.name.asc()).all()
    return {
        "card": card,
        "is_new": is_new,
        "children": children,
        "selected_child": selected_child,
        "years": years,
        "curators": curators,
        "classes": classes,
        "buildings": buildings,
        "iom_type_choices": IOM_TYPE_CHOICES,
        "education_level_choices": EDUCATION_LEVEL_CHOICES,
        "education_level_labels": EDUCATION_LEVEL_LABELS,
        "status_choices": IOM_STATUS_CHOICES,
        "aop_variant_choices": AOP_VARIANT_CHOICES,
        "aop_variant_choices_json": json.dumps(AOP_VARIANT_CHOICES, ensure_ascii=False),
        "child_form_data": child_form_data,
        "readiness": readiness,
    }
