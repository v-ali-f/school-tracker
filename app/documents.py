import os
import mimetypes
from datetime import datetime
from html import escape

from flask import Blueprint, redirect, url_for, request, current_app, abort, send_file, flash, render_template
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from openpyxl import load_workbook

from . import db
from .models import Child, Document, Debt, AcademicYear, User
from .retention import apply_retention_policies
from .permissions import can_view_documents, can_upload_documents, is_admin

documents_bp = Blueprint("documents", __name__)


def _upload_root() -> str:
    root = current_app.config.get("UPLOAD_FOLDER") or "app/uploads"
    return os.path.abspath(root)


def _abs_path(stored_path: str) -> str:
    if not stored_path:
        return ""
    if os.path.isabs(stored_path):
        return stored_path
    return os.path.join(_upload_root(), stored_path)


def _user_can_manage_document(child) -> bool:
    return is_admin() or can_upload_documents(child)


def _render_docx_preview(path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    parts = ['<div class="container-fluid py-3">']
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(f'<p style="margin-bottom:.5rem; white-space:pre-wrap;">{escape(text)}</p>')
    for table in doc.tables:
        parts.append('<div class="table-responsive"><table class="table table-sm table-bordered">')
        for row in table.rows:
            parts.append('<tr>')
            for cell in row.cells:
                parts.append(f'<td>{escape((cell.text or "").strip())}</td>')
            parts.append('</tr>')
        parts.append('</table></div>')
    if len(parts) == 1:
        parts.append('<div class="text-muted">В документе нет читаемого текста.</div>')
    parts.append('</div>')
    return ''.join(parts)


def _render_xlsx_preview(path: str) -> str:
    wb = load_workbook(path, data_only=True)
    parts = ['<div class="container-fluid py-3">']
    for ws in wb.worksheets:
        parts.append(f'<h6 class="mt-2">{escape(ws.title)}</h6>')
        parts.append('<div class="table-responsive mb-3"><table class="table table-sm table-bordered">')
        max_row = min(ws.max_row or 0, 50)
        max_col = min(ws.max_column or 0, 12)
        if max_row == 0 or max_col == 0:
            parts.append('<tr><td class="text-muted">Пустой лист</td></tr>')
        else:
            for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col, values_only=True):
                parts.append('<tr>')
                for val in row:
                    parts.append(f'<td>{escape("" if val is None else str(val))}</td>')
                parts.append('</tr>')
        parts.append('</table></div>')
        if (ws.max_row or 0) > 50 or (ws.max_column or 0) > 12:
            parts.append('<div class="small text-muted mb-3">Показана только часть таблицы для предпросмотра.</div>')
    parts.append('</div>')
    return ''.join(parts)


def _render_text_preview(path: str) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                text = f.read()
            break
        except Exception:
            text = None
    if text is None:
        text = "Не удалось прочитать текст документа."
    return f'<div class="container-fluid py-3"><pre style="white-space:pre-wrap;">{escape(text[:200000])}</pre></div>'


@documents_bp.route("/children/<int:child_id>/documents/upload", methods=["POST"])
@login_required
def upload_document(child_id: int):
    child = Child.query.get_or_404(child_id)
    if not _user_can_manage_document(child):
        abort(403)

    f = request.files.get("file")
    if not f or not f.filename:
        flash("Файл не выбран", "warning")
        return redirect(url_for("children.child_card", child_id=child_id))

    doc_type = (request.form.get("doc_type") or "GENERAL").strip().upper()
    allowed_types = {"GENERAL", "OVZ", "VSHU", "LOW", "AZ", "DISABLED", "MSE", "IPRA"}
    if doc_type not in allowed_types:
        doc_type = "GENERAL"

    debt_id_raw = (request.form.get("debt_id") or "").strip()
    debt_id = int(debt_id_raw) if debt_id_raw.isdigit() else None
    if doc_type != "AZ":
        debt_id = None
    elif debt_id is not None:
        ok = Debt.query.filter_by(id=debt_id, child_id=child_id).first()
        if not ok:
            debt_id = None

    doc_date_raw = (request.form.get("doc_date") or "").strip()
    doc_date = None
    if doc_date_raw:
        try:
            doc_date = datetime.strptime(doc_date_raw, "%Y-%m-%d").date()
        except ValueError:
            doc_date = None

    academic_year_id = request.form.get("academic_year_id", type=int)
    if not academic_year_id:
        current_year = AcademicYear.query.filter_by(is_current=True).first()
        academic_year_id = current_year.id if current_year else None
    retention_until = None
    if academic_year_id:
        year_obj = AcademicYear.query.get(academic_year_id)
        if year_obj and year_obj.end_date:
            try:
                retention_until = year_obj.end_date.replace(year=year_obj.end_date.year + 7)
            except Exception:
                retention_until = None

    safe_name = secure_filename(f.filename) or "file"
    child_rel_folder = str(child_id)
    child_abs_folder = os.path.join(_upload_root(), child_rel_folder)
    os.makedirs(child_abs_folder, exist_ok=True)

    stored_name = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{safe_name}"
    stored_rel_path = os.path.join(child_rel_folder, stored_name)
    stored_abs_path = os.path.join(_upload_root(), stored_rel_path)
    f.save(stored_abs_path)

    doc = Document(
        child_id=child_id,
        debt_id=debt_id,
        doc_type=doc_type,
        doc_date=doc_date,
        original_name=f.filename,
        stored_path=stored_rel_path,
        uploaded_by_user_id=current_user.id,
        uploaded_at=datetime.utcnow(),
        academic_year_id=academic_year_id,
        retention_until=retention_until,
    )
    db.session.add(doc)
    db.session.commit()

    flash("Документ загружен", "success")
    return redirect(url_for("children.child_card", child_id=child_id))


@documents_bp.route("/documents/<int:doc_id>/download")
@login_required
def download_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)
    if not (can_view_documents(child) or can_upload_documents(child) or is_admin()):
        abort(403)

    path = _abs_path(doc.stored_path)
    if not path or not os.path.isfile(path):
        abort(404)

    return send_file(path, as_attachment=True, download_name=(doc.original_name or os.path.basename(path)))


@documents_bp.route("/documents/<int:doc_id>/view")
@login_required
def view_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)
    if not (can_view_documents(child) or can_upload_documents(child) or is_admin()):
        abort(403)

    path = _abs_path(doc.stored_path)
    if not path or not os.path.isfile(path):
        abort(404)

    mime, _ = mimetypes.guess_type(path)
    mime = mime or "application/octet-stream"
    resp = send_file(path, mimetype=mime, as_attachment=False)
    resp.headers["Content-Disposition"] = "inline"
    resp.headers["Cache-Control"] = "no-store"
    return resp


@documents_bp.route("/documents/<int:doc_id>/preview")
@login_required
def preview_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)
    if not (can_view_documents(child) or can_upload_documents(child) or is_admin()):
        abort(403)

    path = _abs_path(doc.stored_path)
    if not path or not os.path.isfile(path):
        abort(404)

    ext = os.path.splitext(doc.original_name or path)[1].lower()
    html = None
    inline_url = None
    mode = "html"
    preview_error = None

    try:
        if ext in {'.pdf', '.png', '.jpg', '.jpeg', '.gif', '.webp', '.txt', '.csv'}:
            if ext in {'.txt', '.csv'}:
                html = _render_text_preview(path)
            else:
                mode = "iframe"
                inline_url = url_for('documents.view_document', doc_id=doc.id)
        elif ext == '.docx':
            html = _render_docx_preview(path)
        elif ext == '.xlsx':
            html = _render_xlsx_preview(path)
        else:
            preview_error = 'Для этого формата доступно скачивание. Полноценный просмотр в окне не поддерживается.'
    except Exception:
        preview_error = 'Не удалось построить предпросмотр документа.'

    return render_template('document_preview.html', doc=doc, mode=mode, inline_url=inline_url, preview_html=html, preview_error=preview_error)


@documents_bp.route("/documents/<int:doc_id>/delete", methods=["POST"])
@login_required
def delete_document(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    child = Child.query.get_or_404(doc.child_id)
    if not _user_can_manage_document(child):
        abort(403)

    doc.is_deleted_soft = True
    doc.deleted_at = datetime.utcnow()
    doc.deleted_by = getattr(current_user, "id", None)
    db.session.commit()

    flash("Документ скрыт из карточки. Файл сохранён на сервере в архиве.", "success")
    return redirect(url_for("children.child_card", child_id=doc.child_id))


@documents_bp.route("/documents/archive")
@login_required
def documents_archive():
    if not is_admin():
        abort(403)
    apply_retention_policies()
    q = (request.args.get("q") or "").strip().lower()
    mode = (request.args.get("mode") or "all").strip().lower()
    docs_query = Document.query
    if mode == "deleted":
        docs_query = docs_query.filter(Document.is_deleted_soft.is_(True))
    elif mode == "retention":
        docs_query = docs_query.filter(Document.is_hidden_by_retention.is_(True))
    else:
        docs_query = docs_query.filter(db.or_(Document.is_deleted_soft.is_(True), Document.is_hidden_by_retention.is_(True), Document.is_archived.is_(True)))
    docs = docs_query.order_by(Document.uploaded_at.desc()).all()
    if q:
        docs = [d for d in docs if q in (d.original_name or '').lower() or q in ((d.child.fio if d.child else '') or '').lower()]
    return render_template("documents_archive.html", docs=docs, mode=mode, q=q)


@documents_bp.route("/documents/retention/run", methods=["POST"])
@login_required
def run_retention():
    if not is_admin():
        abort(403)
    stats = apply_retention_policies()
    flash(
        f"Политика хранения обновлена. Скрыто документов: {stats['documents_hidden']}, архивировано нагрузок: {stats['teacher_load_archived']}, МЦКО: {stats['teacher_mcko_archived']}, курсов: {stats['teacher_courses_archived']}",
        "success"
    )
    return redirect(request.referrer or url_for("documents.documents_archive"))
